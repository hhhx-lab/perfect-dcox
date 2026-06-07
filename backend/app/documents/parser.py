from __future__ import annotations

from pathlib import Path
import re

from docx import Document


class DocumentParseError(RuntimeError):
    pass


def parse_docx(path: Path) -> dict[str, object]:
    try:
        document = Document(path)
    except Exception as exc:  # python-docx raises package-specific errors for corrupt zips.
        raise DocumentParseError(f"DOCX parse failed: {exc}") from exc

    paragraphs = list(document.paragraphs)
    paragraph_styles = [
        {
            "index": index,
            "text": paragraph.text,
            "style": paragraph.style.name if paragraph.style else None,
        }
        for index, paragraph in enumerate(paragraphs)
    ]
    heading_candidates = [
        paragraph.text
        for paragraph in paragraphs
        if paragraph.text.strip() and _is_heading_candidate(paragraph.text, paragraph.style.name if paragraph.style else "")
    ]
    image_count = len(document.part._element.xpath(".//*[local-name()='drawing' or local-name()='pict']"))
    all_text = "\n".join(paragraph.text for paragraph in paragraphs)
    return {
        "paragraph_count": len(paragraphs),
        "table_count": len(document.tables),
        "image_count": image_count,
        "heading_candidates": heading_candidates,
        "paragraph_styles": paragraph_styles,
        "all_text": all_text,
    }


def _is_heading_candidate(text: str, style_name: str) -> bool:
    stripped = text.strip()
    if style_name.lower().startswith("heading") or style_name.startswith("标题"):
        return True
    return bool(re.match(r"^(第[一二三四五六七八九十百0-9]+[章节]|[0-9]+(\.[0-9]+)*\s+)", stripped))
