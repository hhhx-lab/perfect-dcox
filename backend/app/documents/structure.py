from __future__ import annotations

from dataclasses import dataclass
from enum import Enum
import re

from docx.document import Document as DocxDocument


class ParagraphRole(str, Enum):
    DOCUMENT_TITLE = "document_title"
    COVER_OR_METADATA = "cover_or_metadata"
    TOC_TITLE = "toc_title"
    TOC_ITEM = "toc_item"
    ABSTRACT_HEADING = "abstract_heading"
    ABSTRACT_BODY = "abstract_body"
    KEYWORDS = "keywords"
    HEADING = "heading"
    BODY = "body"
    TABLE_CAPTION = "table_caption"
    FIGURE_CAPTION = "figure_caption"
    EQUATION = "equation"
    REFERENCE_HEADING = "reference_heading"
    REFERENCE_ITEM = "reference_item"
    ACKNOWLEDGEMENT_HEADING = "acknowledgement_heading"
    APPENDIX_HEADING = "appendix_heading"
    APPENDIX_BODY = "appendix_body"
    EMPTY = "empty"


@dataclass(frozen=True)
class ParagraphClassification:
    role: ParagraphRole
    heading_level: int | None = None


class DocumentStructure:
    def __init__(self, classifications: list[ParagraphClassification]) -> None:
        self._classifications = classifications

    def role_for(self, paragraph_index: int) -> ParagraphClassification:
        return self._classifications[paragraph_index]


def classify_document(document: DocxDocument) -> DocumentStructure:
    texts = [paragraph.text.strip() for paragraph in document.paragraphs]
    style_names = [paragraph.style.name if paragraph.style else "" for paragraph in document.paragraphs]
    toc_indices = _detect_toc_indices(texts)
    reference_heading = _last_index(
        texts,
        lambda index, text: index not in toc_indices and text in {"参考文献", "References"},
    )
    acknowledgement_heading = _last_index(
        texts,
        lambda index, text: index not in toc_indices and text in {"致谢", "Acknowledgements", "Acknowledgments"},
    )
    abstract_ranges = _detect_abstract_ranges(texts, toc_indices, reference_heading)
    first_nonempty = next((index for index, text in enumerate(texts) if text), None)

    classifications: list[ParagraphClassification] = []
    inside_appendix = False
    for index, (text, style_name) in enumerate(zip(texts, style_names)):
        if not text:
            classifications.append(ParagraphClassification(ParagraphRole.EMPTY))
        elif _paragraph_has_toc_field(document.paragraphs[index]):
            classifications.append(ParagraphClassification(ParagraphRole.TOC_TITLE))
        elif text in {"目录", "Contents"}:
            classifications.append(ParagraphClassification(ParagraphRole.TOC_TITLE))
        elif first_nonempty == index and style_name.lower().startswith("heading"):
            if _looks_like_document_title(text):
                classifications.append(ParagraphClassification(ParagraphRole.DOCUMENT_TITLE, 1))
            else:
                classifications.append(ParagraphClassification(ParagraphRole.HEADING, _heading_level(text, style_name) or 1))
        elif _is_appendix_heading(text):
            inside_appendix = True
            classifications.append(ParagraphClassification(ParagraphRole.APPENDIX_HEADING, 2))
        elif inside_appendix and _is_appendix_terminator(text):
            inside_appendix = False
            classifications.append(ParagraphClassification(ParagraphRole.HEADING, _heading_level(text, style_name) or 2))
        elif index in toc_indices:
            classifications.append(ParagraphClassification(ParagraphRole.TOC_ITEM))
        elif index == reference_heading:
            classifications.append(ParagraphClassification(ParagraphRole.REFERENCE_HEADING, 2))
        elif index == acknowledgement_heading:
            classifications.append(ParagraphClassification(ParagraphRole.ACKNOWLEDGEMENT_HEADING, 2))
        elif index in abstract_ranges["headings"]:
            classifications.append(ParagraphClassification(ParagraphRole.ABSTRACT_HEADING))
        elif index in abstract_ranges["keywords"]:
            classifications.append(ParagraphClassification(ParagraphRole.KEYWORDS))
        elif index in abstract_ranges["bodies"]:
            classifications.append(ParagraphClassification(ParagraphRole.ABSTRACT_BODY))
        elif _is_reference_item(index, text, reference_heading, acknowledgement_heading):
            classifications.append(ParagraphClassification(ParagraphRole.REFERENCE_ITEM))
        elif _is_table_caption(text):
            classifications.append(ParagraphClassification(ParagraphRole.TABLE_CAPTION))
        elif _is_figure_caption(text):
            classifications.append(ParagraphClassification(ParagraphRole.FIGURE_CAPTION))
        elif _is_equation_paragraph(document.paragraphs[index], text):
            classifications.append(ParagraphClassification(ParagraphRole.EQUATION))
        elif (heading_level := _heading_level(text, style_name)) is not None:
            classifications.append(ParagraphClassification(ParagraphRole.HEADING, heading_level))
        elif inside_appendix:
            classifications.append(ParagraphClassification(ParagraphRole.APPENDIX_BODY))
        elif first_nonempty is not None and index < first_nonempty + 7 and "：" in text:
            classifications.append(ParagraphClassification(ParagraphRole.COVER_OR_METADATA))
        else:
            classifications.append(ParagraphClassification(ParagraphRole.BODY))

    return DocumentStructure(classifications)


def _detect_toc_indices(texts: list[str]) -> set[int]:
    toc_start = next((index for index, text in enumerate(texts) if text in {"目录", "Contents"}), None)
    if toc_start is None:
        return set()
    abstract_indices = [index for index, text in enumerate(texts) if text == "摘要" and index > toc_start]
    if len(abstract_indices) >= 2:
        return set(range(toc_start + 1, abstract_indices[1]))
    return set()


def _paragraph_has_toc_field(paragraph) -> bool:
    xml = paragraph._p.xml.upper()
    return "TOC" in xml and ("FLDCHAR" in xml or "FLDSIMPLE" in xml)


def _detect_abstract_ranges(
    texts: list[str],
    toc_indices: set[int],
    reference_heading: int | None,
) -> dict[str, set[int]]:
    headings: set[int] = set()
    bodies: set[int] = set()
    keywords: set[int] = set()
    search_end = reference_heading if reference_heading is not None else len(texts)

    zh_heading = next((index for index, text in enumerate(texts[:search_end]) if index not in toc_indices and text == "摘要"), None)
    if zh_heading is not None:
        headings.add(zh_heading)
        zh_keyword = next(
            (index for index in range(zh_heading + 1, search_end) if texts[index].startswith("关键词")),
            None,
        )
        if zh_keyword is not None:
            keywords.add(zh_keyword)
            bodies.update(index for index in range(zh_heading + 1, zh_keyword) if texts[index])

    en_heading = next(
        (
            index
            for index, text in enumerate(texts[:search_end])
            if index not in toc_indices and text.lower() == "abstract" and (zh_heading is None or index > zh_heading)
        ),
        None,
    )
    if en_heading is not None:
        headings.add(en_heading)
        en_keyword = next(
            (index for index in range(en_heading + 1, search_end) if texts[index].lower().startswith("keywords")),
            None,
        )
        if en_keyword is not None:
            keywords.add(en_keyword)
            bodies.update(index for index in range(en_heading + 1, en_keyword) if texts[index])

    return {"headings": headings, "bodies": bodies, "keywords": keywords}


def _last_index(texts: list[str], predicate) -> int | None:
    for index in range(len(texts) - 1, -1, -1):
        if predicate(index, texts[index]):
            return index
    return None


def _is_reference_item(
    index: int,
    text: str,
    reference_heading: int | None,
    acknowledgement_heading: int | None,
) -> bool:
    if reference_heading is None or index <= reference_heading:
        return False
    if acknowledgement_heading is not None and index >= acknowledgement_heading:
        return False
    return bool(re.match(r"^(\[\d+\]|［\d+］)", text))


def _heading_level(text: str, style_name: str) -> int | None:
    if re.match(r"^[0-9]+\.[0-9]+(?:\.[0-9]+)*\s+", text):
        return 2
    if re.match(r"^[0-9]+、\s*", text):
        return 2
    if re.match(r"^第[一二三四五六七八九十百0-9]+[章节]", text):
        return 1
    if style_name.lower().startswith("heading"):
        match = re.search(r"(\d+)", style_name)
        return int(match.group(1)) if match else 1
    return None


def _is_appendix_heading(text: str) -> bool:
    return bool(re.match(r"^\s*(附录(?:\s*[A-ZＡ-Ｚ一二三四五六七八九十0-9]+)?|Appendix\b)", text.strip(), re.IGNORECASE))


def _is_appendix_terminator(text: str) -> bool:
    return text.strip() in {"参考文献", "References", "致谢", "Acknowledgements", "Acknowledgments"}


def _looks_like_document_title(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if _heading_level(stripped, "") is not None:
        return False
    if len(stripped) <= 8:
        return False
    if stripped.lower() in {
        "abstract",
        "acknowledgements",
        "acknowledgments",
        "references",
        "introduction",
        "preface",
    }:
        return False
    return True


def _is_table_caption(text: str) -> bool:
    if re.search(r"(给出|如下|所示|见|列出|显示|说明|展示|概括)", text):
        return False
    return _caption_number_present(text, prefixes=("表", "Table"))


def _is_figure_caption(text: str) -> bool:
    if re.search(r"(给出|如下|所示|见|列出|显示|说明|展示|概括)", text):
        return False
    return _caption_number_present(text, prefixes=("图", "Figure"))


def _caption_number_present(text: str, *, prefixes: tuple[str, str]) -> bool:
    escaped = "|".join(re.escape(prefix) for prefix in prefixes)
    return bool(
        re.match(
            rf"^(?:{escaped})\s*[\s：:、.\-/]*\s*\d+(?:[\s：:、.\-/]+\S+|$)",
            text.strip(),
            re.IGNORECASE,
        )
    )


def _is_equation_paragraph(paragraph, text: str) -> bool:
    if paragraph._p.xpath(".//*[local-name()='oMath' or local-name()='oMathPara']"):
        return True
    if not text or text.startswith(("http://", "https://", "[", "［")):
        return False
    if "$$" in text or re.search(r"\\(?:frac|sum|int|sqrt|begin|end)", text):
        return True
    if "=" not in text or len(text) > 80:
        return False
    if re.search(r"[\u4e00-\u9fff：；，。]", text):
        return False
    return bool(re.match(r"^[A-Za-z0-9_{}\\^()+\-*/= .,\t]+$", text))
