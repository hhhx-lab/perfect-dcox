from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn


class TemplateBindingError(RuntimeError):
    pass


class TemplateLoader:
    def compose_with_body_slot(
        self,
        template_path: Path,
        source_docx_path: Path,
        output_path: Path,
        body_slot: str | None = "{{BODY}}",
        placeholder_policy: str = "fail",
    ) -> Path:
        slot = body_slot or "{{BODY}}"
        try:
            template_doc = Document(template_path)
            source_doc = Document(source_docx_path)
        except Exception as exc:
            raise TemplateBindingError(f"Template binding failed to open DOCX: {exc}") from exc

        marker = next((paragraph for paragraph in template_doc.paragraphs if slot in paragraph.text), None)
        if marker is None:
            raise TemplateBindingError(f"Template body slot marker not found: {slot}")

        parent = marker._element.getparent()
        insert_at = parent.index(marker._element)
        parent.remove(marker._element)
        inserted = 0
        for element in source_doc.element.body:
            if element.tag == qn("w:sectPr"):
                continue
            parent.insert(insert_at + inserted, deepcopy(element))
            inserted += 1
        if inserted == 0:
            raise TemplateBindingError("Source document has no body content to place into template slot.")
        _handle_remaining_placeholders(template_doc, slot, placeholder_policy)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        template_doc.save(output_path)
        return output_path


def _handle_remaining_placeholders(document, consumed_slot: str, placeholder_policy: str) -> None:
    remaining: list[str] = []
    placeholder_pattern = re.compile(r"\{\{[^{}]+\}\}")
    for paragraph in document.paragraphs:
        matches = [match.group(0) for match in placeholder_pattern.finditer(paragraph.text)]
        matches = [match for match in matches if match != consumed_slot]
        if not matches:
            continue
        if placeholder_policy == "preserve":
            remaining.extend(matches)
            continue
        if placeholder_policy == "remove":
            for run in paragraph.runs:
                run.text = placeholder_pattern.sub("", run.text)
            continue
        remaining.extend(matches)
    if remaining and placeholder_policy == "fail":
        unique = ", ".join(sorted(set(remaining)))
        raise TemplateBindingError(f"Template contains unresolved placeholder(s): {unique}")
