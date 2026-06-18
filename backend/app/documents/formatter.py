from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZipFile
import re
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Mm, Pt, RGBColor

from app.profiles.models import HeadingSettings, TextAlignment, TextFont
from app.profiles.models import FormatProfile
from app.documents.rule_registry import summarize_docx_formatter_dispatch
from app.documents.structure import ParagraphRole, _heading_level, classify_document
from app.documents.ooxml import NS, enable_update_fields


class DocumentFormatError(RuntimeError):
    pass


FORMATTER_PIPELINE_APPLIERS: frozenset[str] = frozenset(
    {
        "_apply_page_settings",
        "_apply_section_size",
        "_apply_document_grid",
        "_apply_basic_page_numbers",
        "_apply_footer_page_number",
        "_ensure_toc",
        "_ensure_table_captions",
        "_ensure_bilingual_caption_near",
        "_ensure_figure_captions",
        "_apply_figure_size_rules",
        "_normalize_body_text",
        "_apply_body_paragraph",
        "_apply_heading_paragraph",
        "_apply_heading_numbering",
        "_apply_abstract_body",
        "_apply_table_rules",
        "_apply_equation_paragraph",
        "_apply_reference_paragraph",
        "_apply_notes",
        "_apply_appendix_heading",
        "_apply_appendix_body",
    }
)


def formatter_pipeline_applier_names() -> list[str]:
    return sorted(FORMATTER_PIPELINE_APPLIERS)


@dataclass
class FormatterExecutionTrace:
    call_counts: dict[str, int] = field(default_factory=dict)

    def record(self, applier: str, count: int = 1) -> None:
        self.call_counts[applier] = self.call_counts.get(applier, 0) + count

    def public_summary(self) -> dict[str, object]:
        return summarize_docx_formatter_dispatch(self.call_counts).public_summary()


def _record_trace(trace: FormatterExecutionTrace | None, applier: str, count: int = 1) -> None:
    if trace is not None:
        trace.record(applier, count)


def format_docx_with_profile(
    input_path: Path,
    output_path: Path,
    profile: FormatProfile,
    *,
    preserve_header_footer: bool = False,
    trace: FormatterExecutionTrace | None = None,
) -> Path:
    try:
        document = Document(input_path)
    except Exception as exc:
        raise DocumentFormatError(f"DOCX formatting failed to open input: {exc}") from exc

    _record_trace(trace, "_apply_page_settings")
    _record_trace(trace, "_apply_section_size", len(document.sections))
    _record_trace(trace, "_apply_document_grid", len(document.sections))
    _apply_page_settings(document, profile, preserve_page_numbering=preserve_header_footer)
    if not preserve_header_footer:
        _record_trace(trace, "_apply_basic_page_numbers")
        if profile.header_footer.footer_page_number and profile.header_footer.page_number_format != "none":
            _record_trace(trace, "_apply_footer_page_number", len(document.sections))
        _apply_basic_page_numbers(document, profile)
    _record_trace(trace, "_ensure_toc")
    _ensure_toc(document, profile)
    _record_trace(trace, "_ensure_table_captions")
    _ensure_table_captions(document, profile, trace=trace)
    _record_trace(trace, "_ensure_figure_captions")
    _ensure_figure_captions(document, profile, trace=trace)
    _record_trace(trace, "_apply_figure_size_rules")
    _apply_figure_size_rules(document, profile)
    heading_counters = [0] * 9
    equation_counter = 0
    inside_appendix = False
    structure = classify_document(document)
    for index, paragraph in enumerate(document.paragraphs):
        if _paragraph_has_toc_field(paragraph):
            continue
        classification = structure.role_for(index)
        if classification.role == ParagraphRole.EMPTY:
            continue
        if _is_appendix_heading_text(paragraph.text):
            inside_appendix = True
            _record_trace(trace, "_apply_appendix_heading")
            _apply_appendix_heading(paragraph, profile)
            continue
        if inside_appendix and _is_appendix_terminator(paragraph.text):
            inside_appendix = False
        if classification.role == ParagraphRole.TOC_TITLE:
            _record_trace(trace, "_apply_heading_paragraph")
            _apply_heading_paragraph(paragraph, profile, 2)
        elif classification.role == ParagraphRole.TOC_ITEM:
            continue
        elif classification.role == ParagraphRole.DOCUMENT_TITLE:
            _record_trace(trace, "_apply_heading_paragraph")
            _apply_heading_paragraph(paragraph, profile, 1)
        elif classification.role in {
            ParagraphRole.HEADING,
            ParagraphRole.REFERENCE_HEADING,
            ParagraphRole.ACKNOWLEDGEMENT_HEADING,
        }:
            if classification.role == ParagraphRole.HEADING:
                _record_trace(trace, "_normalize_body_text")
                _normalize_body_text(paragraph, profile, normalize_list_marker=False)
            _record_trace(trace, "_apply_heading_paragraph")
            _apply_heading_paragraph(paragraph, profile, classification.heading_level or 1)
            if classification.role == ParagraphRole.HEADING:
                _record_trace(trace, "_apply_heading_numbering")
                _apply_heading_numbering(paragraph, profile, classification.heading_level or 1, heading_counters)
        elif classification.role == ParagraphRole.ABSTRACT_HEADING:
            _record_trace(trace, "_normalize_body_text")
            _normalize_body_text(paragraph, profile)
            _apply_abstract_heading(paragraph, profile)
        elif classification.role == ParagraphRole.ABSTRACT_BODY:
            _record_trace(trace, "_normalize_body_text")
            _normalize_body_text(paragraph, profile)
            _record_trace(trace, "_apply_abstract_body")
            _apply_abstract_body(paragraph, profile)
        elif classification.role == ParagraphRole.KEYWORDS:
            _record_trace(trace, "_normalize_body_text")
            _normalize_body_text(paragraph, profile)
            _apply_keywords_paragraph(paragraph, profile)
        elif classification.role == ParagraphRole.TABLE_CAPTION:
            _normalize_caption_paragraph(paragraph, profile, "table")
            _apply_caption_paragraph(paragraph, profile.table.caption.font, profile.document_grid.snap_to_grid)
        elif classification.role == ParagraphRole.FIGURE_CAPTION:
            _normalize_caption_paragraph(paragraph, profile, "figure")
            _apply_caption_paragraph(paragraph, profile.figure.caption.font, profile.document_grid.snap_to_grid)
        elif classification.role == ParagraphRole.EQUATION:
            equation_counter += 1
            _record_trace(trace, "_apply_equation_paragraph")
            _apply_equation_paragraph(paragraph, profile, equation_counter)
        elif classification.role == ParagraphRole.REFERENCE_ITEM:
            _record_trace(trace, "_normalize_body_text")
            _normalize_body_text(paragraph, profile)
            _record_trace(trace, "_apply_reference_paragraph")
            _apply_reference_paragraph(paragraph, profile)
        elif inside_appendix and classification.role in {
            ParagraphRole.APPENDIX_BODY,
            ParagraphRole.BODY,
            ParagraphRole.COVER_OR_METADATA,
        }:
            _record_trace(trace, "_normalize_body_text")
            _normalize_body_text(paragraph, profile)
            _record_trace(trace, "_apply_appendix_body")
            _apply_appendix_body(paragraph, profile)
        else:
            _record_trace(trace, "_normalize_body_text")
            _normalize_body_text(paragraph, profile)
            _record_trace(trace, "_apply_body_paragraph")
            _apply_body_paragraph(paragraph, profile)
    for table in document.tables:
        _record_trace(trace, "_apply_table_rules")
        _apply_table_rules(table, profile)
        _record_trace(trace, "_normalize_body_text")
        _normalize_table_text(table, profile)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    _record_trace(trace, "_apply_notes")
    _apply_notes(output_path, profile)
    enable_update_fields(
        output_path,
        enabled=profile.toc.update_fields_on_open or profile.header_footer.footer_page_number,
        even_and_odd_headers=profile.header_footer.different_odd_even,
    )
    return output_path


def _apply_page_settings(document: Document, profile: FormatProfile, *, preserve_page_numbering: bool = False) -> None:
    for section in document.sections:
        _apply_section_size(section, profile)
        section.top_margin = Cm(profile.page.margins_cm.top)
        section.bottom_margin = Cm(profile.page.margins_cm.bottom)
        section.left_margin = Cm(profile.page.margins_cm.left)
        section.right_margin = Cm(profile.page.margins_cm.right)
        section.gutter = Cm(profile.page.margins_cm.gutter)
        if not preserve_page_numbering:
            section.different_first_page_header_footer = profile.header_footer.different_first_page
            _apply_page_number_type(section, profile)
        _apply_document_grid(section, profile)


def _apply_section_size(section, profile: FormatProfile) -> None:
    if profile.page.size == "A4":
        width_cm, height_cm = 21.0, 29.7
    else:
        width_cm, height_cm = 21.59, 27.94
    if profile.page.orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(height_cm)
        section.page_height = Cm(width_cm)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(width_cm)
        section.page_height = Cm(height_cm)


def _apply_basic_page_numbers(document: Document, profile: FormatProfile) -> None:
    for section in document.sections:
        _apply_header_footer_container(section.header, profile.header_footer.header_text, profile.header_footer.header_alignment, profile)
        if profile.header_footer.different_first_page:
            _apply_header_footer_container(section.first_page_header, profile.header_footer.header_text, profile.header_footer.header_alignment, profile)
        if profile.header_footer.different_odd_even:
            _apply_header_footer_container(section.even_page_header, profile.header_footer.header_text, profile.header_footer.header_alignment, profile)
        if not profile.header_footer.footer_page_number or profile.header_footer.page_number_format == "none":
            _remove_supported_page_number(section.footer)
            if profile.header_footer.different_first_page:
                _remove_supported_page_number(section.first_page_footer)
            if profile.header_footer.different_odd_even:
                _remove_supported_page_number(section.even_page_footer)
            continue
        _apply_footer_page_number(section.footer, profile)
        if profile.header_footer.different_first_page:
            _apply_footer_page_number(section.first_page_footer, profile)
        if profile.header_footer.different_odd_even:
            _apply_footer_page_number(section.even_page_footer, profile)


def _remove_supported_page_number(footer) -> None:
    for paragraph in footer.paragraphs:
        if "PAGE" in paragraph._p.xml and "fldChar" in paragraph._p.xml:
            paragraph.text = ""


def _apply_header_footer_container(container, text: str | None, alignment: TextAlignment, profile: FormatProfile) -> None:
    resolved = (text or "").strip()
    if not resolved:
        return
    paragraph = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
    paragraph.text = resolved
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = None
    _apply_paragraph_alignment(paragraph, alignment)
    _apply_runs_font(paragraph, profile.header_footer.font)


def _apply_footer_page_number(footer, profile: FormatProfile) -> None:
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.text = ""
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = None
    _apply_paragraph_alignment(paragraph, profile.header_footer.footer_alignment)
    footer_text = (profile.header_footer.footer_text or "").strip()
    if footer_text:
        paragraph.add_run(f"{footer_text} ")
    run = paragraph.add_run()
    _append_page_field(run, profile)
    _apply_runs_font(paragraph, profile.header_footer.font)


def _append_page_field(run, profile: FormatProfile) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGE {_page_format_switch(profile.header_footer.page_number_format)} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = _page_number_preview(profile)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def _page_format_switch(format_name: str) -> str:
    return {
        "roman_lower": r"\* roman",
        "roman_upper": r"\* ROMAN",
        "arabic": r"\* Arabic",
    }.get(format_name, "")


def _page_number_preview(profile: FormatProfile) -> str:
    start = profile.header_footer.page_number_start or 1
    if profile.header_footer.page_number_format == "roman_lower":
        return _roman(start).lower()
    if profile.header_footer.page_number_format == "roman_upper":
        return _roman(start).upper()
    return str(start)


def _roman(value: int) -> str:
    numerals = [
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    ]
    remaining = max(1, value)
    result = ""
    for amount, token in numerals:
        while remaining >= amount:
            result += token
            remaining -= amount
    return result


def _apply_page_number_type(section, profile: FormatProfile) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(profile.header_footer.page_number_start))
    fmt = _page_number_ooxml_format(profile.header_footer.page_number_format)
    if fmt:
        pg_num.set(qn("w:fmt"), fmt)
    elif qn("w:fmt") in pg_num.attrib:
        del pg_num.attrib[qn("w:fmt")]


def _page_number_ooxml_format(format_name: str) -> str | None:
    return {
        "arabic": "decimal",
        "roman_lower": "lowerRoman",
        "roman_upper": "upperRoman",
    }.get(format_name)


def _apply_document_grid(section, profile: FormatProfile) -> None:
    sect_pr = section._sectPr
    doc_grid = sect_pr.find(qn("w:docGrid"))
    if doc_grid is None:
        doc_grid = OxmlElement("w:docGrid")
        sect_pr.append(doc_grid)
    if not profile.document_grid.enabled or profile.document_grid.type == "none":
        doc_grid.set(qn("w:type"), "default")
        for attr in (qn("w:linePitch"), qn("w:charSpace")):
            if attr in doc_grid.attrib:
                del doc_grid.attrib[attr]
        return
    doc_grid.set(qn("w:type"), "lines" if profile.document_grid.type == "line" else "linesAndChars")
    if profile.document_grid.lines_per_page:
        usable_height_cm = section.page_height.cm - section.top_margin.cm - section.bottom_margin.cm
        line_pitch = max(120, round((usable_height_cm / profile.document_grid.lines_per_page) * 567))
        doc_grid.set(qn("w:linePitch"), str(line_pitch))
    if profile.document_grid.characters_per_line:
        usable_width_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
        char_space = max(20, round((usable_width_cm / profile.document_grid.characters_per_line) * 567))
        doc_grid.set(qn("w:charSpace"), str(char_space))


def _ensure_toc(document: Document, profile: FormatProfile) -> None:
    if not profile.toc.enabled:
        return
    paragraphs = list(document.paragraphs)
    texts = [paragraph.text.strip() for paragraph in paragraphs]
    toc_titles = {profile.toc.title, "目录", "Contents"}
    toc_index = next((index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip() in toc_titles), None)
    if toc_index is None:
        if _should_generate_missing_toc(profile):
            _insert_generated_toc(document, profile)
        return
    real_abstract_index = _first_real_abstract_heading_index(paragraphs, texts, toc_index)
    if real_abstract_index is not None:
        end_index = real_abstract_index
    else:
        end_index = toc_index + 1
        while end_index < len(paragraphs):
            text = paragraphs[end_index].text.strip()
            if not text:
                end_index += 1
                continue
            if _heading_level(text, paragraphs[end_index].style.name if paragraphs[end_index].style else ""):
                break
            end_index += 1
    if end_index <= toc_index + 1:
        return

    toc_entries = [paragraph.text.strip() for paragraph in paragraphs[toc_index + 1 : end_index] if paragraph.text.strip()]
    toc_paragraph = paragraphs[toc_index]
    _clear_paragraph_runs(toc_paragraph)
    toc_paragraph.text = profile.toc.title
    field_paragraph = toc_paragraph.insert_paragraph_before("")
    toc_paragraph._element.addnext(field_paragraph._element)
    _append_toc_field(field_paragraph.add_run(), toc_entries, profile)
    _apply_paragraph_alignment(toc_paragraph, "left")
    _apply_paragraph_alignment(field_paragraph, "left")
    for paragraph in paragraphs[toc_index + 1 : end_index]:
        paragraph._element.getparent().remove(paragraph._element)


def _first_real_abstract_heading_index(paragraphs: list, texts: list[str], toc_index: int) -> int | None:
    for index, text in enumerate(texts):
        if index <= toc_index or text not in {"摘要", "Abstract"}:
            continue
        style_name = paragraphs[index].style.name if paragraphs[index].style else ""
        following = [item for item in texts[index + 1 : index + 5] if item]
        next_text = following[0] if following else ""
        has_nearby_keywords = any(item.startswith(("关键词", "Keywords:")) for item in following)
        if _looks_like_toc_entry(next_text):
            continue
        if style_name.lower().startswith("heading") or has_nearby_keywords or len(next_text) > 0:
            return index
    return None


def _looks_like_toc_entry(text: str) -> bool:
    if text in {"摘要", "Abstract", "参考文献", "References", "致谢", "Acknowledgements", "Acknowledgments"}:
        return True
    return _heading_level(text, "") is not None


def _insert_generated_toc(document: Document, profile: FormatProfile) -> None:
    paragraphs = list(document.paragraphs)
    anchor = next((paragraph for paragraph in paragraphs if paragraph.text.strip()), None)
    if anchor is None:
        return
    title = anchor.insert_paragraph_before(profile.toc.title)
    anchor._element.addprevious(title._element)
    field = title.insert_paragraph_before("")
    title._element.addnext(field._element)
    _append_toc_field(field.add_run(), _heading_entries(document, profile), profile)
    _apply_paragraph_alignment(title, "left")
    _apply_paragraph_alignment(field, "left")


def _should_generate_missing_toc(profile: FormatProfile) -> bool:
    if not profile.toc.enabled:
        return False
    if profile.schema_version != "1.0.0":
        return True
    return any(section.key.lower() == "toc" and section.required for section in profile.sections)


def _heading_entries(document: Document, profile: FormatProfile) -> list[str]:
    entries: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text or text == profile.toc.title:
            continue
        level = _heading_level(text, paragraph.style.name if paragraph.style else "")
        if level is not None and level <= profile.toc.include_levels:
            entries.append(text)
    return entries


def _append_toc_field(run, entries: list[str], profile: FormatProfile) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    switches = [f'\\o "1-{profile.toc.include_levels}"', r"\u"]
    if profile.toc.use_hyperlinks:
        switches.append(r"\h")
    switches.append(r"\z")
    if not profile.toc.show_page_numbers:
        switches.append(r"\n")
    elif not profile.toc.right_align_page_numbers:
        switches.append(r'\p " "')
    instr.text = " TOC " + " ".join(switches) + " "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "\n".join(entries) if entries else "目录"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def _clear_paragraph_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _paragraph_has_toc_field(paragraph) -> bool:
    xml = paragraph._p.xml.upper()
    return "TOC" in xml and ("FLDCHAR" in xml or "FLDSIMPLE" in xml)


def _ensure_table_captions(document: Document, profile: FormatProfile, *, trace: FormatterExecutionTrace | None = None) -> None:
    existing_numbers = _existing_table_caption_numbers(document)
    table_index = 0
    scope_path: list[int] = []
    scope_counts: dict[str, int] = {}
    previous_text = ""
    caption_position = _table_caption_position(profile)
    for child in list(document.element.body):
        if child.tag == qn("w:p"):
            previous_text = _paragraph_xml_text(child).strip() or previous_text
            scope_path = _updated_caption_scope(scope_path, child, previous_text)
            continue
        if child.tag != qn("w:tbl"):
            continue
        table_index += 1
        body_reference_number = _caption_number_from_text(previous_text, kind="table") or str(table_index)
        inferred_number = _caption_number_for_profile(
            profile.table.caption.numbering,
            table_index,
            scope_path,
            scope_counts,
            kind="table",
        )
        caption_body = _caption_body_from_preceding_sentence(previous_text, body_reference_number, kind="table")
        target_caption = _find_nearby_caption(child, kind="table", position=caption_position)
        if target_caption is not None:
            _normalize_caption_element(target_caption, profile, "table", inferred_number, caption_body)
            existing_numbers.add(inferred_number)
            _record_trace(trace, "_ensure_bilingual_caption_near")
            _ensure_bilingual_caption_near(child, inferred_number, "table", profile, caption_body)
            continue
        opposite_position = "below" if caption_position == "above" else "above"
        opposite_caption = _find_nearby_caption(child, kind="table", position=opposite_position)
        if opposite_caption is not None:
            _normalize_caption_element(opposite_caption, profile, "table", inferred_number, caption_body)
            if caption_position == "above":
                child.addprevious(opposite_caption)
            else:
                child.addnext(opposite_caption)
            existing_numbers.add(inferred_number)
            _record_trace(trace, "_ensure_bilingual_caption_near")
            _ensure_bilingual_caption_near(child, inferred_number, "table", profile, caption_body)
            continue
        existing_group = _find_caption_group_by_number(document, kind="table", number=inferred_number)
        if existing_group:
            primary_caption = _primary_caption_from_group(
                existing_group,
                kind="table",
                preferred_prefix=profile.table.caption.prefix,
            )
            if primary_caption is not None:
                _normalize_caption_element(primary_caption, profile, "table", inferred_number, caption_body)
            _move_caption_group_near(child, existing_group, caption_position)
            existing_numbers.add(inferred_number)
            _record_trace(trace, "_ensure_bilingual_caption_near")
            _ensure_bilingual_caption_near(child, inferred_number, "table", profile, caption_body)
            continue
        caption_text = _build_caption_text(
            profile.table.caption.prefix,
            inferred_number,
            caption_body,
            profile.table.caption.separator,
        )
        caption = _caption_paragraph_xml(caption_text, profile.table.caption.font)
        if caption_position == "above":
            child.addprevious(caption)
        else:
            child.addnext(caption)
        existing_numbers.add(inferred_number)
        _record_trace(trace, "_ensure_bilingual_caption_near")
        _ensure_bilingual_caption_near(child, inferred_number, "table", profile, caption_body)


def _existing_table_caption_numbers(document: Document) -> set[str]:
    numbers: set[str] = set()
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not _is_table_caption(text):
            continue
        parsed = _split_caption_text(text, kind="table")
        if parsed:
            numbers.add(parsed[1])
    return numbers


def _find_nearby_caption(element, *, kind: str, position: str):
    sibling = element.getprevious() if position == "above" else element.getnext()
    while sibling is not None:
        if sibling.tag != qn("w:p"):
            sibling = sibling.getprevious() if position == "above" else sibling.getnext()
            continue
        text = _paragraph_xml_text(sibling).strip()
        if not text:
            sibling = sibling.getprevious() if position == "above" else sibling.getnext()
            continue
        if (kind == "table" and _is_table_caption(text)) or (kind == "figure" and _is_figure_caption(text)):
            return sibling
        return None
    return None


def _find_caption_group_by_number(document: Document, *, kind: str, number: str) -> list[object]:
    body = list(document.element.body)
    for index, child in enumerate(body):
        if not _caption_element_matches_number(child, kind=kind, number=number):
            continue
        start = index
        while start > 0 and _caption_element_matches_number(body[start - 1], kind=kind, number=number):
            start -= 1
        end = index + 1
        while end < len(body) and _caption_element_matches_number(body[end], kind=kind, number=number):
            end += 1
        return [body[item] for item in range(start, end)]
    return []


def _caption_element_matches_number(element, *, kind: str, number: str) -> bool:
    if element.tag != qn("w:p"):
        return False
    parsed = _split_caption_text(_paragraph_xml_text(element).strip(), kind=kind)
    return bool(parsed and parsed[1] == str(number))


def _primary_caption_from_group(group: list[object], *, kind: str, preferred_prefix: str):
    for element in group:
        parsed = _split_caption_text(_paragraph_xml_text(element).strip(), kind=kind)
        if parsed and parsed[0] == preferred_prefix:
            return element
    return group[0] if group else None


def _move_caption_group_near(element, group: list[object], position: str) -> None:
    if position == "above":
        for caption in group:
            element.addprevious(caption)
        return
    for caption in reversed(group):
        element.addnext(caption)


def _caption_body_from_preceding_sentence(text: str, number: str, *, kind: str) -> str | None:
    stripped = text.strip()
    prefix = "表" if kind == "table" else "图"
    if not stripped or not re.search(rf"{re.escape(prefix)}\s*{re.escape(str(number))}(?![\d.-])", stripped):
        return None
    cleaned = re.sub(rf"^.*?{re.escape(prefix)}\s*{re.escape(str(number))}\s*", "", stripped)
    cleaned = re.sub(r"^(给出|根据|如下|所示|见|列出|显示|说明|概括了?|展示|为)\s*", "", cleaned)
    cleaned = re.sub(r"^(一个|一项|一种|一份|若干)\s*", "", cleaned)
    cleaned = re.sub(r"[。；;，,：:]+$", "", cleaned).strip()
    cleaned = re.sub(r"\[[0-9,\]\[]+$", "", cleaned).strip()
    if len(cleaned) < 2:
        cleaned = "表格" if kind == "table" else "插图"
    if len(cleaned) > 32:
        cleaned = cleaned[:32].rstrip("，,、；;。")
    return cleaned


def _build_caption_text(prefix: str, number: str | int, body: str | None, separator: str) -> str:
    core = f"{prefix}{separator}{number}".strip()
    if not body:
        return core
    return f"{core}{separator}{body}".strip()


def _caption_number_from_text(text: str, *, kind: str) -> str | None:
    prefix = "表" if kind == "table" else "图"
    match = re.search(rf"{re.escape(prefix)}\s*[\s：:、.\-/]*\s*([0-9]+(?:[.-][0-9]+)*)", text)
    return match.group(1) if match else None


def _existing_figure_caption_numbers(document: Document) -> set[str]:
    numbers: set[str] = set()
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        parsed = _split_caption_text(text, kind="figure")
        if parsed:
            numbers.add(parsed[1])
    return numbers


def _paragraph_xml_text(paragraph_element) -> str:
    return "".join(node.text or "" for node in paragraph_element.xpath(".//w:t"))


def _caption_paragraph_xml(text: str, font: TextFont | None = None):
    paragraph = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    ppr.append(jc)
    paragraph.append(ppr)
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    paragraph.append(run)
    if font is not None:
        rpr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), font.latin)
        r_fonts.set(qn("w:hAnsi"), font.latin)
        r_fonts.set(qn("w:eastAsia"), font.chinese)
        rpr.append(r_fonts)
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), str(round(font.size_pt * 2)))
        rpr.append(size)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), font.color)
        rpr.append(color)
        if font.weight == "bold":
            bold = OxmlElement("w:b")
            rpr.append(bold)
        run.insert(0, rpr)
    return paragraph


def _ensure_figure_captions(document: Document, profile: FormatProfile, *, trace: FormatterExecutionTrace | None = None) -> None:
    existing_numbers = _existing_figure_caption_numbers(document)
    figure_index = 0
    scope_path: list[int] = []
    scope_counts: dict[str, int] = {}
    caption_position = _figure_caption_position(profile)
    for child in list(document.element.body):
        if child.tag == qn("w:p"):
            text = _paragraph_xml_text(child).strip()
            scope_path = _updated_caption_scope(scope_path, child, text)
        if child.tag != qn("w:p") or not _paragraph_xml_has_inline_image(child):
            continue
        previous_text = _previous_paragraph_text(child)
        target_caption = _find_nearby_caption(child, kind="figure", position=caption_position)
        opposite_position = "below" if caption_position == "above" else "above"
        opposite_caption = _find_nearby_caption(child, kind="figure", position=opposite_position)
        if target_caption is None and opposite_caption is None and not previous_text:
            continue
        figure_index += 1
        body_reference_number = _caption_number_from_text(previous_text, kind="figure") or str(figure_index)
        figure_number = _caption_number_for_profile(
            profile.figure.caption.numbering,
            figure_index,
            scope_path,
            scope_counts,
            kind="figure",
        )
        caption_body = _caption_body_from_preceding_sentence(previous_text, body_reference_number, kind="figure")
        if target_caption is not None:
            _normalize_caption_element(target_caption, profile, "figure", figure_number, caption_body)
            existing_numbers.add(figure_number)
            _record_trace(trace, "_ensure_bilingual_caption_near")
            _ensure_bilingual_caption_near(child, figure_number, "figure", profile, caption_body)
            continue
        if opposite_caption is not None:
            _normalize_caption_element(opposite_caption, profile, "figure", figure_number, caption_body)
            if caption_position == "below":
                child.addnext(opposite_caption)
            else:
                child.addprevious(opposite_caption)
            existing_numbers.add(figure_number)
            _record_trace(trace, "_ensure_bilingual_caption_near")
            _ensure_bilingual_caption_near(child, figure_number, "figure", profile, caption_body)
            continue
        existing_group = _find_caption_group_by_number(document, kind="figure", number=figure_number)
        if existing_group:
            primary_caption = _primary_caption_from_group(
                existing_group,
                kind="figure",
                preferred_prefix=profile.figure.caption.prefix,
            )
            if primary_caption is not None:
                _normalize_caption_element(primary_caption, profile, "figure", figure_number, caption_body)
            _move_caption_group_near(child, existing_group, caption_position)
            existing_numbers.add(figure_number)
            _record_trace(trace, "_ensure_bilingual_caption_near")
            _ensure_bilingual_caption_near(child, figure_number, "figure", profile, caption_body)
            continue
        caption_text = _build_caption_text(
            profile.figure.caption.prefix,
            figure_number,
            caption_body,
            profile.figure.caption.separator,
        )
        caption = _caption_paragraph_xml(caption_text, profile.figure.caption.font)
        if caption_position == "below":
            child.addnext(caption)
        else:
            child.addprevious(caption)
        existing_numbers.add(figure_number)
        _record_trace(trace, "_ensure_bilingual_caption_near")
        _ensure_bilingual_caption_near(child, figure_number, "figure", profile, caption_body)


def _paragraph_xml_has_inline_image(paragraph_element) -> bool:
    return bool(paragraph_element.xpath(".//*[local-name()='inline']"))


def _ensure_bilingual_caption_near(element, number: str, kind: str, profile: FormatProfile, body: str | None) -> None:
    caption = profile.table.caption if kind == "table" else profile.figure.caption
    if not caption.bilingual:
        return
    english_prefix = caption.english_prefix or ("Table" if kind == "table" else "Figure")
    if _english_caption_near(element, english_prefix, number):
        return
    text = _build_caption_text(english_prefix, number, body, caption.separator)
    english_caption = _caption_paragraph_xml(text, caption.font)
    caption_position = _table_caption_position(profile) if kind == "table" else _figure_caption_position(profile)
    primary_caption = _find_nearby_caption(element, kind=kind, position=caption_position)
    if primary_caption is not None:
        primary_caption.addnext(english_caption)
        return
    if kind == "table":
        if caption_position == "above":
            element.addprevious(english_caption)
        else:
            element.addnext(english_caption)
    elif caption_position == "below":
        element.addnext(english_caption)
    else:
        element.addprevious(english_caption)


def _english_caption_near(element, prefix: str, number: str) -> bool:
    pattern = re.compile(rf"^{re.escape(prefix)}\s*[\s：:、.\-/]*\s*{re.escape(str(number))}\b", re.IGNORECASE)
    for sibling in _scan_adjacent_paragraphs(element.getprevious(), step="previous"):
        text = _paragraph_xml_text(sibling).strip()
        if text and pattern.match(text):
            return True
        if text:
            break
    for sibling in _scan_adjacent_paragraphs(element.getnext(), step="next"):
        text = _paragraph_xml_text(sibling).strip()
        if text and pattern.match(text):
            return True
        if text:
            break
    return False


def _apply_figure_size_rules(document: Document, profile: FormatProfile) -> None:
    for shape in document.inline_shapes:
        width_mm = shape.width / 36000
        if width_mm <= profile.figure.half_column_max_mm:
            continue
        target_mm = width_mm
        if width_mm < profile.figure.full_width_min_mm:
            target_mm = profile.figure.full_width_min_mm
        elif width_mm > profile.figure.full_width_max_mm:
            target_mm = profile.figure.full_width_max_mm
        if abs(target_mm - width_mm) < 0.1:
            continue
        ratio = target_mm / width_mm
        shape.width = Mm(target_mm)
        shape.height = Emu(round(shape.height * ratio))


def _apply_body_paragraph(paragraph, profile: FormatProfile) -> None:
    if not paragraph.text.strip():
        return
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = Cm(profile.body.first_line_indent_chars * 0.37)
    paragraph.paragraph_format.line_spacing = profile.body.line_spacing
    paragraph.paragraph_format.space_before = Pt(profile.body.space_before_pt)
    paragraph.paragraph_format.space_after = Pt(profile.body.space_after_pt)
    _set_snap_to_grid(paragraph, profile.document_grid.snap_to_grid)
    _apply_paragraph_alignment(paragraph, profile.body.alignment)
    _apply_runs_font(paragraph, profile.body.font)


def _apply_heading_paragraph(paragraph, profile: FormatProfile, level: int) -> None:
    heading = _heading_rule(profile, level)
    _apply_word_heading_style(paragraph, min(heading.level, 9))
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = Cm(heading.first_line_indent_chars * 0.37)
    paragraph.paragraph_format.line_spacing = heading.line_spacing or profile.body.line_spacing
    paragraph.paragraph_format.space_before = Pt(heading.space_before_pt)
    paragraph.paragraph_format.space_after = Pt(heading.space_after_pt)
    paragraph.paragraph_format.keep_with_next = heading.keep_with_next
    paragraph.paragraph_format.page_break_before = heading.page_break_before
    _set_snap_to_grid(paragraph, profile.document_grid.snap_to_grid)
    _apply_paragraph_alignment(paragraph, heading.alignment)
    _apply_runs_font(paragraph, heading.font)


def _apply_word_heading_style(paragraph, level: int) -> None:
    style_name = f"Heading {min(max(level, 1), 9)}"
    try:
        paragraph.style = style_name
    except KeyError:
        try:
            style = paragraph.part.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            try:
                style.base_style = paragraph.part.styles["Normal"]
            except KeyError:
                pass
            paragraph.style = style
        except Exception:
            pass
    _set_paragraph_outline_level(paragraph, level - 1)


def _set_paragraph_outline_level(paragraph, outline_level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(min(max(outline_level, 0), 8)))


def _apply_abstract_heading(paragraph, profile: FormatProfile) -> None:
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = profile.body.line_spacing
    paragraph.paragraph_format.space_before = Pt(profile.body.space_before_pt)
    paragraph.paragraph_format.space_after = Pt(profile.body.space_after_pt)
    _set_snap_to_grid(paragraph, profile.document_grid.snap_to_grid)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_runs_font(paragraph, profile.abstract.title_font)


def _apply_abstract_body(paragraph, profile: FormatProfile) -> None:
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = profile.body.line_spacing
    paragraph.paragraph_format.space_before = Pt(profile.body.space_before_pt)
    paragraph.paragraph_format.space_after = Pt(profile.body.space_after_pt)
    _set_snap_to_grid(paragraph, profile.document_grid.snap_to_grid)
    _apply_paragraph_alignment(paragraph, profile.body.alignment)
    _apply_runs_font(paragraph, profile.abstract.body_font)


def _apply_keywords_paragraph(paragraph, profile: FormatProfile) -> None:
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = profile.body.line_spacing
    paragraph.paragraph_format.space_before = Pt(profile.body.space_before_pt)
    paragraph.paragraph_format.space_after = Pt(profile.body.space_after_pt)
    _set_snap_to_grid(paragraph, profile.document_grid.snap_to_grid)
    _apply_paragraph_alignment(paragraph, "left")
    _apply_runs_font(paragraph, profile.abstract.body_font)


def _apply_caption_paragraph(paragraph, font: TextFont, snap_to_grid: bool) -> None:
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    _set_snap_to_grid(paragraph, snap_to_grid)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_runs_font(paragraph, font)


def _apply_equation_paragraph(paragraph, profile: FormatProfile, equation_number: int) -> None:
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = Cm(0)
    _ensure_equation_number(paragraph, profile, equation_number)
    _apply_paragraph_alignment(paragraph, profile.equations.alignment)
    _apply_runs_font(paragraph, TextFont(chinese=profile.equations.font, latin=profile.equations.font, size_pt=profile.body.font.size_pt, weight="normal", color=profile.body.font.color))


def _ensure_equation_number(paragraph, profile: FormatProfile, equation_number: int) -> None:
    if profile.equations.numbering == "none":
        return
    text = paragraph.text.strip()
    if _equation_has_visible_number(text):
        return
    if not paragraph.runs:
        paragraph.add_run("")
    if profile.equations.numbering == "left":
        paragraph.runs[0].text = f"({equation_number}) {paragraph.runs[0].text}"
        return
    paragraph.add_run(f"\t({equation_number})")


def _equation_has_visible_number(text: str) -> bool:
    return bool(re.match(r"^\(\d+\)\s+", text) or re.search(r"\s+\(\d+\)$", text))


def _apply_reference_paragraph(paragraph, profile: FormatProfile) -> None:
    paragraph.paragraph_format.left_indent = Cm(profile.references.hanging_indent_chars * 0.37)
    paragraph.paragraph_format.first_line_indent = Cm(-profile.references.hanging_indent_chars * 0.37)
    paragraph.paragraph_format.line_spacing = profile.body.line_spacing
    paragraph.paragraph_format.space_before = Pt(profile.body.space_before_pt)
    paragraph.paragraph_format.space_after = Pt(profile.body.space_after_pt)
    _set_snap_to_grid(paragraph, profile.document_grid.snap_to_grid)
    _apply_runs_font(paragraph, profile.references.font)


def _apply_appendix_heading(paragraph, profile: FormatProfile) -> None:
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = profile.body.line_spacing
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    _set_snap_to_grid(paragraph, profile.document_grid.snap_to_grid)
    _apply_paragraph_alignment(paragraph, profile.appendix.title_alignment)
    _apply_runs_font(paragraph, profile.appendix.title_font)


def _apply_appendix_body(paragraph, profile: FormatProfile) -> None:
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = Pt(profile.appendix.body_first_line_indent_chars * profile.appendix.body_font.size_pt)
    paragraph.paragraph_format.line_spacing = profile.appendix.body_line_spacing
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    _set_snap_to_grid(paragraph, profile.document_grid.snap_to_grid)
    _apply_paragraph_alignment(paragraph, profile.appendix.body_alignment)
    _apply_runs_font(paragraph, profile.appendix.body_font)


def _normalize_body_text(paragraph, profile: FormatProfile, *, normalize_list_marker: bool = True) -> None:
    if not profile.unit_rules.enforce_consistency:
        return
    original = paragraph.text
    normalized = _normalize_fullwidth_numbers(original) if profile.unit_rules.normalize_fullwidth_numbers else original
    if profile.unit_rules.use_si_symbols:
        normalized = _normalize_si_symbols(normalized)
    normalized = _normalize_unit_spacing(normalized, profile)
    if normalize_list_marker:
        normalized = _normalize_list_marker(normalized, profile)
    if normalized != original:
        paragraph.text = normalized


def _normalize_table_text(table, profile: FormatProfile) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _normalize_body_text(paragraph, profile)


def _apply_heading_numbering(paragraph, profile: FormatProfile, level: int, counters: list[int]) -> None:
    heading = _heading_rule(profile, level)
    prefix = _heading_number_prefix(profile, heading, level, counters)
    if not prefix:
        return
    text = paragraph.text.strip()
    if _looks_like_numbered_heading(text):
        return
    if not paragraph.runs:
        paragraph.add_run("")
    first_run = paragraph.runs[0]
    original = first_run.text
    first_run.text = f"{prefix} {original}".strip()


def _heading_number_prefix(profile: FormatProfile, heading: HeadingSettings, level: int, counters: list[int]) -> str | None:
    if not profile.numbering.enabled:
        return None
    if level < 1 or level > len(counters):
        return None
    counters[level - 1] += 1
    for index in range(level, len(counters)):
        counters[index] = 0
    numbers = [str(value) for value in counters[:level] if value > 0]
    if not numbers:
        return None
    pattern = (profile.numbering.heading_pattern or "").strip()
    if pattern:
        return _expand_number_pattern(pattern, numbers)
    style = heading.numbering.strip().lower()
    if style == "none":
        return None
    if style == "chapter" and level == 1:
        return f"第{numbers[0]}章"
    if style in {"decimal", "decimal-dot", "decimal-chinese-pause"}:
        return ".".join(numbers)
    return ".".join(numbers)


def _expand_number_pattern(pattern: str, numbers: list[str]) -> str:
    def replace(match: re.Match[str]) -> str:
        index = int(match.group(1)) - 1
        if 0 <= index < len(numbers):
            return numbers[index]
        return numbers[-1]

    return re.sub(r"%(\d+)", replace, pattern)


def _looks_like_numbered_heading(text: str) -> bool:
    return bool(
        re.match(r"^(第[一二三四五六七八九十百0-9]+[章节])", text)
        or re.match(r"^[0-9]+(?:\.[0-9]+)*(?:[、.）\)])\s*", text)
        or re.match(r"^[0-9]+(?:\.[0-9]+)*\s+", text)
    )


def _normalize_fullwidth_numbers(text: str) -> str:
    table = str.maketrans("０１２３４５６７８９", "0123456789")
    return text.translate(table)


def _normalize_unit_spacing(text: str, profile: FormatProfile) -> str:
    if profile.unit_rules.unit_spacing == "preserve":
        return text
    units = sorted(set(profile.unit_rules.measurement_units + profile.unit_rules.currency_units), key=len, reverse=True)
    if not units:
        return text
    unit_pattern = "|".join(re.escape(unit) for unit in units)
    if profile.unit_rules.unit_spacing == "space":
        return re.sub(rf"(\d)\s*({unit_pattern})\b", r"\1 \2", text)
    return re.sub(rf"(\d)\s+({unit_pattern})\b", r"\1\2", text)


def _normalize_si_symbols(text: str) -> str:
    replacements = {
        "毫米": "mm",
        "厘米": "cm",
        "千米": "km",
        "米": "m",
        "千克": "kg",
        "公斤": "kg",
        "克": "g",
        "秒": "s",
        "分钟": "min",
        "小时": "h",
    }
    normalized = text
    for source, target in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
        normalized = re.sub(rf"(?<=\d)\s*{re.escape(source)}", target, normalized)
    return normalized


def _normalize_list_marker(text: str, profile: FormatProfile) -> str:
    stripped = text.lstrip()
    prefix = text[: len(text) - len(stripped)]
    if re.match(r"^\d+(?:\.\d+)+\s*", stripped):
        return text
    bullet_match = re.match(r"^[·•●\-]\s*(.+)$", stripped)
    if bullet_match:
        return f"{prefix}{profile.list_numbering.unordered_marker} {bullet_match.group(1)}"
    ordered_match = re.match(r"^(\(?)(\d+)([、）\)])\s*(.+)$", stripped)
    if not ordered_match:
        ordered_match = re.match(r"^(\(?)(\d+)(\.)\s*(?!\d)(.+)$", stripped)
    if not ordered_match:
        return text
    marker = _ordered_marker(profile.list_numbering.ordered_pattern, int(ordered_match.group(2)))
    return f"{prefix}{marker} {ordered_match.group(4)}"


def _ordered_marker(pattern: str, number: int) -> str:
    if "%d" in pattern:
        return pattern.replace("%d", str(number))
    if "1" in pattern:
        return pattern.replace("1", str(number), 1)
    return f"{number}{pattern}"


def _updated_caption_scope(scope_path: list[int], paragraph_element, text: str) -> list[int]:
    level = _paragraph_xml_heading_level(paragraph_element, text)
    if level is None:
        return scope_path
    parsed = _heading_numeric_path(text)
    if parsed:
        return parsed[: max(1, level)]
    next_path = list(scope_path)
    while len(next_path) < level:
        next_path.append(0)
    next_path[level - 1] += 1
    return next_path[:level]


def _paragraph_xml_heading_level(paragraph_element, text: str) -> int | None:
    from_text = _heading_level(text, "")
    if from_text is not None:
        return from_text
    ppr = paragraph_element.find(qn("w:pPr"))
    if ppr is None:
        return None
    style = ppr.find(qn("w:pStyle"))
    value = style.get(qn("w:val")) if style is not None else ""
    match = re.search(r"Heading\s*(\d+)", value or "", re.IGNORECASE)
    if match:
        return int(match.group(1))
    return None


def _heading_numeric_path(text: str) -> list[int]:
    stripped = text.strip()
    chapter_match = re.match(r"^第([一二三四五六七八九十百千万0-9]+)[章节]", stripped)
    if chapter_match:
        return [_number_token_to_int(chapter_match.group(1))]
    decimal_match = re.match(r"^([0-9]+(?:\.[0-9]+)*)[、.）\)]?\s*", stripped)
    if decimal_match:
        return [int(item) for item in decimal_match.group(1).split(".") if item]
    return []


def _number_token_to_int(token: str) -> int:
    if token.isdigit():
        return int(token)
    digits = {"零": 0, "一": 1, "二": 2, "两": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}
    units = {"十": 10, "百": 100, "千": 1000, "万": 10000}
    total = 0
    current = 0
    for char in token:
        if char in digits:
            current = digits[char]
        elif char in units:
            unit = units[char]
            total += (current or 1) * unit
            current = 0
    return total + current if total or current else 1


def _caption_number_for_profile(
    numbering: str,
    global_index: int,
    scope_path: list[int],
    scope_counts: dict[str, int],
    *,
    kind: str,
) -> str:
    if numbering == "continuous":
        return str(global_index)
    if numbering == "section":
        prefix_path = scope_path[:2] if len(scope_path) >= 2 else (scope_path[:1] or [1])
        prefix = ".".join(str(item) for item in prefix_path)
        key = f"{kind}:section:{prefix}"
        scope_counts[key] = scope_counts.get(key, 0) + 1
        return f"{prefix}-{scope_counts[key]}"
    chapter = scope_path[0] if scope_path else 1
    key = f"{kind}:chapter:{chapter}"
    scope_counts[key] = scope_counts.get(key, 0) + 1
    return f"{chapter}-{scope_counts[key]}"


def _heading_rule(profile: FormatProfile, level: int) -> HeadingSettings:
    exact = next((heading for heading in profile.headings if heading.level == level), None)
    if exact is not None:
        return exact
    lower_or_equal = [heading for heading in profile.headings if heading.level <= level]
    if lower_or_equal:
        return sorted(lower_or_equal, key=lambda heading: heading.level)[-1]
    return profile.headings[0]


def _matching_heading(profile: FormatProfile, text: str, style_name: str) -> HeadingSettings | None:
    stripped = text.strip()
    if not stripped:
        return None
    level = None
    if style_name.lower().startswith("heading"):
        match = re.search(r"(\d+)", style_name)
        level = int(match.group(1)) if match else 1
    elif re.match(r"^第[一二三四五六七八九十百0-9]+[章节]", stripped):
        level = 1
    elif re.match(r"^[0-9]+\.[0-9]+", stripped):
        level = 2
    if level is None:
        return None
    return next((heading for heading in profile.headings if heading.level == level), profile.headings[0])


def _apply_runs_font(paragraph, font: TextFont) -> None:
    runs = paragraph.runs or [paragraph.add_run("")]
    for run in runs:
        run.font.name = font.latin
        run.font.size = Pt(font.size_pt)
        run.font.bold = font.weight == "bold"
        run.font.color.rgb = RGBColor.from_string(font.color)
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font.chinese)


def _apply_notes(path: Path, profile: FormatProfile) -> None:
    tmp_path: Path | None = None
    try:
        with ZipFile(path) as source_package:
            note_roots = {
                part_name: ET.fromstring(source_package.read(part_name))
                for part_name in ("word/footnotes.xml", "word/endnotes.xml")
                if part_name in source_package.namelist()
            }
            if not note_roots:
                return
            for root in note_roots.values():
                _apply_note_root(root, profile)
            with NamedTemporaryFile(delete=False, suffix=".docx", dir=path.parent) as tmp_file:
                tmp_path = Path(tmp_file.name)
            with ZipFile(tmp_path, "w") as target_package:
                for item in source_package.infolist():
                    if item.filename in note_roots:
                        continue
                    target_package.writestr(item, source_package.read(item.filename))
                for part_name, root in note_roots.items():
                    target_package.writestr(part_name, _xml_bytes(root))
        tmp_path.replace(path)
    except Exception as exc:
        raise DocumentFormatError(f"DOCX notes formatting failed: {exc}") from exc
    finally:
        if tmp_path is not None and tmp_path.exists():
            tmp_path.unlink()


def _apply_note_root(root: ET.Element, profile: FormatProfile) -> None:
    font = profile.notes.font
    for note in [*root.findall("w:footnote", NS), *root.findall("w:endnote", NS)]:
        note_id = note.get(qn("w:id"))
        if note_id in {"-1", "0"}:
            continue
        for paragraph in note.findall(".//w:p", NS):
            _apply_note_paragraph_xml(paragraph, profile)
            for run in paragraph.findall(".//w:r", NS):
                _apply_run_font_xml(run, font)


def _apply_note_paragraph_xml(paragraph: ET.Element, profile: FormatProfile) -> None:
    ppr = paragraph.find(qn("w:pPr"))
    if ppr is None:
        ppr = ET.Element(qn("w:pPr"))
        paragraph.insert(0, ppr)
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = ET.SubElement(ppr, qn("w:spacing"))
    spacing.set(qn("w:before"), str(round(profile.notes.space_before_pt * 20)))
    spacing.set(qn("w:after"), str(round(profile.notes.space_after_pt * 20)))
    spacing.set(qn("w:line"), str(round(profile.notes.line_spacing * 240)))
    spacing.set(qn("w:lineRule"), "auto")


def _apply_run_font_xml(run: ET.Element, font: TextFont) -> None:
    rpr = run.find(qn("w:rPr"))
    if rpr is None:
        rpr = ET.Element(qn("w:rPr"))
        run.insert(0, rpr)
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = ET.SubElement(rpr, qn("w:rFonts"))
    r_fonts.set(qn("w:ascii"), font.latin)
    r_fonts.set(qn("w:hAnsi"), font.latin)
    r_fonts.set(qn("w:eastAsia"), font.chinese)
    size = rpr.find(qn("w:sz"))
    if size is None:
        size = ET.SubElement(rpr, qn("w:sz"))
    size.set(qn("w:val"), str(round(font.size_pt * 2)))
    size_cs = rpr.find(qn("w:szCs"))
    if size_cs is None:
        size_cs = ET.SubElement(rpr, qn("w:szCs"))
    size_cs.set(qn("w:val"), str(round(font.size_pt * 2)))
    color = rpr.find(qn("w:color"))
    if color is None:
        color = ET.SubElement(rpr, qn("w:color"))
    color.set(qn("w:val"), font.color)
    bold = rpr.find(qn("w:b"))
    if font.weight == "bold":
        if bold is None:
            bold = ET.SubElement(rpr, qn("w:b"))
        bold.set(qn("w:val"), "1")
    elif bold is not None:
        bold.set(qn("w:val"), "0")


def _xml_bytes(root: ET.Element) -> bytes:
    ET.register_namespace("w", NS["w"])
    ET.register_namespace("r", NS["r"])
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _set_snap_to_grid(paragraph, enabled: bool) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    snap = ppr.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        ppr.append(snap)
    snap.set(qn("w:val"), "1" if enabled else "0")


def _apply_paragraph_alignment(paragraph, alignment: TextAlignment) -> None:
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    paragraph.alignment = mapping[alignment]


def _is_table_caption(text: str) -> bool:
    if re.search(r"(给出|如下|所示|见|列出|显示|说明|展示|概括)", text):
        return False
    return _split_caption_text(text, kind="table") is not None


def _is_figure_caption(text: str) -> bool:
    if re.search(r"(给出|如下|所示|见|列出|显示|说明|展示|概括)", text):
        return False
    return _split_caption_text(text, kind="figure") is not None


def _is_equation(text: str) -> bool:
    if not text:
        return False
    cleaned = re.sub(r"^\(\d+\)\s*|\s*\(\d+\)$", "", text).strip()
    return any(symbol in cleaned for symbol in ("=", "＋", "+", "-", "*", "/", "^")) and len(cleaned) <= 120


def _is_appendix_heading_text(text: str) -> bool:
    return bool(re.match(r"^\s*(附录(?:\s*[A-ZＡ-Ｚ一二三四五六七八九十0-9]+)?|Appendix\b)", text.strip(), re.IGNORECASE))


def _is_appendix_terminator(text: str) -> bool:
    return text.strip() in {"参考文献", "References", "致谢", "Acknowledgements", "Acknowledgments"}


def _apply_table_rules(table, profile: FormatProfile) -> None:
    table.autofit = profile.table.autofit
    if profile.table.header_repeat and table.rows:
        _set_row_header_repeat(table.rows[0])
    if profile.table.border_style == "three_line":
        _apply_table_borders(table, top=True, bottom=True, inside_h=True, inside_v=False, sides=False)
    elif profile.table.border_style == "minimal":
        _apply_table_borders(table, top=True, bottom=True, inside_h=False, inside_v=False, sides=False)
    elif profile.table.border_style == "full_grid":
        _apply_table_borders(table, top=True, bottom=True, inside_h=True, inside_v=True, sides=True)


def _table_caption_position(profile: FormatProfile) -> str:
    if profile.table.enforce_caption_above:
        return "above"
    return profile.table.caption.position


def _figure_caption_position(profile: FormatProfile) -> str:
    if profile.figure.enforce_caption_below:
        return "below"
    return profile.figure.caption.position


def _previous_paragraph_text(paragraph) -> str:
    sibling = paragraph.getprevious()
    while sibling is not None:
        if sibling.tag == qn("w:p"):
            text = _paragraph_xml_text(sibling).strip()
            if text:
                return text
        sibling = sibling.getprevious()
    return ""


def _scan_adjacent_paragraphs(start, *, step: str):
    sibling = start
    while sibling is not None and sibling.tag == qn("w:p"):
        yield sibling
        sibling = sibling.getprevious() if step == "previous" else sibling.getnext()


def _set_row_header_repeat(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def _apply_table_borders(table, *, top: bool, bottom: bool, inside_h: bool, inside_v: bool, sides: bool) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    enabled_edges = {
        "top": top,
        "bottom": bottom,
        "insideH": inside_h,
        "left": sides,
        "right": sides,
        "insideV": inside_v,
    }
    for edge, enabled in enabled_edges.items():
        element = OxmlElement(f"w:{edge}")
        if enabled:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "8")
            element.set(qn("w:color"), "000000")
        else:
            element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def _normalize_caption_paragraph(paragraph, profile: FormatProfile, kind: str) -> None:
    caption = profile.table.caption if kind == "table" else profile.figure.caption
    text = paragraph.text.strip()
    parsed = _split_caption_text(text, kind=kind)
    if parsed is None:
        return
    parsed_prefix, number, body = parsed
    normalized = _build_caption_text(_caption_output_prefix(caption, kind, parsed_prefix), number, body, caption.separator)
    if normalized != text:
        paragraph.text = normalized


def _normalize_caption_element(element, profile: FormatProfile, kind: str, fallback_number: str, fallback_body: str | None) -> None:
    caption = profile.table.caption if kind == "table" else profile.figure.caption
    text = _paragraph_xml_text(element).strip()
    parsed = _split_caption_text(text, kind=kind)
    number = fallback_number
    body = fallback_body or (parsed[2] if parsed and parsed[2] else None)
    normalized = _build_caption_text(caption.prefix, number, body, caption.separator)
    _replace_paragraph_xml_text(element, normalized)


def _split_caption_text(text: str, *, kind: str) -> tuple[str, str, str | None] | None:
    patterns = {
        "table": re.compile(r"^(?P<prefix>表|Table)\s*[\s：:、.\-/]*\s*(?P<number>\d+(?:[.-]\d+)*)(?P<rest>.*)$", re.IGNORECASE),
        "figure": re.compile(r"^(?P<prefix>图|Figure)\s*[\s：:、.\-/]*\s*(?P<number>\d+(?:[.-]\d+)*)(?P<rest>.*)$", re.IGNORECASE),
    }
    pattern = patterns.get(kind)
    if pattern is None:
        return None
    match = pattern.match(text.strip())
    if match is None:
        return None
    rest = (match.group("rest") or "").strip()
    rest = re.sub(r"^[\s：:、.\-/]+", "", rest).strip()
    body = re.sub(r"[。；;，,：:]+$", "", rest).strip() if rest else ""
    return match.group("prefix"), match.group("number"), body or None


def _replace_paragraph_xml_text(element, text: str) -> None:
    text_nodes = element.xpath(".//w:t")
    if not text_nodes:
        run = OxmlElement("w:r")
        text_node = OxmlElement("w:t")
        text_node.text = text
        run.append(text_node)
        element.append(run)
        return
    text_nodes[0].text = text
    for node in text_nodes[1:]:
        node.text = ""


def _caption_output_prefix(caption, kind: str, parsed_prefix: str) -> str:
    english_prefix = caption.english_prefix or ("Table" if kind == "table" else "Figure")
    if caption.bilingual and parsed_prefix.lower() == english_prefix.lower():
        return english_prefix
    if caption.bilingual and parsed_prefix.lower() in {"table", "figure"}:
        return english_prefix
    return caption.prefix
