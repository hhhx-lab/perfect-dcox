from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile
import re
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor
from pypdf import PdfReader

from app.documents.rule_registry import annotate_docx_quality_issues_with_registry_fields
from app.documents.ooxml import NS, OoxmlDocumentFeatures, OoxmlInspectionError, inspect_ooxml_features
from app.documents.structure import DocumentStructure, ParagraphRole, classify_document
from app.models import QualityIssue
from app.profiles.models import FormatProfile, TextAlignment, TextFont


class QualityInspectionError(RuntimeError):
    pass


def inspect_docx_quality(
    path: Path,
    profile: FormatProfile,
    *,
    inherited_header_footer: bool = False,
) -> list[QualityIssue]:
    try:
        document = Document(path)
    except Exception as exc:
        raise QualityInspectionError(f"DOCX quality inspection failed to open input: {exc}") from exc

    structure = classify_document(document)
    try:
        features = inspect_ooxml_features(path)
    except OoxmlInspectionError as exc:
        features = None
        feature_issue = QualityIssue(
            issue_id="docx_ooxml_features",
            status="fail",
            severity="high",
            check_key="docx.ooxml.features",
            title="DOCX internal OOXML feature inspection failed.",
            description=str(exc),
            recommendation="Regenerate the DOCX before claiming compliance.",
            fixable=False,
        )
    else:
        feature_issue = _inspect_ooxml_feature_inventory(features)
    return annotate_docx_quality_issues_with_registry_fields(
        [
            _inspect_page_setup(document, profile),
            _inspect_margins(document, profile)
            if profile.quality.check_margins
            else _disabled_quality_issue("docx_page_margins", "docx.page.margins", "DOCX page-margin inspection is disabled by the profile.", "quality.check_margins"),
            _inspect_header_footer(document, profile, features)
            if not inherited_header_footer
            else _disabled_quality_issue("docx_header_footer", "docx.header_footer", "DOCX header/footer inspection is delegated to the selected template.", "template_binding.inherit_header_footer"),
            *(
                _inspect_header_footer_fields(document, profile, features)
                if not inherited_header_footer
                else _disabled_header_footer_field_issues("template_binding.inherit_header_footer")
            ),
            _inspect_body_style(document, structure, profile),
            *_inspect_body_style_fields(document, structure, profile),
            _inspect_heading_style(document, structure, profile)
            if profile.quality.check_headings
            else _disabled_quality_issue("docx_heading_style", "docx.heading.style", "DOCX heading-style inspection is disabled by the profile.", "quality.check_headings"),
            *(
                _inspect_heading_style_fields(document, structure, profile)
                if profile.quality.check_headings
                else _disabled_heading_style_field_issues()
            ),
            _inspect_heading_numbering(document, structure, profile)
            if profile.quality.check_headings
            else _disabled_quality_issue("docx_heading_numbering", "docx.heading_numbering", "DOCX heading-numbering inspection is disabled by the profile.", "quality.check_headings"),
            _inspect_document_grid(document, profile, features),
            _inspect_table_borders(document, profile),
            *_inspect_table_rule_fields(document, profile),
            _inspect_captions(document, structure, profile),
            _inspect_equations(document, structure, profile),
            _inspect_role_style_consistency(document, structure, profile),
            _inspect_raw_latex(document),
            _inspect_basic_page_numbers(document, profile, features)
            if not inherited_header_footer
            else _disabled_quality_issue("docx_page_number", "docx.page_number", "DOCX page-number inspection is delegated to the selected template.", "template_binding.inherit_header_footer"),
            *(
                _inspect_page_number_fields(document, profile, features)
                if not inherited_header_footer
                else _disabled_page_number_field_issues("template_binding.inherit_header_footer")
            ),
            feature_issue,
            _inspect_field_update_policy(features),
            _inspect_toc_fields(document, structure, features, profile),
            *_inspect_toc_field_issues(document, structure, features, profile),
            _inspect_template_body_slot(document, profile),
            _inspect_template_placeholders(document),
            _inspect_section_complexity(features),
            _inspect_notes_support(path, features, profile),
            _inspect_appendix_style(document, profile),
            _inspect_figure_sizes(features, profile),
            _inspect_image_caption_pairing(document, structure, features, profile),
            *_inspect_visual_caption_fields(document, features, profile),
            _inspect_list_numbering(features, profile),
            _inspect_unit_rules(document, profile),
        ]
    )


BODY_FIELD_CHECKS: tuple[tuple[str, str, str], ...] = (
    ("body.font.chinese", "docx.body.font.chinese", "DOCX body Chinese font"),
    ("body.font.latin", "docx.body.font.latin", "DOCX body Latin font"),
    ("body.font.size_pt", "docx.body.font.size", "DOCX body font size"),
    ("body.font.color", "docx.body.font.color", "DOCX body font color"),
    ("body.line_spacing", "docx.body.line_spacing", "DOCX body line spacing"),
    ("body.first_line_indent_chars", "docx.body.first_line_indent", "DOCX body first-line indent"),
    ("body.space_before_pt", "docx.body.space_before", "DOCX body space before"),
    ("body.space_after_pt", "docx.body.space_after", "DOCX body space after"),
    ("body.alignment", "docx.body.alignment", "DOCX body alignment"),
)


HEADING_FIELD_CHECKS: tuple[tuple[str, str, str], ...] = (
    ("headings.font.chinese", "docx.heading.font.chinese", "DOCX heading Chinese font"),
    ("headings.font.latin", "docx.heading.font.latin", "DOCX heading Latin font"),
    ("headings.font.size_pt", "docx.heading.font.size", "DOCX heading font size"),
    ("headings.font.color", "docx.heading.font.color", "DOCX heading font color"),
    ("headings.font.weight", "docx.heading.font.weight", "DOCX heading font weight"),
    ("headings.alignment", "docx.heading.alignment", "DOCX heading alignment"),
    ("headings.line_spacing", "docx.heading.line_spacing", "DOCX heading line spacing"),
    ("headings.space_before_pt", "docx.heading.space_before", "DOCX heading space before"),
    ("headings.space_after_pt", "docx.heading.space_after", "DOCX heading space after"),
    ("headings.first_line_indent_chars", "docx.heading.first_line_indent", "DOCX heading first-line indent"),
    ("headings.pagination", "docx.heading.pagination", "DOCX heading pagination"),
)


HEADER_FOOTER_FIELD_CHECKS: tuple[tuple[str, str, str], ...] = (
    ("header_footer.header_text", "docx.header_footer.header_text", "DOCX header text"),
    ("header_footer.header_alignment", "docx.header_footer.header_alignment", "DOCX header alignment"),
    ("header_footer.footer_text", "docx.header_footer.footer_text", "DOCX footer text"),
    ("header_footer.footer_alignment", "docx.header_footer.footer_alignment", "DOCX footer alignment"),
    ("header_footer.different_first_page", "docx.header_footer.different_first_page", "DOCX first-page header/footer setting"),
    ("header_footer.different_odd_even", "docx.header_footer.different_odd_even", "DOCX odd/even header/footer setting"),
)


PAGE_NUMBER_FIELD_CHECKS: tuple[tuple[str, str, str], ...] = (
    ("header_footer.footer_page_number", "docx.page_number.field", "DOCX footer page-number field"),
    ("header_footer.page_number_format", "docx.page_number.format", "DOCX page-number format"),
    ("header_footer.page_number_start", "docx.page_number.start", "DOCX page-number start"),
)


TABLE_RULE_FIELD_CHECKS: tuple[tuple[str, str, str], ...] = (
    ("table.border_style", "docx.table.border_style", "DOCX table border style"),
    ("table.header_repeat", "docx.table.header_repeat", "DOCX table header repeat"),
)


VISUAL_CAPTION_FIELD_CHECKS: tuple[tuple[str, str, str], ...] = (
    ("table.caption.position", "docx.table.caption.position", "DOCX table caption position"),
    ("table.caption.bilingual", "docx.table.caption.bilingual", "DOCX table bilingual captions"),
    ("figure.caption.position", "docx.figure.caption.position", "DOCX figure caption position"),
    ("figure.caption.bilingual", "docx.figure.caption.bilingual", "DOCX figure bilingual captions"),
)


TOC_FIELD_CHECKS: tuple[tuple[str, str, str], ...] = (
    ("toc.enabled", "docx.toc.enabled", "DOCX TOC enabled state"),
    ("toc.title", "docx.toc.title", "DOCX TOC title"),
    ("toc.include_levels", "docx.toc.include_levels", "DOCX TOC included heading levels"),
    ("toc.show_page_numbers", "docx.toc.show_page_numbers", "DOCX TOC page-number display"),
    ("toc.right_align_page_numbers", "docx.toc.right_align_page_numbers", "DOCX TOC right-aligned page numbers"),
    ("toc.use_hyperlinks", "docx.toc.use_hyperlinks", "DOCX TOC hyperlinks"),
    ("toc.update_fields_on_open", "docx.toc.update_fields_on_open", "DOCX TOC field update-on-open"),
)


def _disabled_quality_issue(issue_id: str, check_key: str, title: str, profile_rule_ref: str) -> QualityIssue:
    return QualityIssue(
        issue_id=issue_id,
        status="pass",
        severity="info",
        check_key=check_key,
        title=title,
        profile_rule_ref=profile_rule_ref,
        details={"disabled_by_profile": True},
    )


def _quality_issue_id_from_check_key(check_key: str) -> str:
    return check_key.replace(".", "_")


def _field_issue_from_mismatches(
    field_path: str,
    check_key: str,
    title: str,
    mismatches: list[str],
    *,
    location: str,
    fixable: bool,
) -> QualityIssue:
    if mismatches:
        return QualityIssue(
            issue_id=_quality_issue_id_from_check_key(check_key),
            status="warning",
            severity="high",
            check_key=check_key,
            title=f"{title} needs review.",
            description="; ".join(mismatches[:20]),
            profile_rule_ref=field_path,
            location=location,
            recommendation=f"Apply profile rule `{field_path}` before final delivery.",
            fixable=fixable,
            details={"mismatch_count": len(mismatches)},
        )
    return QualityIssue(
        issue_id=_quality_issue_id_from_check_key(check_key),
        status="pass",
        severity="info",
        check_key=check_key,
        title=f"{title} matches the profile.",
        profile_rule_ref=field_path,
        location=location,
    )


def inspect_pdf_quality(path: Path) -> list[QualityIssue]:
    try:
        data = path.read_bytes()
    except OSError as exc:
        return [
            _pdf_openability_issue(False, f"PDF cannot be read from disk: {exc}"),
            _pdf_page_count_issue(0, readable=False),
            _pdf_text_extractability_issue(False, readable=False),
            _pdf_blank_pages_issue(0, has_text=False, readable=False),
        ]

    readable = data.startswith(b"%PDF") and b"%%EOF" in data[-2048:]
    pypdf_info = _read_pdf_text_info(data) if readable else None
    page_count = pypdf_info[0] if pypdf_info is not None else (_count_pdf_pages(data) if readable else 0)
    has_text = pypdf_info[1] if pypdf_info is not None else (_pdf_has_literal_text(data) if readable else False)
    return [
        _pdf_openability_issue(readable, None),
        _pdf_page_count_issue(page_count, readable=readable),
        _pdf_text_extractability_issue(has_text, readable=readable),
        _pdf_blank_pages_issue(page_count, has_text, readable=readable),
    ]


def _pdf_openability_issue(readable: bool, detail: str | None) -> QualityIssue:
    if readable:
        return QualityIssue(
            issue_id="pdf_openability",
            status="pass",
            severity="info",
            check_key="pdf.openability",
            title="PDF file has a readable PDF envelope.",
        )
    return QualityIssue(
        issue_id="pdf_openability",
        status="fail",
        severity="high",
        check_key="pdf.openability",
        title="PDF file cannot be opened by the lightweight checker.",
        description=detail or "The file does not look like a complete PDF.",
        recommendation="Regenerate the PDF and inspect the export tool logs.",
        fixable=False,
    )


def _pdf_page_count_issue(page_count: int, readable: bool) -> QualityIssue:
    if not readable or page_count <= 0:
        return QualityIssue(
            issue_id="pdf_page_count",
            status="fail",
            severity="high",
            check_key="pdf.page_count",
            title="PDF page count could not be confirmed.",
            description=f"Detected page_count={page_count}.",
            recommendation="Regenerate the PDF and confirm it contains at least one page.",
            details={"page_count": page_count},
            fixable=False,
        )
    return QualityIssue(
        issue_id="pdf_page_count",
        status="pass",
        severity="info",
        check_key="pdf.page_count",
        title="PDF page count is greater than zero.",
        details={"page_count": page_count},
    )


def _pdf_text_extractability_issue(has_text: bool, readable: bool) -> QualityIssue:
    if readable and has_text:
        return QualityIssue(
            issue_id="pdf_text_extractability",
            status="pass",
            severity="info",
            check_key="pdf.text_extractability",
            title="PDF contains extractable text.",
        )
    return QualityIssue(
        issue_id="pdf_text_extractability",
        status="fail",
        severity="high",
        check_key="pdf.text_extractability",
        title="PDF text extractability could not be confirmed.",
        description="The lightweight checker did not find literal text drawing operators.",
        recommendation="Inspect the PDF with codex-pdf-inspect or regenerate from DOCX.",
        fixable=False,
    )


def _pdf_blank_pages_issue(page_count: int, has_text: bool, readable: bool) -> QualityIssue:
    if not readable:
        return QualityIssue(
            issue_id="pdf_blank_pages",
            status="unsupported",
            check_key="pdf.blank_pages",
            title="Blank-page check cannot run because the PDF is unreadable.",
            recommendation="Regenerate the PDF first.",
        )
    if page_count > 0 and not has_text:
        return QualityIssue(
            issue_id="pdf_blank_pages",
            status="warning",
            check_key="pdf.blank_pages",
            title="PDF may contain blank or image-only pages.",
            description="A page was detected, but no literal text was found.",
            recommendation="Open the PDF and review whether pages are blank or scanned-only.",
            fixable=False,
        )
    return QualityIssue(
        issue_id="pdf_blank_pages",
        status="pass",
        severity="info",
        check_key="pdf.blank_pages",
        title="No obvious blank-page warning detected by the lightweight checker.",
    )


def _inspect_margins(document: Document, profile: FormatProfile) -> QualityIssue:
    expected = profile.page.margins_cm
    mismatches: list[str] = []
    for index, section in enumerate(document.sections):
        actual = {
            "top": section.top_margin.cm,
            "bottom": section.bottom_margin.cm,
            "left": section.left_margin.cm,
            "right": section.right_margin.cm,
            "gutter": section.gutter.cm if section.gutter else 0,
        }
        for key, expected_value in expected.model_dump().items():
            if not _close_cm(actual[key], expected_value):
                mismatches.append(f"section[{index}].{key}: expected {expected_value:.2f} cm, found {actual[key]:.2f} cm")
    if mismatches:
        return QualityIssue(
            issue_id="docx_page_margins",
            status="fail",
            severity="high",
            check_key="docx.page.margins",
            title="DOCX page margins do not match the profile.",
            description="; ".join(mismatches),
            profile_rule_ref="page.margins_cm",
            location="sections",
            recommendation="Reapply profile page settings.",
            fixable=True,
        )
    return QualityIssue(
        issue_id="docx_page_margins",
        status="pass",
        severity="info",
        check_key="docx.page.margins",
        title="DOCX page margins match the profile.",
        profile_rule_ref="page.margins_cm",
    )


def _inspect_page_setup(document: Document, profile: FormatProfile) -> QualityIssue:
    expected_width, expected_height = _expected_page_size_cm(profile)
    mismatches: list[str] = []
    for index, section in enumerate(document.sections):
        if not _close_cm(section.page_width.cm, expected_width, tolerance=0.15):
            mismatches.append(f"section[{index}].width expected {expected_width:.2f} cm, found {section.page_width.cm:.2f} cm")
        if not _close_cm(section.page_height.cm, expected_height, tolerance=0.15):
            mismatches.append(f"section[{index}].height expected {expected_height:.2f} cm, found {section.page_height.cm:.2f} cm")
    if mismatches:
        return QualityIssue(
            issue_id="docx_page_setup",
            status="fail",
            severity="high",
            check_key="docx.page.setup",
            title="DOCX page size or orientation does not match the profile.",
            description="; ".join(mismatches),
            profile_rule_ref="page.size; page.orientation",
            location="sections",
            recommendation="Reapply profile page size and orientation settings.",
            fixable=True,
        )
    return QualityIssue(
        issue_id="docx_page_setup",
        status="pass",
        severity="info",
        check_key="docx.page.setup",
        title="DOCX page size and orientation match the profile.",
        profile_rule_ref="page.size; page.orientation",
    )


def _expected_page_size_cm(profile: FormatProfile) -> tuple[float, float]:
    if profile.page.size == "A4":
        width_cm, height_cm = 21.0, 29.7
    else:
        width_cm, height_cm = 21.59, 27.94
    if profile.page.orientation == "landscape":
        return height_cm, width_cm
    return width_cm, height_cm


def _inspect_header_footer(document: Document, profile: FormatProfile, features: OoxmlDocumentFeatures | None) -> QualityIssue:
    mismatches: list[str] = []
    expected_header = (profile.header_footer.header_text or "").strip()
    expected_footer = (profile.header_footer.footer_text or "").strip()
    if features is not None and features.even_and_odd_headers != profile.header_footer.different_odd_even:
        mismatches.append(
            "document settings evenAndOddHeaders does not match profile expectation"
        )
    for index, section in enumerate(document.sections):
        if bool(section.different_first_page_header_footer) != profile.header_footer.different_first_page:
            mismatches.append(
                f"section[{index}] different_first_page_header_footer does not match the profile"
            )
        header_text = _container_text(section.header)
        if expected_header and expected_header not in header_text:
            mismatches.append(f"section[{index}] missing header text")
        footer_text = _container_text(section.footer)
        if expected_footer and expected_footer not in footer_text:
            mismatches.append(f"section[{index}] missing footer text")
        has_page = _container_has_page_field(section.footer)
        if profile.header_footer.footer_page_number and not has_page:
            mismatches.append(f"section[{index}] missing footer page number")
        if not profile.header_footer.footer_page_number and has_page:
            mismatches.append(f"section[{index}] has unexpected footer page number")
        if profile.header_footer.different_first_page:
            first_header_text = _container_text(section.first_page_header)
            if expected_header and expected_header not in first_header_text:
                mismatches.append(f"section[{index}] missing first-page header text")
            first_footer_text = _container_text(section.first_page_footer)
            if expected_footer and expected_footer not in first_footer_text:
                mismatches.append(f"section[{index}] missing first-page footer text")
            first_page_has_page = _container_has_page_field(section.first_page_footer)
            if profile.header_footer.footer_page_number and not first_page_has_page:
                mismatches.append(f"section[{index}] missing first-page footer page number")
        if profile.header_footer.different_odd_even:
            even_header_text = _container_text(section.even_page_header)
            if expected_header and expected_header not in even_header_text:
                mismatches.append(f"section[{index}] missing even-page header text")
            even_footer_text = _container_text(section.even_page_footer)
            if expected_footer and expected_footer not in even_footer_text:
                mismatches.append(f"section[{index}] missing even-page footer text")
            even_page_has_page = _container_has_page_field(section.even_page_footer)
            if profile.header_footer.footer_page_number and not even_page_has_page:
                mismatches.append(f"section[{index}] missing even-page footer page number")
    if mismatches:
        return QualityIssue(
            issue_id="docx_header_footer",
            status="fail",
            severity="high",
            check_key="docx.header_footer",
            title="DOCX header/footer rules do not match the profile.",
            description="; ".join(mismatches),
            profile_rule_ref="header_footer",
            location="sections",
            recommendation="Apply supported header text and footer page-number rules.",
            fixable=True,
        )
    return QualityIssue(
        issue_id="docx_header_footer",
        status="pass",
        severity="info",
        check_key="docx.header_footer",
        title="DOCX header/footer rules match supported profile checks.",
        profile_rule_ref="header_footer",
    )


def _inspect_header_footer_fields(
    document: Document,
    profile: FormatProfile,
    features: OoxmlDocumentFeatures | None,
) -> list[QualityIssue]:
    mismatches_by_field = {field_path: [] for field_path, _, _ in HEADER_FOOTER_FIELD_CHECKS}
    expected_header = (profile.header_footer.header_text or "").strip()
    expected_footer = (profile.header_footer.footer_text or "").strip()
    if features is not None and features.even_and_odd_headers != profile.header_footer.different_odd_even:
        mismatches_by_field["header_footer.different_odd_even"].append(
            "document settings evenAndOddHeaders does not match profile expectation"
        )
    for index, section in enumerate(document.sections):
        if bool(section.different_first_page_header_footer) != profile.header_footer.different_first_page:
            mismatches_by_field["header_footer.different_first_page"].append(
                f"section[{index}] different_first_page_header_footer does not match the profile"
            )
        header_text = _container_text(section.header)
        if expected_header and expected_header not in header_text:
            mismatches_by_field["header_footer.header_text"].append(f"section[{index}] missing header text")
        if not _container_alignment_matches(section.header, profile.header_footer.header_alignment):
            mismatches_by_field["header_footer.header_alignment"].append(
                f"section[{index}] header alignment expected {profile.header_footer.header_alignment}"
            )
        footer_text = _container_text(section.footer)
        if expected_footer and expected_footer not in footer_text:
            mismatches_by_field["header_footer.footer_text"].append(f"section[{index}] missing footer text")
        if not _container_alignment_matches(section.footer, profile.header_footer.footer_alignment):
            mismatches_by_field["header_footer.footer_alignment"].append(
                f"section[{index}] footer alignment expected {profile.header_footer.footer_alignment}"
            )
        if profile.header_footer.different_first_page:
            first_header_text = _container_text(section.first_page_header)
            if expected_header and expected_header not in first_header_text:
                mismatches_by_field["header_footer.header_text"].append(f"section[{index}] missing first-page header text")
            if not _container_alignment_matches(section.first_page_header, profile.header_footer.header_alignment):
                mismatches_by_field["header_footer.header_alignment"].append(
                    f"section[{index}] first-page header alignment expected {profile.header_footer.header_alignment}"
                )
            first_footer_text = _container_text(section.first_page_footer)
            if expected_footer and expected_footer not in first_footer_text:
                mismatches_by_field["header_footer.footer_text"].append(f"section[{index}] missing first-page footer text")
            if not _container_alignment_matches(section.first_page_footer, profile.header_footer.footer_alignment):
                mismatches_by_field["header_footer.footer_alignment"].append(
                    f"section[{index}] first-page footer alignment expected {profile.header_footer.footer_alignment}"
                )
        if profile.header_footer.different_odd_even:
            even_header_text = _container_text(section.even_page_header)
            if expected_header and expected_header not in even_header_text:
                mismatches_by_field["header_footer.header_text"].append(f"section[{index}] missing even-page header text")
            if not _container_alignment_matches(section.even_page_header, profile.header_footer.header_alignment):
                mismatches_by_field["header_footer.header_alignment"].append(
                    f"section[{index}] even-page header alignment expected {profile.header_footer.header_alignment}"
                )
            even_footer_text = _container_text(section.even_page_footer)
            if expected_footer and expected_footer not in even_footer_text:
                mismatches_by_field["header_footer.footer_text"].append(f"section[{index}] missing even-page footer text")
            if not _container_alignment_matches(section.even_page_footer, profile.header_footer.footer_alignment):
                mismatches_by_field["header_footer.footer_alignment"].append(
                    f"section[{index}] even-page footer alignment expected {profile.header_footer.footer_alignment}"
                )
    issues: list[QualityIssue] = []
    for field_path, check_key, title in HEADER_FOOTER_FIELD_CHECKS:
        field_mismatches = mismatches_by_field[field_path]
        if field_mismatches:
            issues.append(
                QualityIssue(
                    issue_id=_quality_issue_id_from_check_key(check_key),
                    status="warning",
                    severity="high",
                    check_key=check_key,
                    title=f"{title} needs review.",
                    description="; ".join(field_mismatches),
                    profile_rule_ref=field_path,
                    location="sections",
                    recommendation=f"Apply profile rule `{field_path}` to every section.",
                    fixable=True,
                    details={"mismatch_count": len(field_mismatches)},
                )
            )
            continue
        issues.append(
            QualityIssue(
                issue_id=_quality_issue_id_from_check_key(check_key),
                status="pass",
                severity="info",
                check_key=check_key,
                title=f"{title} matches the profile.",
                profile_rule_ref=field_path,
                location=f"{len(document.sections)} section(s)",
            )
        )
    return issues


def _disabled_header_footer_field_issues(reason: str) -> list[QualityIssue]:
    return [
        QualityIssue(
            issue_id=_quality_issue_id_from_check_key(check_key),
            status="pass",
            severity="info",
            check_key=check_key,
            title=f"{title} inspection is delegated.",
            profile_rule_ref=field_path,
            details={"disabled_by_profile": True, "reason": reason},
        )
        for field_path, check_key, title in HEADER_FOOTER_FIELD_CHECKS
    ]


def _container_text(container) -> str:
    return "\n".join(paragraph.text.strip() for paragraph in container.paragraphs if paragraph.text.strip())


def _container_alignment_matches(container, alignment: TextAlignment) -> bool:
    expected = _alignment_value(alignment)
    selected = [paragraph for paragraph in container.paragraphs if paragraph.text.strip() or "PAGE" in paragraph._p.xml.upper()]
    if not selected:
        return True
    return all(paragraph.alignment == expected for paragraph in selected)


def _count_pdf_pages(data: bytes) -> int:
    text = data.decode("latin-1", errors="ignore")
    type_page_refs = len(re.findall(r"/Type\s*/Page\b", text))
    if type_page_refs:
        return type_page_refs
    count_values = [int(match) for match in re.findall(r"/Count\s+(\d+)", text)]
    return max(count_values, default=0)


def _pdf_has_literal_text(data: bytes) -> bool:
    text = data.decode("latin-1", errors="ignore")
    return bool(re.search(r"\([^()]{2,}\)\s*T[Jj]", text) or re.search(r"<[0-9A-Fa-f]{4,}>\s*T[Jj]", text))


def _read_pdf_text_info(data: bytes) -> tuple[int, bool] | None:
    try:
        reader = PdfReader(BytesIO(data))
        page_count = len(reader.pages)
        text_chars = 0
        for page in reader.pages:
            text_chars += len((page.extract_text() or "").strip())
        return page_count, text_chars > 0
    except Exception:
        return None


def _inspect_body_style(document: Document, structure: DocumentStructure, profile: FormatProfile) -> QualityIssue:
    paragraphs = [
        paragraph
        for paragraph, classification in _paragraphs_with_classifications(document, structure)
        if classification.role == ParagraphRole.BODY
    ]
    if not paragraphs:
        return QualityIssue(
            issue_id="docx_body_style",
            status="unsupported",
            check_key="docx.body.style",
            title="No body paragraph could be selected for style inspection.",
            profile_rule_ref="body",
            recommendation="Review body text manually.",
        )

    mismatches: list[str] = []
    for paragraph in paragraphs:
        location = _paragraph_location(document, paragraph)
        mismatches.extend(f"{location}: {item}" for item in _body_paragraph_style_mismatches(paragraph, profile))

    if mismatches:
        return QualityIssue(
            issue_id="docx_body_style",
            status="warning",
            check_key="docx.body.style",
            title="DOCX body paragraph style needs review.",
            description="; ".join(mismatches[:20]),
            profile_rule_ref="body",
            location="body paragraphs",
            recommendation="Apply the body paragraph style from the selected profile.",
            fixable=True,
        )
    return QualityIssue(
        issue_id="docx_body_style",
        status="pass",
        severity="info",
        check_key="docx.body.style",
        title="DOCX body paragraph style matches the profile.",
        profile_rule_ref="body",
        location=f"{len(paragraphs)} body paragraph(s)",
    )


def _inspect_body_style_fields(
    document: Document,
    structure: DocumentStructure,
    profile: FormatProfile,
) -> list[QualityIssue]:
    paragraphs = [
        paragraph
        for paragraph, classification in _paragraphs_with_classifications(document, structure)
        if classification.role == ParagraphRole.BODY
    ]
    if not paragraphs:
        return [
            QualityIssue(
                issue_id=_quality_issue_id_from_check_key(check_key),
                status="unsupported",
                check_key=check_key,
                title=f"{title} could not be selected for inspection.",
                profile_rule_ref=field_path,
                recommendation="Review body text manually.",
            )
            for field_path, check_key, title in BODY_FIELD_CHECKS
        ]

    mismatches_by_field = {field_path: [] for field_path, _, _ in BODY_FIELD_CHECKS}
    disabled_fields = _disabled_body_field_paths(profile)
    for paragraph in paragraphs:
        location = _paragraph_location(document, paragraph)
        for field_path, mismatches in _body_paragraph_field_mismatches(paragraph, profile).items():
            mismatches_by_field[field_path].extend(f"{location}: {item}" for item in mismatches)

    issues: list[QualityIssue] = []
    for field_path, check_key, title in BODY_FIELD_CHECKS:
        if field_path in disabled_fields:
            issues.append(
                QualityIssue(
                    issue_id=_quality_issue_id_from_check_key(check_key),
                    status="pass",
                    severity="info",
                    check_key=check_key,
                    title=f"{title} inspection is disabled by the profile.",
                    profile_rule_ref=field_path,
                    location=f"{len(paragraphs)} body paragraph(s)",
                    details={"disabled_by_profile": True},
                )
            )
            continue
        field_mismatches = mismatches_by_field[field_path]
        if field_mismatches:
            issues.append(
                QualityIssue(
                    issue_id=_quality_issue_id_from_check_key(check_key),
                    status="warning",
                    check_key=check_key,
                    title=f"{title} needs review.",
                    description="; ".join(field_mismatches[:20]),
                    profile_rule_ref=field_path,
                    location="body paragraphs",
                    recommendation=f"Apply profile rule `{field_path}` to every body paragraph.",
                    fixable=True,
                    details={"mismatch_count": len(field_mismatches)},
                )
            )
            continue
        issues.append(
            QualityIssue(
                issue_id=_quality_issue_id_from_check_key(check_key),
                status="pass",
                severity="info",
                check_key=check_key,
                title=f"{title} matches the profile.",
                profile_rule_ref=field_path,
                location=f"{len(paragraphs)} body paragraph(s)",
            )
        )
    return issues


def _body_paragraph_style_mismatches(paragraph, profile: FormatProfile) -> list[str]:
    mismatches: list[str] = []
    fmt = paragraph.paragraph_format
    expected_indent = profile.body.first_line_indent_chars * 0.37
    actual_indent = fmt.first_line_indent.cm if fmt.first_line_indent else 0
    if not _close_cm(actual_indent, expected_indent):
        mismatches.append(f"first line indent expected {expected_indent:.2f} cm, found {actual_indent:.2f} cm")
    if profile.quality.check_line_spacing and fmt.line_spacing != profile.body.line_spacing:
        mismatches.append(f"line spacing expected {profile.body.line_spacing}, found {fmt.line_spacing}")
    expected_before = profile.body.space_before_pt
    actual_before = _spacing_pt(fmt.space_before)
    if profile.quality.check_line_spacing and abs((actual_before or 0) - expected_before) > 0.1:
        mismatches.append(f"space before expected {expected_before:g} pt, found {actual_before if actual_before is not None else 'auto'}")
    expected_after = profile.body.space_after_pt
    actual_after = _spacing_pt(fmt.space_after)
    if profile.quality.check_line_spacing and abs((actual_after or 0) - expected_after) > 0.1:
        mismatches.append(f"space after expected {expected_after:g} pt, found {actual_after if actual_after is not None else 'auto'}")
    if profile.document_grid.snap_to_grid and not _paragraph_snap_to_grid(paragraph):
        mismatches.append("body paragraph is not snapped to document grid")
    run = next((r for r in paragraph.runs if r.text.strip()), None)
    if run is None:
        mismatches.append("body paragraph has no text run")
    elif profile.quality.check_fonts:
        mismatches.extend(_run_style_mismatches(run, profile.body.font, "body run"))
    return mismatches


def _body_paragraph_field_mismatches(paragraph, profile: FormatProfile) -> dict[str, list[str]]:
    mismatches: dict[str, list[str]] = {field_path: [] for field_path, _, _ in BODY_FIELD_CHECKS}
    fmt = paragraph.paragraph_format
    expected_indent = profile.body.first_line_indent_chars * 0.37
    actual_indent = fmt.first_line_indent.cm if fmt.first_line_indent else 0
    if not _close_cm(actual_indent, expected_indent):
        mismatches["body.first_line_indent_chars"].append(
            f"first line indent expected {expected_indent:.2f} cm, found {actual_indent:.2f} cm"
        )
    if profile.quality.check_line_spacing and fmt.line_spacing != profile.body.line_spacing:
        mismatches["body.line_spacing"].append(
            f"line spacing expected {profile.body.line_spacing}, found {fmt.line_spacing}"
        )
    expected_before = profile.body.space_before_pt
    actual_before = _spacing_pt(fmt.space_before)
    if profile.quality.check_line_spacing and abs((actual_before or 0) - expected_before) > 0.1:
        mismatches["body.space_before_pt"].append(
            f"space before expected {expected_before:g} pt, found {actual_before if actual_before is not None else 'auto'}"
        )
    expected_after = profile.body.space_after_pt
    actual_after = _spacing_pt(fmt.space_after)
    if profile.quality.check_line_spacing and abs((actual_after or 0) - expected_after) > 0.1:
        mismatches["body.space_after_pt"].append(
            f"space after expected {expected_after:g} pt, found {actual_after if actual_after is not None else 'auto'}"
        )
    if paragraph.alignment != _alignment_value(profile.body.alignment):
        mismatches["body.alignment"].append(f"alignment expected {profile.body.alignment}")
    run = next((r for r in paragraph.runs if r.text.strip()), None)
    if run is None:
        for field_path in ("body.font.chinese", "body.font.latin", "body.font.size_pt", "body.font.color"):
            mismatches[field_path].append("body paragraph has no text run")
        return mismatches
    if not profile.quality.check_fonts:
        return mismatches
    east_asia = run._element.rPr.rFonts.get(qn("w:eastAsia")) if run._element.rPr is not None and run._element.rPr.rFonts is not None else None
    if east_asia != profile.body.font.chinese:
        mismatches["body.font.chinese"].append(
            f"Chinese font expected {profile.body.font.chinese}, found {east_asia or 'auto/inherited'}"
        )
    if run.font.name != profile.body.font.latin:
        mismatches["body.font.latin"].append(
            f"Latin font expected {profile.body.font.latin}, found {run.font.name or 'auto/inherited'}"
        )
    actual_size = run.font.size.pt if run.font.size else None
    if actual_size is None or abs(actual_size - profile.body.font.size_pt) > 0.1:
        found = f"{actual_size:g} pt" if actual_size else "inherited/auto"
        mismatches["body.font.size_pt"].append(f"size expected {profile.body.font.size_pt:g} pt, found {found}")
    if not _run_color_matches(run, profile.body.font.color):
        actual = _run_color_value(run)
        mismatches["body.font.color"].append(f"color expected {profile.body.font.color}, found {actual or 'auto/inherited'}")
    return mismatches


def _disabled_body_field_paths(profile: FormatProfile) -> set[str]:
    disabled: set[str] = set()
    if not profile.quality.check_fonts:
        disabled.update({"body.font.chinese", "body.font.latin", "body.font.size_pt", "body.font.color"})
    if not profile.quality.check_line_spacing:
        disabled.update({"body.line_spacing", "body.space_before_pt", "body.space_after_pt"})
    return disabled


def _inspect_heading_style(document: Document, structure: DocumentStructure, profile: FormatProfile) -> QualityIssue:
    selected = _selected_heading_paragraphs(document, structure)
    if not selected:
        return QualityIssue(
            issue_id="docx_heading_style",
            status="unsupported",
            check_key="docx.heading.style",
            title="No heading paragraph could be selected for style inspection.",
            profile_rule_ref="headings",
            recommendation="Review heading styles manually.",
        )

    mismatches: list[str] = []
    for paragraph, classification in selected:
        heading_level = classification.heading_level if classification and classification.heading_level else 1
        heading_rule = _heading_rule(profile, heading_level)
        location = _paragraph_location(document, paragraph)
        mismatches.extend(f"{location}: {item}" for item in _heading_paragraph_style_mismatches(paragraph, heading_rule, profile))

    if mismatches:
        return QualityIssue(
            issue_id="docx_heading_style",
            status="warning",
            check_key="docx.heading.style",
            title="DOCX heading style needs review.",
            description="; ".join(mismatches[:20]),
            profile_rule_ref="headings",
            location="heading paragraphs",
            recommendation="Apply heading style rules from the selected profile.",
            fixable=True,
        )
    return QualityIssue(
        issue_id="docx_heading_style",
        status="pass",
        severity="info",
        check_key="docx.heading.style",
        title="DOCX heading style matches the profile.",
        profile_rule_ref="headings",
        location=f"{len(selected)} heading paragraph(s)",
    )


def _inspect_heading_style_fields(
    document: Document,
    structure: DocumentStructure,
    profile: FormatProfile,
) -> list[QualityIssue]:
    selected = _selected_heading_paragraphs(document, structure)
    if not selected:
        return [
            QualityIssue(
                issue_id=_quality_issue_id_from_check_key(check_key),
                status="unsupported",
                check_key=check_key,
                title=f"{title} could not be selected for inspection.",
                profile_rule_ref=field_path,
                recommendation="Review heading styles manually.",
            )
            for field_path, check_key, title in HEADING_FIELD_CHECKS
        ]

    mismatches_by_field = {field_path: [] for field_path, _, _ in HEADING_FIELD_CHECKS}
    disabled_fields = _disabled_heading_field_paths(profile)
    for paragraph, classification in selected:
        heading_level = classification.heading_level if classification and classification.heading_level else 1
        heading_rule = _heading_rule(profile, heading_level)
        location = _paragraph_location(document, paragraph)
        for field_path, mismatches in _heading_paragraph_field_mismatches(paragraph, heading_rule, profile).items():
            mismatches_by_field[field_path].extend(f"{location}: {item}" for item in mismatches)

    issues: list[QualityIssue] = []
    for field_path, check_key, title in HEADING_FIELD_CHECKS:
        if field_path in disabled_fields:
            issues.append(
                QualityIssue(
                    issue_id=_quality_issue_id_from_check_key(check_key),
                    status="pass",
                    severity="info",
                    check_key=check_key,
                    title=f"{title} inspection is disabled by the profile.",
                    profile_rule_ref=field_path,
                    location=f"{len(selected)} heading paragraph(s)",
                    details={"disabled_by_profile": True},
                )
            )
            continue
        field_mismatches = mismatches_by_field[field_path]
        if field_mismatches:
            issues.append(
                QualityIssue(
                    issue_id=_quality_issue_id_from_check_key(check_key),
                    status="warning",
                    check_key=check_key,
                    title=f"{title} needs review.",
                    description="; ".join(field_mismatches[:20]),
                    profile_rule_ref=field_path,
                    location="heading paragraphs",
                    recommendation=f"Apply profile rule `{field_path}` to every heading paragraph.",
                    fixable=True,
                    details={"mismatch_count": len(field_mismatches)},
                )
            )
            continue
        issues.append(
            QualityIssue(
                issue_id=_quality_issue_id_from_check_key(check_key),
                status="pass",
                severity="info",
                check_key=check_key,
                title=f"{title} matches the profile.",
                profile_rule_ref=field_path,
                location=f"{len(selected)} heading paragraph(s)",
            )
        )
    return issues


def _selected_heading_paragraphs(document: Document, structure: DocumentStructure):
    return [
        (paragraph, classification)
        for paragraph, classification in _paragraphs_with_classifications(document, structure)
        if classification.role
        in {
            ParagraphRole.DOCUMENT_TITLE,
            ParagraphRole.HEADING,
            ParagraphRole.REFERENCE_HEADING,
            ParagraphRole.ACKNOWLEDGEMENT_HEADING,
        }
    ]


def _heading_paragraph_style_mismatches(paragraph, heading_rule, profile: FormatProfile) -> list[str]:
    run = next((r for r in paragraph.runs if r.text.strip()), None)
    mismatches: list[str] = []
    if paragraph.alignment != _alignment_value(heading_rule.alignment):
        mismatches.append(f"heading alignment expected {heading_rule.alignment}")
    actual_line_spacing = paragraph.paragraph_format.line_spacing
    expected_line_spacing = heading_rule.line_spacing or profile.body.line_spacing
    if profile.quality.check_line_spacing and actual_line_spacing != expected_line_spacing:
        mismatches.append(f"heading line spacing expected {expected_line_spacing}, found {actual_line_spacing}")
    actual_before = _spacing_pt(paragraph.paragraph_format.space_before)
    if profile.quality.check_line_spacing and abs((actual_before or 0) - heading_rule.space_before_pt) > 0.1:
        mismatches.append(f"heading space before expected {heading_rule.space_before_pt:g} pt, found {actual_before if actual_before is not None else 'auto'}")
    actual_after = _spacing_pt(paragraph.paragraph_format.space_after)
    if profile.quality.check_line_spacing and abs((actual_after or 0) - heading_rule.space_after_pt) > 0.1:
        mismatches.append(f"heading space after expected {heading_rule.space_after_pt:g} pt, found {actual_after if actual_after is not None else 'auto'}")
    actual_indent = paragraph.paragraph_format.first_line_indent.cm if paragraph.paragraph_format.first_line_indent else 0
    expected_indent = heading_rule.first_line_indent_chars * 0.37
    if not _close_cm(actual_indent, expected_indent):
        mismatches.append(f"heading first line indent expected {expected_indent:.2f} cm, found {actual_indent:.2f} cm")
    if bool(paragraph.paragraph_format.keep_with_next) != heading_rule.keep_with_next:
        mismatches.append(f"heading keep_with_next expected {heading_rule.keep_with_next}")
    if bool(paragraph.paragraph_format.page_break_before) != heading_rule.page_break_before:
        mismatches.append(f"heading page_break_before expected {heading_rule.page_break_before}")
    if profile.document_grid.snap_to_grid and not _paragraph_snap_to_grid(paragraph):
        mismatches.append("heading paragraph is not snapped to document grid")
    if run is None:
        mismatches.append("heading paragraph has no text run")
    elif profile.quality.check_fonts:
        mismatches.extend(_run_style_mismatches(run, heading_rule.font, "heading run"))
    return mismatches


def _heading_paragraph_field_mismatches(paragraph, heading_rule, profile: FormatProfile) -> dict[str, list[str]]:
    mismatches: dict[str, list[str]] = {field_path: [] for field_path, _, _ in HEADING_FIELD_CHECKS}
    if paragraph.alignment != _alignment_value(heading_rule.alignment):
        mismatches["headings.alignment"].append(f"alignment expected {heading_rule.alignment}")
    actual_line_spacing = paragraph.paragraph_format.line_spacing
    expected_line_spacing = heading_rule.line_spacing or profile.body.line_spacing
    if profile.quality.check_line_spacing and actual_line_spacing != expected_line_spacing:
        mismatches["headings.line_spacing"].append(f"line spacing expected {expected_line_spacing}, found {actual_line_spacing}")
    actual_before = _spacing_pt(paragraph.paragraph_format.space_before)
    if profile.quality.check_line_spacing and abs((actual_before or 0) - heading_rule.space_before_pt) > 0.1:
        mismatches["headings.space_before_pt"].append(
            f"space before expected {heading_rule.space_before_pt:g} pt, found {actual_before if actual_before is not None else 'auto'}"
        )
    actual_after = _spacing_pt(paragraph.paragraph_format.space_after)
    if profile.quality.check_line_spacing and abs((actual_after or 0) - heading_rule.space_after_pt) > 0.1:
        mismatches["headings.space_after_pt"].append(
            f"space after expected {heading_rule.space_after_pt:g} pt, found {actual_after if actual_after is not None else 'auto'}"
        )
    actual_indent = paragraph.paragraph_format.first_line_indent.cm if paragraph.paragraph_format.first_line_indent else 0
    expected_indent = heading_rule.first_line_indent_chars * 0.37
    if not _close_cm(actual_indent, expected_indent):
        mismatches["headings.first_line_indent_chars"].append(
            f"first line indent expected {expected_indent:.2f} cm, found {actual_indent:.2f} cm"
        )
    if bool(paragraph.paragraph_format.keep_with_next) != heading_rule.keep_with_next:
        mismatches["headings.pagination"].append(f"keep_with_next expected {heading_rule.keep_with_next}")
    if bool(paragraph.paragraph_format.page_break_before) != heading_rule.page_break_before:
        mismatches["headings.pagination"].append(f"page_break_before expected {heading_rule.page_break_before}")
    run = next((r for r in paragraph.runs if r.text.strip()), None)
    if run is None:
        for field_path in (
            "headings.font.chinese",
            "headings.font.latin",
            "headings.font.size_pt",
            "headings.font.color",
            "headings.font.weight",
        ):
            mismatches[field_path].append("heading paragraph has no text run")
        return mismatches
    if not profile.quality.check_fonts:
        return mismatches
    east_asia = run._element.rPr.rFonts.get(qn("w:eastAsia")) if run._element.rPr is not None and run._element.rPr.rFonts is not None else None
    if east_asia != heading_rule.font.chinese:
        mismatches["headings.font.chinese"].append(
            f"Chinese font expected {heading_rule.font.chinese}, found {east_asia or 'auto/inherited'}"
        )
    if run.font.name != heading_rule.font.latin:
        mismatches["headings.font.latin"].append(
            f"Latin font expected {heading_rule.font.latin}, found {run.font.name or 'auto/inherited'}"
        )
    actual_size = run.font.size.pt if run.font.size else None
    if actual_size is None or abs(actual_size - heading_rule.font.size_pt) > 0.1:
        found = f"{actual_size:g} pt" if actual_size else "inherited/auto"
        mismatches["headings.font.size_pt"].append(f"size expected {heading_rule.font.size_pt:g} pt, found {found}")
    if not _run_color_matches(run, heading_rule.font.color):
        actual = _run_color_value(run)
        mismatches["headings.font.color"].append(f"color expected {heading_rule.font.color}, found {actual or 'auto/inherited'}")
    actual_bold = bool(run.font.bold)
    expected_bold = heading_rule.font.weight == "bold"
    if actual_bold != expected_bold:
        mismatches["headings.font.weight"].append(f"bold expected {expected_bold}, found {actual_bold}")
    return mismatches


def _disabled_heading_field_paths(profile: FormatProfile) -> set[str]:
    disabled: set[str] = set()
    if not profile.quality.check_fonts:
        disabled.update(
            {
                "headings.font.chinese",
                "headings.font.latin",
                "headings.font.size_pt",
                "headings.font.color",
                "headings.font.weight",
            }
        )
    if not profile.quality.check_line_spacing:
        disabled.update({"headings.line_spacing", "headings.space_before_pt", "headings.space_after_pt"})
    return disabled


def _disabled_heading_style_field_issues() -> list[QualityIssue]:
    return [
        QualityIssue(
            issue_id=_quality_issue_id_from_check_key(check_key),
            status="pass",
            severity="info",
            check_key=check_key,
            title=f"{title} inspection is disabled by the profile.",
            profile_rule_ref=field_path,
            details={"disabled_by_profile": True},
        )
        for field_path, check_key, title in HEADING_FIELD_CHECKS
    ]


def _inspect_table_borders(document: Document, profile: FormatProfile) -> QualityIssue:
    if not document.tables:
        return QualityIssue(
            issue_id="docx_table_borders",
            status="pass",
            severity="info",
            check_key="docx.table.borders",
            title="DOCX table-border check is not applicable because no tables were detected.",
            profile_rule_ref="table",
        )
    if profile.table.border_style == "custom":
        return QualityIssue(
            issue_id="docx_table_borders",
            status="unsupported",
            severity="high",
            check_key="docx.table.borders",
            title="DOCX table border style is custom and cannot be auto-verified.",
            profile_rule_ref="table.border_style",
            recommendation="Use a supported border style or review the table manually.",
            fixable=False,
        )
    for index, table in enumerate(document.tables):
        borders = table._tbl.tblPr.xpath("./w:tblBorders")
        if not borders:
            return QualityIssue(
                issue_id="docx_table_borders",
                status="warning",
                check_key="docx.table.borders",
                title="DOCX table border rules need review.",
                description=f"table[{index}] does not define tblBorders.",
                profile_rule_ref="table",
                location=f"table[{index}]",
                recommendation=f"Apply the {profile.table.border_style} table border style.",
                fixable=True,
            )
        border = borders[0]
        expected_edges = {
            "three_line": {"top", "bottom", "insideH"},
            "minimal": {"top", "bottom"},
            "full_grid": {"top", "bottom", "insideH", "insideV", "left", "right"},
        }[profile.table.border_style]
        missing = [edge for edge in expected_edges if not _table_border_enabled(border, edge)]
        extra = [edge for edge in {"top", "bottom", "insideH", "insideV", "left", "right"} - expected_edges if _table_border_enabled(border, edge)]
        if missing or extra:
            return QualityIssue(
                issue_id="docx_table_borders",
                status="warning",
                check_key="docx.table.borders",
                title="DOCX table border rules need review.",
                description=f"table[{index}] missing {', '.join(missing) if missing else 'none'}; unexpected {', '.join(extra) if extra else 'none'}.",
                profile_rule_ref="table",
                location=f"table[{index}]",
                recommendation=f"Apply the {profile.table.border_style} table border style.",
                fixable=True,
            )
    return QualityIssue(
        issue_id="docx_table_borders",
        status="pass",
        severity="info",
        check_key="docx.table.borders",
        title="DOCX table border rules are present.",
        profile_rule_ref="table",
    )


def _inspect_table_rule_fields(document: Document, profile: FormatProfile) -> list[QualityIssue]:
    if not document.tables:
        return [
            QualityIssue(
                issue_id=_quality_issue_id_from_check_key(check_key),
                status="pass",
                severity="info",
                check_key=check_key,
                title=f"{title} is not applicable because no tables were detected.",
                profile_rule_ref=field_path,
                details={"not_applicable": True},
            )
            for field_path, check_key, title in TABLE_RULE_FIELD_CHECKS
        ]
    border_mismatches = _table_border_mismatches(document, profile)
    header_repeat_mismatches = _table_header_repeat_mismatches(document, profile)
    return [
        _field_issue_from_mismatches(
            "table.border_style",
            "docx.table.border_style",
            "DOCX table border style",
            border_mismatches,
            location="tables",
            fixable=True,
        ),
        _field_issue_from_mismatches(
            "table.header_repeat",
            "docx.table.header_repeat",
            "DOCX table header repeat",
            header_repeat_mismatches,
            location="tables",
            fixable=True,
        ),
    ]


def _table_border_mismatches(document: Document, profile: FormatProfile) -> list[str]:
    if profile.table.border_style == "custom":
        return ["custom table border style cannot be auto-verified"]
    mismatches: list[str] = []
    expected_edges = {
        "three_line": {"top", "bottom", "insideH"},
        "minimal": {"top", "bottom"},
        "full_grid": {"top", "bottom", "insideH", "insideV", "left", "right"},
    }[profile.table.border_style]
    for index, table in enumerate(document.tables):
        borders = table._tbl.tblPr.xpath("./w:tblBorders")
        if not borders:
            mismatches.append(f"table[{index}] does not define tblBorders")
            continue
        border = borders[0]
        missing = [edge for edge in expected_edges if not _table_border_enabled(border, edge)]
        extra = [edge for edge in {"top", "bottom", "insideH", "insideV", "left", "right"} - expected_edges if _table_border_enabled(border, edge)]
        if missing or extra:
            mismatches.append(
                f"table[{index}] missing {', '.join(missing) if missing else 'none'}; unexpected {', '.join(extra) if extra else 'none'}"
            )
    return mismatches


def _table_header_repeat_mismatches(document: Document, profile: FormatProfile) -> list[str]:
    mismatches: list[str] = []
    for index, table in enumerate(document.tables):
        if not table.rows:
            continue
        repeats = _table_row_has_header_repeat(table.rows[0])
        if profile.table.header_repeat and not repeats:
            mismatches.append(f"table[{index}] first row is not marked as repeating header")
        if not profile.table.header_repeat and repeats:
            mismatches.append(f"table[{index}] first row unexpectedly repeats as header")
    return mismatches


def _table_row_has_header_repeat(row) -> bool:
    header = row._tr.get_or_add_trPr().find(qn("w:tblHeader"))
    if header is None:
        return False
    return header.get(qn("w:val")) not in {"0", "false"}


def _inspect_captions(document: Document, structure: DocumentStructure, profile: FormatProfile) -> QualityIssue:
    captions = [
        (paragraph, classification.role)
        for paragraph, classification in _paragraphs_with_classifications(document, structure)
        if classification.role in {ParagraphRole.TABLE_CAPTION, ParagraphRole.FIGURE_CAPTION}
    ]
    caption_required_image_count = _caption_required_inline_image_count(document, profile)
    if not captions:
        if not document.tables and caption_required_image_count == 0:
            return QualityIssue(
                issue_id="docx_captions",
                status="pass",
                severity="info",
                check_key="docx.captions",
                title="DOCX caption check is not applicable because no tables or images were detected.",
                profile_rule_ref="table.caption; figure.caption",
            )
        missing_targets: list[str] = []
        if document.tables:
            missing_targets.append(f"{len(document.tables)} table(s)")
        if caption_required_image_count:
            missing_targets.append(f"{caption_required_image_count} image(s)")
        return QualityIssue(
            issue_id="docx_captions",
            status="warning",
            check_key="docx.captions",
            title="No figure or table captions were detected for existing visual objects.",
            description=f"Detected {', '.join(missing_targets)} but no supported captions.",
            recommendation="Add table/figure captions or review whether captions are intentionally omitted.",
            fixable=False,
        )
    mismatches: list[str] = []
    for paragraph, role in captions:
        location = _paragraph_location(document, paragraph)
        expected_font = profile.table.caption.font if role == ParagraphRole.TABLE_CAPTION else profile.figure.caption.font
        expected_prefixes = _expected_caption_prefixes(profile, role)
        text = paragraph.text.strip()
        if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            mismatches.append(f"{location} caption is not centered")
        run = next((r for r in paragraph.runs if r.text.strip()), None)
        if profile.quality.check_fonts and run is not None and _run_style_mismatches(run, expected_font, f"{location} caption"):
            mismatches.append(f"{location} caption font does not match profile")
        if text and not any(_caption_text_matches(text, prefix, profile.table.caption.separator if role == ParagraphRole.TABLE_CAPTION else profile.figure.caption.separator) for prefix in expected_prefixes):
            mismatches.append(f"{location} caption text does not match configured prefix or separator")
    if mismatches:
        return QualityIssue(
            issue_id="docx_captions",
            status="warning",
            check_key="docx.captions",
            title="DOCX caption alignment needs review.",
            description="; ".join(mismatches[:12]),
            profile_rule_ref="table.caption; figure.caption",
            location=", ".join(mismatches[:8]),
            recommendation="Center figure and table captions and keep the configured prefix/separator.",
            fixable=True,
        )
    return QualityIssue(
        issue_id="docx_captions",
        status="pass",
        severity="info",
        check_key="docx.captions",
        title="DOCX captions match supported profile checks.",
        profile_rule_ref="table.caption; figure.caption",
    )


def _inspect_equations(document: Document, structure: DocumentStructure, profile: FormatProfile) -> QualityIssue:
    equations = [
        paragraph
        for paragraph, classification in _paragraphs_with_classifications(document, structure)
        if classification.role == ParagraphRole.EQUATION
    ]
    if not equations:
        return QualityIssue(
            issue_id="docx_equations",
            status="pass",
            severity="info",
            check_key="docx.equations",
            title="DOCX equation check is not applicable because no equations were detected.",
            profile_rule_ref="equations",
        )
    mismatches: list[str] = []
    for paragraph in equations:
        location = _paragraph_location(document, paragraph)
        if paragraph.alignment != _alignment_value(profile.equations.alignment):
            mismatches.append(f"{location} equation alignment expected {profile.equations.alignment}")
        if profile.equations.numbering != "none" and not _equation_has_visible_number(paragraph.text.strip()):
            mismatches.append(f"{location} equation is missing visible {profile.equations.numbering}-side numbering")
        if profile.quality.check_fonts:
            run = next((r for r in paragraph.runs if r.text.strip()), None)
            if run is not None:
                expected_font = TextFont(
                    chinese=profile.equations.font,
                    latin=profile.equations.font,
                    size_pt=profile.body.font.size_pt,
                    weight="normal",
                    color=profile.body.font.color,
                )
                mismatches.extend(_run_style_mismatches(run, expected_font, f"{location} equation"))
    if mismatches:
        return QualityIssue(
            issue_id="docx_equations",
            status="warning",
            severity="high",
            check_key="docx.equations",
            title="DOCX equation formatting needs review.",
            description="; ".join(mismatches[:8]),
            profile_rule_ref="equations",
            recommendation="Apply equation alignment, font, and numbering rules before final delivery.",
            fixable=True,
        )
    return QualityIssue(
        issue_id="docx_equations",
        status="pass",
        severity="info",
        check_key="docx.equations",
        title="DOCX equation formatting matches supported profile checks.",
        profile_rule_ref="equations",
    )


def _inspect_role_style_consistency(
    document: Document,
    structure: DocumentStructure,
    profile: FormatProfile,
) -> QualityIssue:
    mismatches: list[str] = []
    for paragraph, classification in _paragraphs_with_classifications(document, structure):
        role = classification.role
        if role == ParagraphRole.BODY:
            fmt = paragraph.paragraph_format
            first_line = fmt.first_line_indent.cm if fmt.first_line_indent else 0
            left = fmt.left_indent.cm if fmt.left_indent else 0
            if first_line < -0.05 and left > 0.05:
                mismatches.append(f"{_paragraph_location(document, paragraph)} body has reference-style hanging indent")
            if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                mismatches.append(f"{_paragraph_location(document, paragraph)} body is centered like a caption or equation")
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                mismatches.append(f"{_paragraph_location(document, paragraph)} body is using a Word heading style")
        elif role == ParagraphRole.TOC_ITEM:
            run = next((r for r in paragraph.runs if r.text.strip()), None)
            if run:
                mismatches.extend(_run_style_mismatches(run, profile.body.font, f"{_paragraph_location(document, paragraph)} toc item"))
        elif role in {
            ParagraphRole.DOCUMENT_TITLE,
            ParagraphRole.HEADING,
            ParagraphRole.REFERENCE_HEADING,
            ParagraphRole.ACKNOWLEDGEMENT_HEADING,
        }:
            if not profile.quality.check_headings:
                continue
            if not (paragraph.style and paragraph.style.name.startswith("Heading")):
                mismatches.append(f"{_paragraph_location(document, paragraph)} heading is not using a Word heading style")
        elif role in {ParagraphRole.TABLE_CAPTION, ParagraphRole.FIGURE_CAPTION}:
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                mismatches.append(f"{_paragraph_location(document, paragraph)} caption is not centered")
        elif role == ParagraphRole.EQUATION:
            if paragraph.alignment != _alignment_value(profile.equations.alignment):
                mismatches.append(f"{_paragraph_location(document, paragraph)} equation alignment expected {profile.equations.alignment}")
        elif role == ParagraphRole.REFERENCE_ITEM:
            if not profile.quality.check_references:
                continue
            fmt = paragraph.paragraph_format
            first_line = fmt.first_line_indent.cm if fmt.first_line_indent else 0
            left = fmt.left_indent.cm if fmt.left_indent else 0
            if not (first_line < -0.05 and left > 0.05):
                mismatches.append(f"{_paragraph_location(document, paragraph)} reference item is missing hanging indent")

    if mismatches:
        return QualityIssue(
            issue_id="docx_role_styles",
            status="warning",
            check_key="docx.role_styles",
            title="DOCX paragraph roles and applied styles need review.",
            description="; ".join(mismatches[:12]),
            profile_rule_ref="body; headings; captions; equations; references",
            recommendation="Re-run profile formatting or review paragraphs whose detected document role does not match their applied Word style.",
            fixable=False,
            details={"mismatch_count": len(mismatches)},
        )
    return QualityIssue(
        issue_id="docx_role_styles",
        status="pass",
        severity="info",
        check_key="docx.role_styles",
        title="DOCX paragraph roles match supported applied style checks.",
        profile_rule_ref="body; headings; captions; equations; references",
    )


def _inspect_heading_numbering(document: Document, structure: DocumentStructure, profile: FormatProfile) -> QualityIssue:
    if not profile.numbering.enabled and not profile.numbering.heading_pattern:
        return QualityIssue(
            issue_id="docx_heading_numbering",
            status="pass",
            severity="info",
            check_key="docx.heading_numbering",
            title="DOCX heading numbering is disabled by profile.",
        )
    mismatches: list[str] = []
    for paragraph, classification in _paragraphs_with_classifications(document, structure):
        if classification.role != ParagraphRole.HEADING:
            continue
        level = classification.heading_level or 1
        heading_rule = _heading_rule(profile, level)
        if heading_rule.numbering == "none" and not profile.numbering.heading_pattern:
            continue
        if not _looks_like_numbered_heading(paragraph.text.strip()):
            mismatches.append(
                f"{_paragraph_location(document, paragraph)} heading missing visible numbering"
            )
    if mismatches:
        return QualityIssue(
            issue_id="docx_heading_numbering",
            status="warning",
            severity="high",
            check_key="docx.heading_numbering",
            title="DOCX heading numbering needs review.",
            description="; ".join(mismatches[:8]),
            profile_rule_ref="numbering.heading_pattern; headings.numbering",
            recommendation="Apply the profile heading numbering pattern before final delivery.",
            fixable=True,
        )
    return QualityIssue(
        issue_id="docx_heading_numbering",
        status="pass",
        severity="info",
        check_key="docx.heading_numbering",
        title="DOCX heading numbering matches the profile.",
        profile_rule_ref="numbering.heading_pattern; headings.numbering",
    )


def _heading_rule(profile: FormatProfile, level: int):
    exact = next((heading for heading in profile.headings if heading.level == level), None)
    if exact is not None:
        return exact
    lower_or_equal = [heading for heading in profile.headings if heading.level <= level]
    if lower_or_equal:
        return sorted(lower_or_equal, key=lambda heading: heading.level)[-1]
    return profile.headings[0]


def _looks_like_numbered_heading(text: str) -> bool:
    return bool(
        re.match(r"^(第[一二三四五六七八九十百0-9]+[章节])", text)
        or re.match(r"^[0-9]+(?:\.[0-9]+)*(?:[、.）\)])\s*", text)
        or re.match(r"^[0-9]+(?:\.[0-9]+)*\s+", text)
    )


def _inspect_raw_latex(document: Document) -> QualityIssue:
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text
        if _contains_raw_latex(text):
            return QualityIssue(
                issue_id="docx_raw_latex",
                status="fail",
                severity="high",
                check_key="docx.raw_latex",
                title="DOCX contains raw LaTeX residue.",
                description=text,
                location=f"paragraph[{index}]",
                recommendation="Convert formulas to Word-readable equation formatting before final delivery.",
                fixable=False,
            )
    return QualityIssue(
        issue_id="docx_raw_latex",
        status="pass",
        severity="info",
        check_key="docx.raw_latex",
        title="No raw LaTeX residue detected in DOCX paragraphs.",
    )


def _inspect_basic_page_numbers(
    document: Document,
    profile: FormatProfile,
    features: OoxmlDocumentFeatures | None,
) -> QualityIssue:
    if not profile.header_footer.footer_page_number:
        return QualityIssue(
            issue_id="docx_page_number",
            status="pass",
            severity="info",
            check_key="docx.page_number",
            title="DOCX page number check is disabled by the selected profile.",
            profile_rule_ref="header_footer.footer_page_number",
        )
    missing_sections: list[str] = []
    mismatches: list[str] = []
    expected_fmt = _page_number_ooxml_format(profile.header_footer.page_number_format)
    for index, section in enumerate(document.sections):
        has_page_field = _section_has_page_field(section)
        if profile.header_footer.page_number_format == "none":
            if has_page_field:
                mismatches.append(f"section[{index}] has an unexpected PAGE field")
        elif not has_page_field:
            missing_sections.append(f"section[{index}]")
        if features is not None:
            actual_fmt = features.page_number_formats[index] if index < len(features.page_number_formats) else None
            actual_start = features.page_number_starts[index] if index < len(features.page_number_starts) else None
            if profile.header_footer.page_number_format == "none":
                if actual_fmt is not None:
                    mismatches.append(f"section[{index}] has unexpected page-number format {actual_fmt}")
            else:
                if actual_fmt != expected_fmt:
                    mismatches.append(
                        f"section[{index}] page-number format expected {expected_fmt}, found {actual_fmt}"
                    )
            if actual_start is not None and actual_start != profile.header_footer.page_number_start:
                mismatches.append(
                    f"section[{index}] page-number start expected {profile.header_footer.page_number_start}, found {actual_start}"
                )
    if missing_sections:
        return QualityIssue(
            issue_id="docx_page_number",
            status="fail",
            severity="high",
            check_key="docx.page_number",
            title="DOCX page number field was not found in every section footer.",
            description=", ".join(missing_sections),
            profile_rule_ref="page.numbering",
            recommendation="Apply supported centered footer PAGE fields before final delivery.",
            fixable=True,
        )
    if mismatches:
        return QualityIssue(
            issue_id="docx_page_number",
            status="warning",
            severity="high",
            check_key="docx.page_number",
            title="DOCX page number format or start value needs review.",
            description="; ".join(mismatches),
            profile_rule_ref="header_footer.page_number_format; header_footer.page_number_start",
            recommendation="Apply the profile page-number format and start value.",
            fixable=True,
        )
    return QualityIssue(
        issue_id="docx_page_number",
        status="pass",
        severity="info",
        check_key="docx.page_number",
        title="DOCX contains supported footer PAGE fields.",
        profile_rule_ref="page.numbering",
    )


def _inspect_page_number_fields(
    document: Document,
    profile: FormatProfile,
    features: OoxmlDocumentFeatures | None,
) -> list[QualityIssue]:
    mismatches_by_field = {field_path: [] for field_path, _, _ in PAGE_NUMBER_FIELD_CHECKS}
    expected_fmt = _page_number_ooxml_format(profile.header_footer.page_number_format)
    for index, section in enumerate(document.sections):
        has_page_field = _section_has_page_field(section)
        if profile.header_footer.footer_page_number and not has_page_field:
            mismatches_by_field["header_footer.footer_page_number"].append(f"section[{index}] missing footer page number")
        if not profile.header_footer.footer_page_number and has_page_field:
            mismatches_by_field["header_footer.footer_page_number"].append(f"section[{index}] has unexpected footer page number")
        if features is None:
            continue
        actual_fmt = features.page_number_formats[index] if index < len(features.page_number_formats) else None
        actual_start = features.page_number_starts[index] if index < len(features.page_number_starts) else None
        if profile.header_footer.page_number_format == "none":
            if actual_fmt is not None:
                mismatches_by_field["header_footer.page_number_format"].append(
                    f"section[{index}] has unexpected page-number format {actual_fmt}"
                )
        elif actual_fmt != expected_fmt:
            mismatches_by_field["header_footer.page_number_format"].append(
                f"section[{index}] page-number format expected {expected_fmt}, found {actual_fmt}"
            )
        if actual_start is not None and actual_start != profile.header_footer.page_number_start:
            mismatches_by_field["header_footer.page_number_start"].append(
                f"section[{index}] page-number start expected {profile.header_footer.page_number_start}, found {actual_start}"
            )

    issues: list[QualityIssue] = []
    for field_path, check_key, title in PAGE_NUMBER_FIELD_CHECKS:
        field_mismatches = mismatches_by_field[field_path]
        if field_mismatches:
            issues.append(
                QualityIssue(
                    issue_id=_quality_issue_id_from_check_key(check_key),
                    status="warning",
                    severity="high",
                    check_key=check_key,
                    title=f"{title} needs review.",
                    description="; ".join(field_mismatches),
                    profile_rule_ref=field_path,
                    location="sections",
                    recommendation=f"Apply profile rule `{field_path}` to every section.",
                    fixable=True,
                    details={"mismatch_count": len(field_mismatches)},
                )
            )
            continue
        issues.append(
            QualityIssue(
                issue_id=_quality_issue_id_from_check_key(check_key),
                status="pass",
                severity="info",
                check_key=check_key,
                title=f"{title} matches the profile.",
                profile_rule_ref=field_path,
                location=f"{len(document.sections)} section(s)",
            )
        )
    return issues


def _disabled_page_number_field_issues(reason: str) -> list[QualityIssue]:
    return [
        QualityIssue(
            issue_id=_quality_issue_id_from_check_key(check_key),
            status="pass",
            severity="info",
            check_key=check_key,
            title=f"{title} inspection is delegated.",
            profile_rule_ref=field_path,
            details={"disabled_by_profile": True, "reason": reason},
        )
        for field_path, check_key, title in PAGE_NUMBER_FIELD_CHECKS
    ]


def _inspect_ooxml_feature_inventory(features: OoxmlDocumentFeatures) -> QualityIssue:
    return QualityIssue(
        issue_id="docx_ooxml_features",
        status="pass",
        severity="info",
        check_key="docx.ooxml.features",
        title="DOCX internal OOXML feature inventory completed.",
        details={
            "section_count": features.section_count,
            "toc_field_count": features.toc_field_count,
            "footnote_count": features.footnote_count,
            "endnote_count": features.endnote_count,
            "inline_image_count": features.inline_image_count,
            "anchored_image_count": features.anchored_image_count,
            "numbering_reference_count": features.numbering_reference_count,
            "omml_equation_count": features.omml_equation_count,
        },
    )


def _inspect_field_update_policy(features: OoxmlDocumentFeatures | None) -> QualityIssue:
    if features is None:
        return QualityIssue(
            issue_id="docx_field_update_policy",
            status="unsupported",
            check_key="docx.fields.update_policy",
            title="DOCX field update policy cannot be inspected.",
            recommendation="Regenerate the DOCX and rerun quality inspection.",
        )
    fields_requiring_refresh = features.toc_field_count + features.page_field_count
    if fields_requiring_refresh > 0 and not features.has_update_fields:
        return QualityIssue(
            issue_id="docx_field_update_policy",
            status="warning",
            check_key="docx.fields.update_policy",
            title="DOCX fields exist but automatic field refresh is not enabled.",
            description=f"Detected {fields_requiring_refresh} supported field reference(s).",
            profile_rule_ref="fields.update_on_open",
            recommendation="Enable Word updateFields so TOC and PAGE fields can refresh during final export/open.",
            fixable=True,
            details={"fields_requiring_refresh": fields_requiring_refresh},
        )
    return QualityIssue(
        issue_id="docx_field_update_policy",
        status="pass",
        severity="info",
        check_key="docx.fields.update_policy",
        title="DOCX field update policy is safe for supported fields.",
        details={"has_update_fields": features.has_update_fields, "fields_requiring_refresh": fields_requiring_refresh},
    )


def _inspect_document_grid(
    document: Document,
    profile: FormatProfile,
    features: OoxmlDocumentFeatures | None,
) -> QualityIssue:
    if features is None:
        return QualityIssue(
            issue_id="docx_document_grid",
            status="unsupported",
            check_key="docx.document_grid",
            title="DOCX document-grid settings cannot be inspected.",
            recommendation="Regenerate the DOCX and rerun quality inspection.",
        )
    expected_enabled = profile.document_grid.enabled and profile.document_grid.type != "none"
    expected_type = (
        "default"
        if not expected_enabled
        else ("lines" if profile.document_grid.type == "line" else "linesAndChars")
    )
    mismatches: list[str] = []
    for index, section in enumerate(document.sections):
        actual_type = features.document_grid_types[index] if index < len(features.document_grid_types) else None
        if actual_type != expected_type:
            mismatches.append(f"section[{index}] docGrid type expected {expected_type}, found {actual_type}")
        actual_line_pitch = features.document_grid_line_pitches[index] if index < len(features.document_grid_line_pitches) else None
        actual_char_space = features.document_grid_char_spaces[index] if index < len(features.document_grid_char_spaces) else None
        if expected_enabled and profile.document_grid.lines_per_page and actual_line_pitch is None:
            mismatches.append(f"section[{index}] missing docGrid linePitch")
        if expected_enabled and profile.document_grid.characters_per_line and actual_char_space is None:
            mismatches.append(f"section[{index}] missing docGrid charSpace")
    if mismatches:
        return QualityIssue(
            issue_id="docx_document_grid",
            status="warning",
            severity="high",
            check_key="docx.document_grid",
            title="DOCX document-grid rules need review.",
            description="; ".join(mismatches),
            profile_rule_ref="document_grid",
            recommendation="Apply the profile document-grid settings.",
            fixable=True,
        )
    return QualityIssue(
        issue_id="docx_document_grid",
        status="pass",
        severity="info",
        check_key="docx.document_grid",
        title="DOCX document-grid settings match the profile.",
        profile_rule_ref="document_grid",
    )


def _inspect_toc_fields(
    document: Document,
    structure: DocumentStructure,
    features: OoxmlDocumentFeatures | None,
    profile: FormatProfile,
) -> QualityIssue:
    if features is None:
        return QualityIssue(
            issue_id="docx_toc_fields",
            status="unsupported",
            check_key="docx.toc.fields",
            title="DOCX table-of-contents fields cannot be inspected.",
            recommendation="Regenerate the DOCX and rerun quality inspection.",
        )

    toc_classified_paragraphs = [
        (paragraph, classification)
        for paragraph, classification in _paragraphs_with_classifications(document, structure)
        if classification.role in {ParagraphRole.TOC_TITLE, ParagraphRole.TOC_ITEM}
    ]
    toc_text_paragraphs = [_paragraph_location(document, paragraph) for paragraph, _classification in toc_classified_paragraphs]
    toc_title_texts = [
        paragraph.text.strip()
        for paragraph, classification in toc_classified_paragraphs
        if classification.role == ParagraphRole.TOC_TITLE and paragraph.text.strip()
    ]
    visible_paragraph_texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    if not toc_text_paragraphs and features.toc_field_count == 0:
        return QualityIssue(
            issue_id="docx_toc_fields",
            status="pass",
            severity="info",
            check_key="docx.toc.fields",
            title="DOCX TOC check is not applicable because no TOC was detected.",
        )
    mismatches: list[str] = []
    if features.toc_field_count > 0 and not features.has_update_fields:
        mismatches.append("TOC field codes were detected without updateFields enabled")
    if features.toc_field_count == 0:
        if toc_text_paragraphs:
            mismatches.append("Manual TOC-like text was detected without a Word TOC field")
        elif profile.toc.enabled:
            mismatches.append("TOC is enabled in the profile but no TOC field or TOC text was detected")
    elif profile.toc.enabled:
        expected_title = profile.toc.title.strip()
        if expected_title and expected_title not in visible_paragraph_texts:
            found_titles = ", ".join(toc_title_texts[:3]) if toc_title_texts else "none"
            mismatches.append(f"TOC title expected {expected_title}, found {found_titles}")
    if features.toc_field_count > 0 and profile.toc.include_levels > 0:
        expected_switch = f'\\o "1-{profile.toc.include_levels}"'
        if not any(expected_switch in instr for instr in features.toc_instructions):
            mismatches.append(f"TOC field missing expected heading-depth switch {expected_switch}")
        if profile.toc.use_hyperlinks and not any(r"\h" in instr for instr in features.toc_instructions):
            mismatches.append("TOC field is missing hyperlink support")
        if not profile.toc.show_page_numbers and not any(r"\n" in instr for instr in features.toc_instructions):
            mismatches.append("TOC field is missing no-page-number switch")
        if profile.toc.show_page_numbers and not profile.toc.right_align_page_numbers and not any(r'\p " "' in instr for instr in features.toc_instructions):
            mismatches.append('TOC field is missing left-flow page-number separator switch \\p " "')
        if profile.toc.show_page_numbers and profile.toc.right_align_page_numbers and any(r'\p " "' in instr for instr in features.toc_instructions):
            mismatches.append('TOC field has non-right-aligned page-number separator switch \\p " "')
    if mismatches:
        return QualityIssue(
            issue_id="docx_toc_fields",
            status="warning",
            check_key="docx.toc.fields",
            title="Manual TOC-like text or TOC field options need review.",
            description="; ".join(mismatches[:8]) if mismatches else ", ".join(toc_text_paragraphs[:8]),
            profile_rule_ref="toc",
            recommendation="Review the table of contents and refresh the TOC field before delivery.",
            fixable=True,
            details={"toc_text_paragraph_count": len(toc_text_paragraphs), "toc_field_count": features.toc_field_count},
        )
    return QualityIssue(
        issue_id="docx_toc_fields",
        status="pass",
        severity="info",
        check_key="docx.toc.fields",
        title="DOCX TOC settings match supported profile checks.",
        profile_rule_ref="toc",
        details={
            "toc_field_count": features.toc_field_count,
            "simple_toc_field_count": features.simple_toc_field_count,
            "complex_toc_field_count": features.complex_toc_field_count,
        },
    )


def _inspect_toc_field_issues(
    document: Document,
    structure: DocumentStructure,
    features: OoxmlDocumentFeatures | None,
    profile: FormatProfile,
) -> list[QualityIssue]:
    if features is None:
        return [
            QualityIssue(
                issue_id=_quality_issue_id_from_check_key(check_key),
                status="unsupported",
                check_key=check_key,
                title=f"{title} cannot be inspected.",
                profile_rule_ref=field_path,
                recommendation="Regenerate the DOCX and rerun quality inspection.",
            )
            for field_path, check_key, title in TOC_FIELD_CHECKS
        ]

    toc_classified_paragraphs = [
        (paragraph, classification)
        for paragraph, classification in _paragraphs_with_classifications(document, structure)
        if classification.role in {ParagraphRole.TOC_TITLE, ParagraphRole.TOC_ITEM}
    ]
    toc_text_paragraphs = [_paragraph_location(document, paragraph) for paragraph, _classification in toc_classified_paragraphs]
    toc_title_texts = [
        paragraph.text.strip()
        for paragraph, classification in toc_classified_paragraphs
        if classification.role == ParagraphRole.TOC_TITLE and paragraph.text.strip()
    ]
    visible_paragraph_texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    toc_detected = bool(toc_text_paragraphs or features.toc_field_count > 0)
    toc_required = _profile_requires_toc(profile)
    mismatches_by_field = {field_path: [] for field_path, _, _ in TOC_FIELD_CHECKS}
    if profile.toc.enabled and toc_required and not toc_detected:
        mismatches_by_field["toc.enabled"].append("TOC is enabled in the profile but no TOC field or TOC text was detected")
    if not profile.toc.enabled and toc_detected:
        mismatches_by_field["toc.enabled"].append("TOC is disabled in the profile but TOC field or TOC text was detected")
    if not toc_detected:
        return [
            _field_issue_from_mismatches(
                field_path,
                check_key,
                title,
                mismatches_by_field[field_path],
                location="toc",
                fixable=True,
            )
            for field_path, check_key, title in TOC_FIELD_CHECKS
        ]

    if features.toc_field_count == 0 and toc_text_paragraphs:
        mismatches_by_field["toc.enabled"].append("Manual TOC-like text was detected without a Word TOC field")
    if profile.toc.enabled:
        expected_title = profile.toc.title.strip()
        if expected_title and expected_title not in visible_paragraph_texts:
            found_titles = ", ".join(toc_title_texts[:3]) if toc_title_texts else "none"
            mismatches_by_field["toc.title"].append(f"TOC title expected {expected_title}, found {found_titles}")
    if features.toc_field_count > 0 and profile.toc.include_levels > 0:
        expected_switch = f'\\o "1-{profile.toc.include_levels}"'
        if not any(expected_switch in instr for instr in features.toc_instructions):
            mismatches_by_field["toc.include_levels"].append(f"TOC field missing expected heading-depth switch {expected_switch}")
        has_no_page_number_switch = any(r"\n" in instr for instr in features.toc_instructions)
        if not profile.toc.show_page_numbers and not has_no_page_number_switch:
            mismatches_by_field["toc.show_page_numbers"].append("TOC field is missing no-page-number switch")
        if profile.toc.show_page_numbers and has_no_page_number_switch:
            mismatches_by_field["toc.show_page_numbers"].append("TOC field unexpectedly hides page numbers")
        has_left_flow_separator = any(r'\p " "' in instr for instr in features.toc_instructions)
        if profile.toc.show_page_numbers and not profile.toc.right_align_page_numbers and not has_left_flow_separator:
            mismatches_by_field["toc.right_align_page_numbers"].append('TOC field is missing left-flow page-number separator switch \\p " "')
        if profile.toc.show_page_numbers and profile.toc.right_align_page_numbers and has_left_flow_separator:
            mismatches_by_field["toc.right_align_page_numbers"].append('TOC field has non-right-aligned page-number separator switch \\p " "')
        has_hyperlink_switch = any(r"\h" in instr for instr in features.toc_instructions)
        if profile.toc.use_hyperlinks and not has_hyperlink_switch:
            mismatches_by_field["toc.use_hyperlinks"].append("TOC field is missing hyperlink support")
        if not profile.toc.use_hyperlinks and has_hyperlink_switch:
            mismatches_by_field["toc.use_hyperlinks"].append("TOC field unexpectedly includes hyperlink support")
        if profile.toc.update_fields_on_open and not features.has_update_fields:
            mismatches_by_field["toc.update_fields_on_open"].append("TOC field codes were detected without updateFields enabled")

    return [
        _field_issue_from_mismatches(
            field_path,
            check_key,
            title,
            mismatches_by_field[field_path],
            location="toc",
            fixable=True,
        )
        for field_path, check_key, title in TOC_FIELD_CHECKS
    ]


def _profile_requires_toc(profile: FormatProfile) -> bool:
    if not profile.toc.enabled:
        return False
    if profile.schema_version != "1.0.0":
        return True
    return any(section.key.lower() == "toc" and section.required for section in profile.sections)


def _inspect_section_complexity(features: OoxmlDocumentFeatures | None) -> QualityIssue:
    if features is None:
        return QualityIssue(
            issue_id="docx_sections",
            status="unsupported",
            check_key="docx.sections",
            title="DOCX sections cannot be inspected.",
            recommendation="Regenerate the DOCX and rerun quality inspection.",
        )
    if features.section_count <= 1:
        return QualityIssue(
            issue_id="docx_sections",
            status="pass",
            severity="info",
            check_key="docx.sections",
            title="DOCX has a single supported section.",
            details={"section_count": features.section_count},
        )
    return QualityIssue(
        issue_id="docx_sections",
        status="pass",
        severity="info",
        check_key="docx.sections",
        title="DOCX has multiple sections and supported section-level checks ran across all sections.",
        recommendation="If the profile expects mixed portrait/landscape sections or section-specific page numbering, review those requirements manually.",
        details={"section_count": features.section_count},
    )


def _inspect_template_placeholders(document: Document) -> QualityIssue:
    placeholder_pattern = re.compile(r"\{\{[^{}]+\}\}")
    matches: list[str] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        matches.extend(f"paragraph[{index}] {match.group(0)}" for match in placeholder_pattern.finditer(paragraph.text))
    for section_index, section in enumerate(document.sections):
        matches.extend(_container_placeholder_matches(section.header, placeholder_pattern, f"section[{section_index}].header"))
        matches.extend(_container_placeholder_matches(section.footer, placeholder_pattern, f"section[{section_index}].footer"))
        matches.extend(_container_placeholder_matches(section.first_page_header, placeholder_pattern, f"section[{section_index}].first_page_header"))
        matches.extend(_container_placeholder_matches(section.first_page_footer, placeholder_pattern, f"section[{section_index}].first_page_footer"))
        matches.extend(_container_placeholder_matches(section.even_page_header, placeholder_pattern, f"section[{section_index}].even_page_header"))
        matches.extend(_container_placeholder_matches(section.even_page_footer, placeholder_pattern, f"section[{section_index}].even_page_footer"))
    if matches:
        return QualityIssue(
            issue_id="docx_template_placeholders",
            status="warning",
            severity="high",
            check_key="docx.template.placeholders",
            title="DOCX contains unresolved template placeholders.",
            description=f"Unresolved template placeholder(s): {'; '.join(matches[:12])}",
            profile_rule_ref="template_binding.placeholder_policy",
            recommendation="Resolve, remove, or fill all template placeholders before final delivery.",
            fixable=False,
            details={"placeholder_count": len(matches), "placeholders": matches[:50]},
        )
    return QualityIssue(
        issue_id="docx_template_placeholders",
        status="pass",
        severity="info",
        check_key="docx.template.placeholders",
        title="DOCX has no unresolved template placeholders.",
        profile_rule_ref="template_binding.placeholder_policy",
    )


def _inspect_template_body_slot(document: Document, profile: FormatProfile) -> QualityIssue:
    body_slot = profile.template_binding.body_slot
    if not body_slot:
        return QualityIssue(
            issue_id="docx_template_body_slot",
            status="pass",
            severity="info",
            check_key="docx.template.body_slot",
            title="DOCX template body-slot inspection is not applicable.",
            profile_rule_ref="template_binding.body_slot",
            details={"not_applicable": True},
        )
    matches: list[str] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        if body_slot in paragraph.text:
            matches.append(f"paragraph[{index}] {body_slot}")
    for section_index, section in enumerate(document.sections):
        matches.extend(_container_literal_matches(section.header, body_slot, f"section[{section_index}].header"))
        matches.extend(_container_literal_matches(section.footer, body_slot, f"section[{section_index}].footer"))
        matches.extend(_container_literal_matches(section.first_page_header, body_slot, f"section[{section_index}].first_page_header"))
        matches.extend(_container_literal_matches(section.first_page_footer, body_slot, f"section[{section_index}].first_page_footer"))
        matches.extend(_container_literal_matches(section.even_page_header, body_slot, f"section[{section_index}].even_page_header"))
        matches.extend(_container_literal_matches(section.even_page_footer, body_slot, f"section[{section_index}].even_page_footer"))
    if matches:
        return QualityIssue(
            issue_id="docx_template_body_slot",
            status="warning",
            severity="high",
            check_key="docx.template.body_slot",
            title="DOCX still contains the configured template body slot marker.",
            description=f"Unconsumed template body slot marker(s): {'; '.join(matches[:12])}",
            profile_rule_ref="template_binding.body_slot",
            recommendation="Compose the document through TemplateLoader and consume the configured body slot before delivery.",
            fixable=False,
            details={"slot": body_slot, "slot_count": len(matches), "locations": matches[:50]},
        )
    return QualityIssue(
        issue_id="docx_template_body_slot",
        status="pass",
        severity="info",
        check_key="docx.template.body_slot",
        title="DOCX has no remaining configured template body slot marker.",
        profile_rule_ref="template_binding.body_slot",
        details={"slot": body_slot},
    )


def _container_placeholder_matches(container, pattern: re.Pattern[str], label: str) -> list[str]:
    matches: list[str] = []
    for index, paragraph in enumerate(container.paragraphs, start=1):
        matches.extend(f"{label}.paragraph[{index}] {match.group(0)}" for match in pattern.finditer(paragraph.text))
    return matches


def _container_literal_matches(container, literal: str, label: str) -> list[str]:
    matches: list[str] = []
    for index, paragraph in enumerate(container.paragraphs, start=1):
        if literal in paragraph.text:
            matches.append(f"{label}.paragraph[{index}] {literal}")
    return matches


def _inspect_notes_support(path: Path, features: OoxmlDocumentFeatures | None, profile: FormatProfile) -> QualityIssue:
    if features is None:
        return QualityIssue(
            issue_id="docx_notes",
            status="unsupported",
            check_key="docx.notes",
            title="DOCX footnotes/endnotes cannot be inspected.",
            recommendation="Regenerate the DOCX and rerun quality inspection.",
        )
    count = features.footnote_count + features.endnote_count
    if count == 0:
        return QualityIssue(
            issue_id="docx_notes",
            status="pass",
            severity="info",
            check_key="docx.notes",
            title="DOCX footnote/endnote check is not applicable because none were detected.",
        )
    try:
        mismatches = _inspect_notes_xml_format(path, profile)
    except (BadZipFile, KeyError, ET.ParseError, OSError) as exc:
        return QualityIssue(
            issue_id="docx_notes",
            status="unsupported",
            check_key="docx.notes",
            title="DOCX footnote/endnote formatting cannot be inspected.",
            description=str(exc),
            profile_rule_ref="notes",
            recommendation="Regenerate the DOCX and rerun quality inspection.",
            fixable=False,
        )
    if mismatches:
        return QualityIssue(
            issue_id="docx_notes",
            status="fail",
            severity="high",
            check_key="docx.notes",
            title="DOCX footnote/endnote formatting does not match the profile.",
            description="; ".join(mismatches[:10]),
            profile_rule_ref="notes.font; notes.line_spacing",
            recommendation="Apply profile note formatting to footnotes/endnotes before delivery.",
            fixable=True,
            details={
                "footnote_count": features.footnote_count,
                "endnote_count": features.endnote_count,
                "mismatch_count": len(mismatches),
            },
        )
    return QualityIssue(
        issue_id="docx_notes",
        status="pass",
        severity="info",
        check_key="docx.notes",
        title="DOCX footnotes/endnotes match the profile.",
        description=f"Detected and formatted {features.footnote_count} footnote(s) and {features.endnote_count} endnote(s).",
        profile_rule_ref="notes",
        details={"footnote_count": features.footnote_count, "endnote_count": features.endnote_count},
    )


def _inspect_notes_xml_format(path: Path, profile: FormatProfile) -> list[str]:
    mismatches: list[str] = []
    expected_font = profile.notes.font
    expected_size = str(round(expected_font.size_pt * 2))
    expected_line = str(round(profile.notes.line_spacing * 240))
    expected_before = str(round(profile.notes.space_before_pt * 20))
    expected_after = str(round(profile.notes.space_after_pt * 20))
    with ZipFile(path) as package:
        for part_name, note_tag in (("word/footnotes.xml", "footnote"), ("word/endnotes.xml", "endnote")):
            if part_name not in package.namelist():
                continue
            root = ET.fromstring(package.read(part_name))
            for note in root.findall(f"w:{note_tag}", NS):
                note_id = note.get(qn("w:id"))
                if note_id in {"-1", "0"}:
                    continue
                for p_index, paragraph in enumerate(note.findall(".//w:p", NS), start=1):
                    spacing = paragraph.find("w:pPr/w:spacing", NS)
                    if spacing is None:
                        mismatches.append(f"{part_name} note[{note_id}] paragraph[{p_index}] missing spacing")
                    else:
                        _append_note_spacing_mismatches(
                            mismatches,
                            spacing,
                            part_name,
                            note_id or "?",
                            p_index,
                            expected_before,
                            expected_after,
                            expected_line,
                        )
                    runs = paragraph.findall(".//w:r", NS)
                    if not runs:
                        mismatches.append(f"{part_name} note[{note_id}] paragraph[{p_index}] has no runs")
                    for r_index, run in enumerate(runs, start=1):
                        _append_note_run_mismatches(
                            mismatches,
                            run,
                            part_name,
                            note_id or "?",
                            p_index,
                            r_index,
                            expected_font,
                            expected_size,
                        )
    return mismatches


def _append_note_spacing_mismatches(
    mismatches: list[str],
    spacing: ET.Element,
    part_name: str,
    note_id: str,
    paragraph_index: int,
    expected_before: str,
    expected_after: str,
    expected_line: str,
) -> None:
    prefix = f"{part_name} note[{note_id}] paragraph[{paragraph_index}]"
    if spacing.get(qn("w:before")) != expected_before:
        mismatches.append(f"{prefix} spacing before expected {expected_before}, found {spacing.get(qn('w:before'))}")
    if spacing.get(qn("w:after")) != expected_after:
        mismatches.append(f"{prefix} spacing after expected {expected_after}, found {spacing.get(qn('w:after'))}")
    if spacing.get(qn("w:line")) != expected_line:
        mismatches.append(f"{prefix} line spacing expected {expected_line}, found {spacing.get(qn('w:line'))}")


def _append_note_run_mismatches(
    mismatches: list[str],
    run: ET.Element,
    part_name: str,
    note_id: str,
    paragraph_index: int,
    run_index: int,
    expected_font: TextFont,
    expected_size: str,
) -> None:
    prefix = f"{part_name} note[{note_id}] paragraph[{paragraph_index}] run[{run_index}]"
    rpr = run.find("w:rPr", NS)
    if rpr is None:
        mismatches.append(f"{prefix} missing run properties")
        return
    r_fonts = rpr.find("w:rFonts", NS)
    if r_fonts is None:
        mismatches.append(f"{prefix} missing run fonts")
    else:
        if r_fonts.get(qn("w:eastAsia")) != expected_font.chinese:
            mismatches.append(f"{prefix} eastAsia font expected {expected_font.chinese}, found {r_fonts.get(qn('w:eastAsia'))}")
        if r_fonts.get(qn("w:ascii")) != expected_font.latin:
            mismatches.append(f"{prefix} ascii font expected {expected_font.latin}, found {r_fonts.get(qn('w:ascii'))}")
        if r_fonts.get(qn("w:hAnsi")) != expected_font.latin:
            mismatches.append(f"{prefix} hAnsi font expected {expected_font.latin}, found {r_fonts.get(qn('w:hAnsi'))}")
    size = rpr.find("w:sz", NS)
    if size is None or size.get(qn("w:val")) != expected_size:
        mismatches.append(f"{prefix} size expected {expected_size}, found {size.get(qn('w:val')) if size is not None else None}")
    color = rpr.find("w:color", NS)
    if color is None or color.get(qn("w:val")) != expected_font.color:
        mismatches.append(f"{prefix} color expected {expected_font.color}, found {color.get(qn('w:val')) if color is not None else None}")


def _inspect_appendix_style(document: Document, profile: FormatProfile) -> QualityIssue:
    appendix_paragraphs = _appendix_paragraphs(document)
    if not appendix_paragraphs:
        return QualityIssue(
            issue_id="docx_appendix",
            status="pass",
            severity="info",
            check_key="docx.appendix",
            title="DOCX appendix check is not applicable because no appendix was detected.",
            profile_rule_ref="appendix",
            details={"not_applicable": True},
        )
    mismatches: list[str] = []
    heading_count = 0
    body_count = 0
    for paragraph, kind in appendix_paragraphs:
        location = _paragraph_location(document, paragraph)
        if kind == "heading":
            heading_count += 1
            mismatches.extend(f"{location}: {item}" for item in _appendix_heading_mismatches(paragraph, profile))
        else:
            body_count += 1
            mismatches.extend(f"{location}: {item}" for item in _appendix_body_mismatches(paragraph, profile))
    if mismatches:
        return QualityIssue(
            issue_id="docx_appendix",
            status="fail",
            severity="high",
            check_key="docx.appendix",
            title="DOCX appendix formatting does not match the profile.",
            description="; ".join(mismatches[:20]),
            profile_rule_ref="appendix.title_font; appendix.body_font",
            recommendation="Apply profile appendix formatting before delivery.",
            fixable=True,
            details={"heading_count": heading_count, "body_count": body_count, "mismatch_count": len(mismatches)},
        )
    return QualityIssue(
        issue_id="docx_appendix",
        status="pass",
        severity="info",
        check_key="docx.appendix",
        title="DOCX appendix formatting matches the profile.",
        profile_rule_ref="appendix",
        details={"heading_count": heading_count, "body_count": body_count},
    )


def _appendix_paragraphs(document: Document) -> list[tuple[object, str]]:
    selected: list[tuple[object, str]] = []
    inside_appendix = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if _is_appendix_heading_text(text):
            inside_appendix = True
            selected.append((paragraph, "heading"))
            continue
        if inside_appendix and _is_appendix_terminator_text(text):
            inside_appendix = False
        elif inside_appendix:
            selected.append((paragraph, "body"))
    return selected


def _appendix_heading_mismatches(paragraph, profile: FormatProfile) -> list[str]:
    mismatches: list[str] = []
    if paragraph.alignment != _alignment_value(profile.appendix.title_alignment):
        mismatches.append(f"appendix title alignment expected {profile.appendix.title_alignment}")
    mismatches.extend(_paragraph_first_run_font_mismatches(paragraph, profile.appendix.title_font, "appendix title"))
    return mismatches


def _appendix_body_mismatches(paragraph, profile: FormatProfile) -> list[str]:
    mismatches: list[str] = []
    fmt = paragraph.paragraph_format
    expected_indent = profile.appendix.body_first_line_indent_chars * profile.appendix.body_font.size_pt
    actual_indent = fmt.first_line_indent.pt if fmt.first_line_indent else 0
    if abs(actual_indent - expected_indent) > 0.5:
        mismatches.append(f"appendix body first-line indent expected {expected_indent:.1f} pt, found {actual_indent:.1f} pt")
    if fmt.line_spacing != profile.appendix.body_line_spacing:
        mismatches.append(f"appendix body line spacing expected {profile.appendix.body_line_spacing}, found {fmt.line_spacing}")
    if paragraph.alignment != _alignment_value(profile.appendix.body_alignment):
        mismatches.append(f"appendix body alignment expected {profile.appendix.body_alignment}")
    mismatches.extend(_paragraph_first_run_font_mismatches(paragraph, profile.appendix.body_font, "appendix body"))
    return mismatches


def _paragraph_first_run_font_mismatches(paragraph, expected_font: TextFont, label: str) -> list[str]:
    run = next((item for item in paragraph.runs if item.text.strip()), None)
    if run is None:
        return [f"{label} paragraph has no text run"]
    mismatches: list[str] = []
    east_asia = run._element.rPr.rFonts.get(qn("w:eastAsia")) if run._element.rPr is not None and run._element.rPr.rFonts is not None else None
    if east_asia != expected_font.chinese:
        mismatches.append(f"{label} Chinese font expected {expected_font.chinese}, found {east_asia or 'auto/inherited'}")
    if run.font.name != expected_font.latin:
        mismatches.append(f"{label} Latin font expected {expected_font.latin}, found {run.font.name or 'auto/inherited'}")
    actual_size = run.font.size.pt if run.font.size else None
    if actual_size is None or abs(actual_size - expected_font.size_pt) > 0.1:
        found = f"{actual_size:g} pt" if actual_size else "inherited/auto"
        mismatches.append(f"{label} size expected {expected_font.size_pt:g} pt, found {found}")
    if not _run_color_matches(run, expected_font.color):
        actual = _run_color_value(run)
        mismatches.append(f"{label} color expected {expected_font.color}, found {actual or 'auto/inherited'}")
    actual_bold = bool(run.font.bold)
    expected_bold = expected_font.weight == "bold"
    if actual_bold != expected_bold:
        mismatches.append(f"{label} bold expected {expected_bold}, found {actual_bold}")
    return mismatches


def _is_appendix_heading_text(text: str) -> bool:
    return bool(re.match(r"^\s*(附录(?:\s*[A-ZＡ-Ｚ一二三四五六七八九十0-9]+)?|Appendix\b)", text.strip(), re.IGNORECASE))


def _is_appendix_terminator_text(text: str) -> bool:
    return text.strip() in {"参考文献", "References", "致谢", "Acknowledgements", "Acknowledgments"}


def _inspect_figure_sizes(features: OoxmlDocumentFeatures | None, profile: FormatProfile) -> QualityIssue:
    if features is None:
        return QualityIssue(
            issue_id="docx_figure_size",
            status="unsupported",
            check_key="docx.figure.size",
            title="DOCX inline figure sizes cannot be inspected.",
            recommendation="Regenerate the DOCX and rerun quality inspection.",
        )
    widths = list(features.inline_image_width_mm)
    if not widths:
        return QualityIssue(
            issue_id="docx_figure_size",
            status="pass",
            severity="info",
            check_key="docx.figure.size",
            title="DOCX figure-size check is not applicable because no inline images were detected.",
            profile_rule_ref="figure.size_rules",
            details={"inline_image_count": features.inline_image_count},
        )
    if features.inline_image_count > len(widths):
        return QualityIssue(
            issue_id="docx_figure_size",
            status="unsupported",
            severity="high",
            check_key="docx.figure.size",
            title="DOCX figure-size inspection could not read every inline image width.",
            description=f"Detected {features.inline_image_count} inline image(s), but only {len(widths)} width value(s) could be inspected.",
            profile_rule_ref="figure.size_rules",
            recommendation="Regenerate image markup or review image sizing manually before delivery.",
            fixable=False,
            details={"inline_image_count": features.inline_image_count, "inspected_width_count": len(widths)},
        )
    mismatches = [
        f"image[{index}] width {width:.1f} mm is outside half-column <= {profile.figure.half_column_max_mm:.1f} mm "
        f"or full-width {profile.figure.full_width_min_mm:.1f}-{profile.figure.full_width_max_mm:.1f} mm"
        for index, width in enumerate(widths, start=1)
        if not _figure_width_matches_profile(width, profile)
    ]
    if mismatches:
        return QualityIssue(
            issue_id="docx_figure_size",
            status="warning",
            severity="high",
            check_key="docx.figure.size",
            title="DOCX inline figure widths do not match the profile.",
            description="; ".join(mismatches[:12]),
            profile_rule_ref="figure.size_rules",
            recommendation="Resize inline figures to the profile half-column or full-width ranges.",
            fixable=True,
            details={
                "inline_image_count": features.inline_image_count,
                "inline_image_width_mm": widths,
                "half_column_max_mm": profile.figure.half_column_max_mm,
                "full_width_min_mm": profile.figure.full_width_min_mm,
                "full_width_max_mm": profile.figure.full_width_max_mm,
            },
        )
    return QualityIssue(
        issue_id="docx_figure_size",
        status="pass",
        severity="info",
        check_key="docx.figure.size",
        title="DOCX inline figure widths match the profile.",
        profile_rule_ref="figure.size_rules",
        details={
            "inline_image_count": features.inline_image_count,
            "inline_image_width_mm": widths,
            "half_column_max_mm": profile.figure.half_column_max_mm,
            "full_width_min_mm": profile.figure.full_width_min_mm,
            "full_width_max_mm": profile.figure.full_width_max_mm,
        },
    )


def _figure_width_matches_profile(width_mm: float, profile: FormatProfile) -> bool:
    tolerance_mm = 0.5
    return (
        width_mm <= profile.figure.half_column_max_mm + tolerance_mm
        or profile.figure.full_width_min_mm - tolerance_mm <= width_mm <= profile.figure.full_width_max_mm + tolerance_mm
    )


def _inspect_image_caption_pairing(
    document: Document,
    structure: DocumentStructure,
    features: OoxmlDocumentFeatures | None,
    profile: FormatProfile,
) -> QualityIssue:
    if features is None:
        return QualityIssue(
            issue_id="docx_visual_caption_pairing",
            status="unsupported",
            check_key="docx.visuals.caption_pairing",
            title="DOCX visual caption pairing cannot be inspected.",
            recommendation="Regenerate the DOCX and rerun quality inspection.",
        )
    table_captions = [
        paragraph
        for paragraph, classification in _paragraphs_with_classifications(document, structure)
        if classification.role == ParagraphRole.TABLE_CAPTION
    ]
    figure_captions = [
        paragraph
        for paragraph, classification in _paragraphs_with_classifications(document, structure)
        if classification.role == ParagraphRole.FIGURE_CAPTION
    ]
    total_images = features.inline_image_count + features.anchored_image_count
    if not document.tables and total_images == 0:
        return QualityIssue(
            issue_id="docx_visual_caption_pairing",
            status="pass",
            severity="info",
            check_key="docx.visuals.caption_pairing",
            title="DOCX visual caption pairing is not applicable because no tables or images were detected.",
        )
    if profile.figure.placement != "inline":
        if total_images > 0:
            return QualityIssue(
                issue_id="docx_visual_caption_pairing",
                status="unsupported",
                severity="high",
                check_key="docx.visuals.caption_pairing",
                title="DOCX figure placement is not fully auto-verified for non-inline layouts.",
                description=f"Profile requests {profile.figure.placement} figures and the current formatter only auto-verifies inline captions.",
                profile_rule_ref="figure.placement",
                recommendation="Use inline figures for fully automated delivery or review non-inline figures manually.",
                fixable=False,
                details={"figure_placement": profile.figure.placement, "inline_image_count": features.inline_image_count, "anchored_image_count": features.anchored_image_count},
            )
        return QualityIssue(
            issue_id="docx_visual_caption_pairing",
            status="pass",
            severity="info",
            check_key="docx.visuals.caption_pairing",
            title="DOCX figure placement is not applicable because no figures were detected.",
        )
    if features.anchored_image_count:
        return QualityIssue(
            issue_id="docx_visual_caption_pairing",
            status="unsupported",
            severity="high",
            check_key="docx.visuals.caption_pairing",
            title="DOCX contains floating/anchored images that require manual layout review.",
            description=f"Detected {features.anchored_image_count} anchored image(s).",
            profile_rule_ref="figure.caption; image.layout",
            recommendation="Review image anchoring and nearby captions in Word before treating the export as compliant.",
            fixable=False,
            details={"anchored_image_count": features.anchored_image_count},
        )
    table_texts = [
        paragraph.text.strip()
        for paragraph, classification in _paragraphs_with_classifications(document, structure)
        if classification.role == ParagraphRole.TABLE_CAPTION
    ]
    figure_texts = [
        paragraph.text.strip()
        for paragraph, classification in _paragraphs_with_classifications(document, structure)
        if classification.role == ParagraphRole.FIGURE_CAPTION
    ]
    caption_required_image_count = _caption_required_inline_image_count(document, profile)
    missing: list[str] = []
    chinese_table_count = sum(1 for text in table_texts if text.startswith(profile.table.caption.prefix))
    english_table_prefix = profile.table.caption.english_prefix or "Table"
    english_table_count = sum(1 for text in table_texts if text.startswith(english_table_prefix))
    chinese_figure_count = sum(1 for text in figure_texts if text.startswith(profile.figure.caption.prefix))
    english_figure_prefix = profile.figure.caption.english_prefix or "Figure"
    english_figure_count = sum(1 for text in figure_texts if text.startswith(english_figure_prefix))
    if len(document.tables) > chinese_table_count:
        missing.append(f"{len(document.tables) - chinese_table_count} table caption(s)")
    if caption_required_image_count > chinese_figure_count:
        missing.append(f"{caption_required_image_count - chinese_figure_count} figure caption(s)")
    if profile.table.caption.bilingual and len(document.tables) > english_table_count:
        missing.append(f"{len(document.tables) - english_table_count} English table caption(s)")
    if profile.figure.caption.bilingual and caption_required_image_count > english_figure_count:
        missing.append(f"{caption_required_image_count - english_figure_count} English figure caption(s)")
    position_mismatches = _visual_caption_position_mismatches(document, profile)
    bilingual_mismatches = _visual_bilingual_caption_mismatches(document, profile)
    if missing:
        return QualityIssue(
            issue_id="docx_visual_caption_pairing",
            status="warning",
            check_key="docx.visuals.caption_pairing",
            title="DOCX visual objects may be missing captions.",
            description=", ".join(missing),
            profile_rule_ref="table.caption; figure.caption",
            recommendation="Add or review captions so each table/image has a profile-compliant caption.",
            fixable=False,
            details={
                "table_count": len(document.tables),
                "table_caption_count": chinese_table_count,
                "inline_image_count": features.inline_image_count,
                "caption_required_inline_image_count": caption_required_image_count,
                "figure_caption_count": chinese_figure_count,
                "table_caption_position": profile.table.caption.position,
                "figure_caption_position": profile.figure.caption.position,
            },
        )
    if bilingual_mismatches:
        return QualityIssue(
            issue_id="docx_visual_caption_pairing",
            status="warning",
            check_key="docx.visuals.caption_pairing",
            title="DOCX bilingual visual captions are not paired with each visual object.",
            description="; ".join(bilingual_mismatches[:12]),
            profile_rule_ref="table.caption.bilingual; figure.caption.bilingual",
            recommendation="Keep both Chinese and English captions adjacent to the table/image they describe.",
            fixable=True,
            details={
                "table_count": len(document.tables),
                "inline_image_count": features.inline_image_count,
                "caption_required_inline_image_count": caption_required_image_count,
                "table_caption_position": _expected_table_caption_position(profile),
                "figure_caption_position": _expected_figure_caption_position(profile),
            },
        )
    if position_mismatches:
        return QualityIssue(
            issue_id="docx_visual_caption_pairing",
            status="warning",
            check_key="docx.visuals.caption_pairing",
            title="DOCX visual captions are not adjacent to the expected visual objects.",
            description="; ".join(position_mismatches[:12]),
            profile_rule_ref="table.caption.position; figure.caption.position",
            recommendation="Move table captions above tables and figure captions below images according to the selected profile.",
            fixable=True,
            details={
                "table_count": len(document.tables),
                "table_caption_count": chinese_table_count,
                "inline_image_count": features.inline_image_count,
                "caption_required_inline_image_count": caption_required_image_count,
                "figure_caption_count": chinese_figure_count,
                "table_caption_position": _expected_table_caption_position(profile),
                "figure_caption_position": _expected_figure_caption_position(profile),
            },
        )
    return QualityIssue(
        issue_id="docx_visual_caption_pairing",
        status="pass",
        severity="info",
        check_key="docx.visuals.caption_pairing",
        title="DOCX table/image counts have matching supported captions.",
        profile_rule_ref="table.caption; figure.caption",
        details={
            "table_count": len(document.tables),
            "table_caption_count": chinese_table_count,
            "inline_image_count": features.inline_image_count,
            "caption_required_inline_image_count": caption_required_image_count,
            "figure_caption_count": chinese_figure_count,
        },
    )


def _visual_caption_position_mismatches(document: Document, profile: FormatProfile) -> list[str]:
    mismatches: list[str] = []
    table_index = 0
    image_index = 0
    for child in document.element.body:
        if child.tag == qn("w:tbl"):
            table_index += 1
            expected_position = _expected_table_caption_position(profile)
            if not _element_has_nearby_caption(child, "table", expected_position, profile):
                mismatches.append(f"table[{table_index}] caption is not {expected_position} the table")
            continue
        if child.tag == qn("w:p") and child.xpath(".//*[local-name()='inline']"):
            if not _image_requires_caption(child, profile):
                continue
            image_index += 1
            expected_position = _expected_figure_caption_position(profile)
            if not _element_has_nearby_caption(child, "figure", expected_position, profile):
                mismatches.append(f"image[{image_index}] caption is not {expected_position} the image")
    return mismatches


def _visual_bilingual_caption_mismatches(document: Document, profile: FormatProfile) -> list[str]:
    mismatches: list[str] = []
    table_index = 0
    image_index = 0
    for child in document.element.body:
        if child.tag == qn("w:tbl"):
            table_index += 1
            if profile.table.caption.bilingual:
                captions = _nearby_caption_texts(child, "table", _expected_table_caption_position(profile), profile)
                if not _caption_group_has_prefix(captions, profile.table.caption.prefix, profile.table.caption.separator):
                    mismatches.append(f"table[{table_index}] missing Chinese caption near table")
                english_prefix = profile.table.caption.english_prefix or "Table"
                if not _caption_group_has_prefix(captions, english_prefix, profile.table.caption.separator):
                    mismatches.append(f"table[{table_index}] missing English caption near table")
            continue
        if child.tag == qn("w:p") and child.xpath(".//*[local-name()='inline']"):
            if not _image_requires_caption(child, profile):
                continue
            image_index += 1
            if profile.figure.caption.bilingual:
                captions = _nearby_caption_texts(child, "figure", _expected_figure_caption_position(profile), profile)
                if not _caption_group_has_prefix(captions, profile.figure.caption.prefix, profile.figure.caption.separator):
                    mismatches.append(f"image[{image_index}] missing Chinese caption near image")
                english_prefix = profile.figure.caption.english_prefix or "Figure"
                if not _caption_group_has_prefix(captions, english_prefix, profile.figure.caption.separator):
                    mismatches.append(f"image[{image_index}] missing English caption near image")
    return mismatches


def _inspect_visual_caption_fields(
    document: Document,
    features: OoxmlDocumentFeatures | None,
    profile: FormatProfile,
) -> list[QualityIssue]:
    if features is None:
        return [
            QualityIssue(
                issue_id=_quality_issue_id_from_check_key(check_key),
                status="unsupported",
                check_key=check_key,
                title=f"{title} cannot be inspected.",
                profile_rule_ref=field_path,
                recommendation="Regenerate the DOCX and rerun quality inspection.",
            )
            for field_path, check_key, title in VISUAL_CAPTION_FIELD_CHECKS
        ]
    total_images = features.inline_image_count + features.anchored_image_count
    if not document.tables and total_images == 0:
        return [
            QualityIssue(
                issue_id=_quality_issue_id_from_check_key(check_key),
                status="pass",
                severity="info",
                check_key=check_key,
                title=f"{title} is not applicable because no tables or images were detected.",
                profile_rule_ref=field_path,
                details={"not_applicable": True},
            )
            for field_path, check_key, title in VISUAL_CAPTION_FIELD_CHECKS
        ]
    position_mismatches = _visual_caption_position_mismatches_by_field(document, profile)
    bilingual_mismatches = _visual_bilingual_caption_mismatches_by_field(document, profile)
    if profile.figure.placement != "inline" and total_images > 0:
        message = (
            f"Profile requests {profile.figure.placement} figures and the current formatter only auto-verifies inline captions."
        )
        position_mismatches["figure.caption.position"].append(message)
        bilingual_mismatches["figure.caption.bilingual"].append(message)
    if features.anchored_image_count:
        message = f"Detected {features.anchored_image_count} anchored image(s)."
        position_mismatches["figure.caption.position"].append(message)
        bilingual_mismatches["figure.caption.bilingual"].append(message)
    issues: list[QualityIssue] = []
    for field_path, check_key, title in VISUAL_CAPTION_FIELD_CHECKS:
        mismatches = position_mismatches.get(field_path, []) + bilingual_mismatches.get(field_path, [])
        issues.append(
            _field_issue_from_mismatches(
                field_path,
                check_key,
                title,
                mismatches,
                location="visual objects",
                fixable=True,
            )
        )
    return issues


def _visual_caption_position_mismatches_by_field(document: Document, profile: FormatProfile) -> dict[str, list[str]]:
    mismatches = {
        "table.caption.position": [],
        "figure.caption.position": [],
    }
    table_index = 0
    image_index = 0
    for child in document.element.body:
        if child.tag == qn("w:tbl"):
            table_index += 1
            expected_position = _expected_table_caption_position(profile)
            if not _element_has_nearby_caption(child, "table", expected_position, profile):
                mismatches["table.caption.position"].append(f"table[{table_index}] caption is not {expected_position} the table")
            continue
        if child.tag == qn("w:p") and child.xpath(".//*[local-name()='inline']"):
            if not _image_requires_caption(child, profile):
                continue
            image_index += 1
            expected_position = _expected_figure_caption_position(profile)
            if not _element_has_nearby_caption(child, "figure", expected_position, profile):
                mismatches["figure.caption.position"].append(f"image[{image_index}] caption is not {expected_position} the image")
    return mismatches


def _visual_bilingual_caption_mismatches_by_field(document: Document, profile: FormatProfile) -> dict[str, list[str]]:
    mismatches = {
        "table.caption.bilingual": [],
        "figure.caption.bilingual": [],
    }
    table_index = 0
    image_index = 0
    for child in document.element.body:
        if child.tag == qn("w:tbl"):
            table_index += 1
            if profile.table.caption.bilingual:
                captions = _nearby_caption_texts(child, "table", _expected_table_caption_position(profile), profile)
                if not _caption_group_has_prefix(captions, profile.table.caption.prefix, profile.table.caption.separator):
                    mismatches["table.caption.bilingual"].append(f"table[{table_index}] missing Chinese caption near table")
                english_prefix = profile.table.caption.english_prefix or "Table"
                if not _caption_group_has_prefix(captions, english_prefix, profile.table.caption.separator):
                    mismatches["table.caption.bilingual"].append(f"table[{table_index}] missing English caption near table")
            continue
        if child.tag == qn("w:p") and child.xpath(".//*[local-name()='inline']"):
            if not _image_requires_caption(child, profile):
                continue
            image_index += 1
            if profile.figure.caption.bilingual:
                captions = _nearby_caption_texts(child, "figure", _expected_figure_caption_position(profile), profile)
                if not _caption_group_has_prefix(captions, profile.figure.caption.prefix, profile.figure.caption.separator):
                    mismatches["figure.caption.bilingual"].append(f"image[{image_index}] missing Chinese caption near image")
                english_prefix = profile.figure.caption.english_prefix or "Figure"
                if not _caption_group_has_prefix(captions, english_prefix, profile.figure.caption.separator):
                    mismatches["figure.caption.bilingual"].append(f"image[{image_index}] missing English caption near image")
    return mismatches


def _nearby_caption_texts(element, kind: str, position: str, profile: FormatProfile) -> list[str]:
    texts: list[str] = []
    sibling = element.getprevious() if position == "above" else element.getnext()
    while sibling is not None:
        if sibling.tag != qn("w:p") or sibling.xpath(".//*[local-name()='inline']"):
            break
        text = _paragraph_xml_text(sibling).strip()
        if not text:
            sibling = sibling.getprevious() if position == "above" else sibling.getnext()
            continue
        matches_caption = _table_caption_matches_profile(text, profile) if kind == "table" else _figure_caption_matches_profile(text, profile)
        if not matches_caption:
            break
        texts.append(text)
        sibling = sibling.getprevious() if position == "above" else sibling.getnext()
    return texts


def _caption_group_has_prefix(texts: list[str], prefix: str, separator: str) -> bool:
    return any(_caption_text_matches(text, prefix, separator) for text in texts)


def _expected_table_caption_position(profile: FormatProfile) -> str:
    return "above" if profile.table.enforce_caption_above else profile.table.caption.position


def _expected_figure_caption_position(profile: FormatProfile) -> str:
    return "below" if profile.figure.enforce_caption_below else profile.figure.caption.position


def _image_requires_caption(paragraph_element, profile: FormatProfile) -> bool:
    expected_position = _expected_figure_caption_position(profile)
    opposite_position = "below" if expected_position == "above" else "above"
    has_caption = _element_has_nearby_caption(paragraph_element, "figure", expected_position, profile) or _element_has_nearby_caption(
        paragraph_element,
        "figure",
        opposite_position,
        profile,
    )
    return has_caption or bool(_previous_paragraph_text(paragraph_element))


def _element_has_nearby_caption(element, kind: str, position: str, profile: FormatProfile) -> bool:
    sibling = element.getprevious() if position == "above" else element.getnext()
    while sibling is not None:
        if sibling.tag != qn("w:p"):
            return False
        text = _paragraph_xml_text(sibling).strip()
        if not text:
            sibling = sibling.getprevious() if position == "above" else sibling.getnext()
            continue
        if kind == "table":
            return _table_caption_matches_profile(text, profile)
        return _figure_caption_matches_profile(text, profile)
    return False


def _table_caption_matches_profile(text: str, profile: FormatProfile) -> bool:
    return any(
        _caption_text_matches(text, prefix, profile.table.caption.separator)
        for prefix in _expected_caption_prefixes(profile, ParagraphRole.TABLE_CAPTION)
    )


def _caption_required_inline_image_count(document: Document, profile: FormatProfile) -> int:
    count = 0
    caption_position = _expected_figure_caption_position(profile)
    opposite_position = "below" if caption_position == "above" else "above"
    for paragraph in document.paragraphs:
        if not _paragraph_has_inline_image(paragraph):
            continue
        has_caption = _has_nearby_figure_caption(paragraph._p, profile, caption_position) or _has_nearby_figure_caption(
            paragraph._p,
            profile,
            opposite_position,
        )
        # A leading logo/cover image with no nearby caption and no preceding prose is decorative front matter.
        if not has_caption and not _previous_paragraph_text(paragraph._p):
            continue
        count += 1
    return count


def _paragraph_has_inline_image(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//*[local-name()='inline']"))


def _has_nearby_figure_caption(paragraph_element, profile: FormatProfile, position: str) -> bool:
    sibling = paragraph_element.getprevious() if position == "above" else paragraph_element.getnext()
    while sibling is not None:
        if sibling.tag != qn("w:p"):
            sibling = sibling.getprevious() if position == "above" else sibling.getnext()
            continue
        text = _paragraph_xml_text(sibling).strip()
        if not text:
            sibling = sibling.getprevious() if position == "above" else sibling.getnext()
            continue
        return _figure_caption_matches_profile(text, profile)
    return False


def _figure_caption_matches_profile(text: str, profile: FormatProfile) -> bool:
    return any(
        _caption_text_matches(text, prefix, profile.figure.caption.separator)
        for prefix in _expected_caption_prefixes(profile, ParagraphRole.FIGURE_CAPTION)
    )


def _previous_paragraph_text(paragraph_element) -> str:
    sibling = paragraph_element.getprevious()
    while sibling is not None:
        if sibling.tag == qn("w:p"):
            text = _paragraph_xml_text(sibling).strip()
            if text:
                return text
        sibling = sibling.getprevious()
    return ""


def _paragraph_xml_text(paragraph_element) -> str:
    return "".join(node.text or "" for node in paragraph_element.xpath(".//w:t"))


def _inspect_list_numbering(features: OoxmlDocumentFeatures | None, profile: FormatProfile) -> QualityIssue:
    if features is None:
        return QualityIssue(
            issue_id="docx_numbering",
            status="unsupported",
            check_key="docx.numbering",
            title="DOCX numbering cannot be inspected.",
            recommendation="Regenerate the DOCX and rerun quality inspection.",
        )
    unsupported_requests: list[str] = []
    if profile.list_numbering.restart_per_section:
        unsupported_requests.append("list_numbering.restart_per_section")
    if profile.numbering.restart_per_section:
        unsupported_requests.append("numbering.restart_per_section")
    if unsupported_requests:
        return QualityIssue(
            issue_id="docx_numbering",
            status="unsupported",
            severity="high",
            check_key="docx.numbering",
            title="DOCX section-scoped numbering restart is not yet fully automated.",
            description=", ".join(unsupported_requests),
            profile_rule_ref="; ".join(unsupported_requests),
            recommendation="Disable section restart rules or handle numbering restart manually before final delivery.",
            fixable=False,
        )
    if features.numbering_reference_count == 0:
        return QualityIssue(
            issue_id="docx_numbering",
            status="pass",
            severity="info",
            check_key="docx.numbering",
            title="DOCX list-numbering check is not applicable because no numbering references were detected.",
        )
    if not profile.list_numbering.multilevel_enabled and features.numbering_reference_count > 0:
        return QualityIssue(
            issue_id="docx_numbering",
            status="warning",
            severity="high",
            check_key="docx.numbering",
            title="DOCX numbering references are present even though multilevel lists are disabled.",
            description=f"Detected {features.numbering_reference_count} numbering reference(s).",
            profile_rule_ref="list_numbering.multilevel_enabled",
            recommendation="Enable multilevel numbering or simplify the list structure.",
            fixable=False,
        )
    return QualityIssue(
        issue_id="docx_numbering",
        status="pass",
        severity="info",
        check_key="docx.numbering",
        title="DOCX numbering references are preserved for supported list paragraphs.",
        recommendation="Review numbering manually if the profile requires a specific multilevel numbering scheme.",
        details={"numbering_reference_count": features.numbering_reference_count},
    )


def _inspect_unit_rules(document: Document, profile: FormatProfile) -> QualityIssue:
    if not profile.unit_rules.enforce_consistency:
        return QualityIssue(
            issue_id="docx_unit_rules",
            status="pass",
            severity="info",
            check_key="docx.unit_rules",
            title="DOCX unit normalization is disabled by profile.",
        )
    mismatches: list[str] = []
    segments = _document_text_segments(document)
    if profile.unit_rules.normalize_fullwidth_numbers:
        fullwidth_locations = [location for location, text in segments if re.search(r"[０-９]", text)]
        if fullwidth_locations:
            mismatches.append(f"fullwidth digits remain in {', '.join(fullwidth_locations[:8])}")
    units = sorted(set(profile.unit_rules.measurement_units + profile.unit_rules.currency_units), key=len, reverse=True)
    if units:
        unit_pattern = "|".join(re.escape(unit) for unit in units)
        if profile.unit_rules.unit_spacing == "space":
            spacing_locations = [location for location, text in segments if re.search(rf"\d(?={unit_pattern})", text)]
            if spacing_locations:
                mismatches.append(f"number-unit spacing is missing in {', '.join(spacing_locations[:8])}")
        if profile.unit_rules.unit_spacing == "no_space":
            spacing_locations = [location for location, text in segments if re.search(rf"\d\s+(?:{unit_pattern})\b", text)]
            if spacing_locations:
                mismatches.append(f"number-unit spacing is present in {', '.join(spacing_locations[:8])}")
    if mismatches:
        return QualityIssue(
            issue_id="docx_unit_rules",
            status="warning",
            severity="high",
            check_key="docx.unit_rules",
            title="DOCX unit normalization needs review.",
            description="; ".join(mismatches),
            profile_rule_ref="unit_rules",
            recommendation="Re-run unit normalization before delivery.",
            fixable=True,
        )
    return QualityIssue(
        issue_id="docx_unit_rules",
        status="pass",
        severity="info",
        check_key="docx.unit_rules",
        title="DOCX unit normalization matches the profile.",
        profile_rule_ref="unit_rules",
    )


def _document_text_segments(document: Document) -> list[tuple[str, str]]:
    segments: list[tuple[str, str]] = [
        (f"paragraph[{index}]", paragraph.text)
        for index, paragraph in enumerate(document.paragraphs, start=1)
        if paragraph.text
    ]
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                for paragraph_index, paragraph in enumerate(cell.paragraphs, start=1):
                    if paragraph.text:
                        segments.append(
                            (
                                f"table[{table_index}].row[{row_index}].cell[{cell_index}].paragraph[{paragraph_index}]",
                                paragraph.text,
                            )
                        )
    return segments


def _section_has_page_field(section) -> bool:
    return _container_has_page_field(section.footer)


def _container_has_page_field(container) -> bool:
    for paragraph in container.paragraphs:
        xml = paragraph._p.xml
        if "PAGE" in xml and "fldChar" in xml:
            return True
    return False


def _close_cm(actual: float, expected: float, tolerance: float = 0.08) -> bool:
    return abs(actual - expected) <= tolerance


def _run_font_matches(run, expected_chinese: str, expected_latin: str) -> bool:
    east_asia = run._element.rPr.rFonts.get(qn("w:eastAsia")) if run._element.rPr is not None and run._element.rPr.rFonts is not None else None
    return run.font.name == expected_latin and east_asia == expected_chinese


def _run_style_mismatches(run, font: TextFont, label: str) -> list[str]:
    mismatches: list[str] = []
    if not _run_font_matches(run, font.chinese, font.latin):
        mismatches.append(f"{label} font does not match profile")
    actual_size = run.font.size.pt if run.font.size else None
    if actual_size is None or abs(actual_size - font.size_pt) > 0.1:
        mismatches.append(f"{label} size expected {font.size_pt:g} pt, found {actual_size:g} pt" if actual_size else f"{label} size expected {font.size_pt:g} pt, found inherited/auto")
    actual_bold = bool(run.font.bold)
    expected_bold = font.weight == "bold"
    if actual_bold != expected_bold:
        mismatches.append(f"{label} bold expected {expected_bold}, found {actual_bold}")
    if not _run_color_matches(run, font.color):
        actual = _run_color_value(run)
        mismatches.append(f"{label} color expected {font.color}, found {actual or 'auto/inherited'}")
    return mismatches


def _run_color_matches(run, expected_hex: str) -> bool:
    return run.font.color.rgb == RGBColor.from_string(expected_hex)


def _run_color_value(run) -> str | None:
    rgb = run.font.color.rgb
    return str(rgb) if rgb is not None else None


def _alignment_value(alignment: TextAlignment):
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping[alignment]


def _paragraphs_with_classifications(document: Document, structure: DocumentStructure):
    for index, paragraph in enumerate(document.paragraphs):
        yield paragraph, structure.role_for(index)


def _first_paragraph_with_classification(
    document: Document,
    structure: DocumentStructure,
    roles: set[ParagraphRole],
):
    return next(
        (
            (paragraph, classification)
            for paragraph, classification in _paragraphs_with_classifications(document, structure)
            if classification.role in roles
        ),
        None,
    )


def _first_paragraph_with_role(
    document: Document,
    structure: DocumentStructure,
    roles: set[ParagraphRole],
):
    selected = _first_paragraph_with_classification(document, structure, roles)
    return selected[0] if selected else None


def _paragraph_location(document: Document, paragraph) -> str:
    for index, candidate in enumerate(document.paragraphs, start=1):
        if candidate._p is paragraph._p:
            return f"paragraph[{index}]"
    return "paragraph[unknown]"


def _page_number_ooxml_format(format_name: str) -> str | None:
    return {
        "arabic": "decimal",
        "roman_lower": "lowerRoman",
        "roman_upper": "upperRoman",
    }.get(format_name)


def _spacing_pt(value) -> float | None:
    if value is None:
        return None
    try:
        return float(value.pt)
    except Exception:
        return None


def _paragraph_snap_to_grid(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:snapToGrid[@w:val='1' or not(@w:val)]"))


def _table_border_enabled(border, edge: str) -> bool:
    items = border.xpath(f"./*[local-name()='{edge}']")
    if not items:
        return False
    value = items[0].get(f"{{{NS['w']}}}val")
    return value not in {"nil", "none", "0"}


def _expected_caption_prefixes(profile: FormatProfile, role: ParagraphRole) -> list[str]:
    if role == ParagraphRole.TABLE_CAPTION:
        prefixes = [profile.table.caption.prefix]
        if profile.table.caption.bilingual:
            prefixes.append(profile.table.caption.english_prefix or "Table")
        return prefixes
    prefixes = [profile.figure.caption.prefix]
    if profile.figure.caption.bilingual:
        prefixes.append(profile.figure.caption.english_prefix or "Figure")
    return prefixes


def _caption_text_matches(text: str, prefix: str, separator: str) -> bool:
    escaped_separator = re.escape(separator or " ")
    return bool(
        re.match(rf"^{re.escape(prefix)}{escaped_separator}\d+(?:[.-]\d+)*\b", text)
        or re.match(rf"^{re.escape(prefix)}\d+(?:[.-]\d+)*\b", text)
    )


def _equation_has_visible_number(text: str) -> bool:
    return bool(re.match(r"^\(\d+\)\s+", text) or re.search(r"\s+\(\d+\)$", text))


def _contains_raw_latex(text: str) -> bool:
    return bool(
        re.search(r"(?<!\\)\$[^$]+\$(?!\\)", text)
        or re.search(r"\\(?:begin|end|frac|sum|int|sqrt|left|right|[()[\]])", text)
    )
