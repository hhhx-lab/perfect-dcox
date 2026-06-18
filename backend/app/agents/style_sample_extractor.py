from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from app.agents.extraction import ExtractionSourceError
from app.documents.converter import DocumentConversionError, convert_doc_to_docx
from app.models import ExtractionEvidence, FileRecord


@dataclass(frozen=True)
class StyleSampleAnalysis:
    text: str
    evidence: list[ExtractionEvidence] = field(default_factory=list)


def extract_style_sample_analysis(record: FileRecord, work_dir: Path, soffice_bin: str | None) -> StyleSampleAnalysis:
    input_path = Path(record.storage_path)
    try:
        docx_path = convert_doc_to_docx(input_path, work_dir, soffice_bin)
    except DocumentConversionError as exc:
        raise ExtractionSourceError(str(exc)) from exc

    try:
        document = Document(docx_path)
    except Exception as exc:  # noqa: BLE001 - python-docx wraps multiple OOXML parse failures.
        raise ExtractionSourceError(f"Style sample DOCX inspection failed: {exc}") from exc

    lines: list[str] = ["STYLE_SAMPLE_DOCX_INSPECTION", f"filename: {record.filename}"]
    evidence: list[ExtractionEvidence] = []

    if document.sections:
        section = document.sections[0]
        page_facts = {
            "page_width_mm": round(section.page_width.mm, 2),
            "page_height_mm": round(section.page_height.mm, 2),
            "top_margin_cm": round(section.top_margin.cm, 2),
            "bottom_margin_cm": round(section.bottom_margin.cm, 2),
            "left_margin_cm": round(section.left_margin.cm, 2),
            "right_margin_cm": round(section.right_margin.cm, 2),
        }
        lines.append(f"page: {page_facts}")
        evidence.append(
            ExtractionEvidence(
                field_path="page",
                source="style_sample_docx",
                quote=str(page_facts),
                note="第一节页面尺寸与页边距。",
                confidence=0.9,
            )
        )

    body_samples = _paragraph_samples(document.paragraphs, exclude_headings=True)
    if body_samples:
        lines.append(f"body_style_samples: {body_samples[:8]}")
        evidence.append(
            ExtractionEvidence(
                field_path="body",
                source="style_sample_docx",
                quote=str(body_samples[:8]),
                note="非标题段落 run/style 统计样本。",
                confidence=0.78,
            )
        )

    heading_samples = _paragraph_samples(document.paragraphs, headings_only=True)
    if heading_samples:
        lines.append(f"heading_style_samples: {heading_samples[:12]}")
        evidence.append(
            ExtractionEvidence(
                field_path="headings",
                source="style_sample_docx",
                quote=str(heading_samples[:12]),
                note="标题段落 run/style 统计样本。",
                confidence=0.82,
            )
        )

    header_footer = _header_footer_facts(document)
    if header_footer:
        lines.append(f"header_footer: {header_footer}")
        evidence.append(
            ExtractionEvidence(
                field_path="header_footer",
                source="style_sample_docx",
                quote=str(header_footer),
                note="各 section 页眉页脚文本。",
                confidence=0.78,
            )
        )

    if document.tables:
        table_facts = {"table_count": len(document.tables), "first_table_rows": len(document.tables[0].rows)}
        lines.append(f"tables: {table_facts}")
        evidence.append(
            ExtractionEvidence(
                field_path="table",
                source="style_sample_docx",
                quote=str(table_facts),
                note="样本文档表格数量与首表结构。",
                confidence=0.66,
            )
        )

    if document.inline_shapes:
        image_facts = [
            {"width_mm": round(shape.width.mm, 2), "height_mm": round(shape.height.mm, 2)}
            for shape in document.inline_shapes[:10]
        ]
        lines.append(f"figures: {image_facts}")
        evidence.append(
            ExtractionEvidence(
                field_path="figure",
                source="style_sample_docx",
                quote=str(image_facts),
                note="内嵌图片尺寸样本。",
                confidence=0.74,
            )
        )

    text_samples = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    if text_samples:
        lines.append(f"text_samples: {text_samples[:20]}")

    return StyleSampleAnalysis(text="\n".join(lines), evidence=evidence)


def _paragraph_samples(
    paragraphs: list[Any],
    *,
    headings_only: bool = False,
    exclude_headings: bool = False,
) -> list[dict[str, object]]:
    samples: list[dict[str, object]] = []
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        is_heading = style_name.lower().startswith("heading") or style_name.startswith("标题")
        if headings_only and not is_heading:
            continue
        if exclude_headings and is_heading:
            continue
        run = next((item for item in paragraph.runs if item.text.strip()), None)
        samples.append(
            {
                "paragraph_index": index,
                "text": text[:80],
                "style": style_name,
                "alignment": str(paragraph.alignment) if paragraph.alignment is not None else None,
                "font": _run_font_facts(run) if run is not None else _style_font_facts(paragraph),
                "paragraph": {
                    "first_line_indent_cm": _length_cm(paragraph.paragraph_format.first_line_indent),
                    "line_spacing": paragraph.paragraph_format.line_spacing,
                    "space_before_pt": _length_pt(paragraph.paragraph_format.space_before),
                    "space_after_pt": _length_pt(paragraph.paragraph_format.space_after),
                },
            }
        )
        if len(samples) >= 24:
            break
    return samples


def _run_font_facts(run: Any) -> dict[str, object]:
    font = run.font
    return {
        "east_asia": _east_asia_font(run),
        "name": font.name,
        "size_pt": round(font.size.pt, 2) if font.size is not None else None,
        "bold": font.bold,
        "italic": font.italic,
        "color": str(font.color.rgb) if font.color is not None and font.color.rgb is not None else None,
    }


def _style_font_facts(paragraph: Any) -> dict[str, object]:
    font = paragraph.style.font if paragraph.style is not None else None
    if font is None:
        return {}
    return {
        "name": font.name,
        "size_pt": round(font.size.pt, 2) if font.size is not None else None,
        "bold": font.bold,
        "italic": font.italic,
        "color": str(font.color.rgb) if font.color is not None and font.color.rgb is not None else None,
    }


def _east_asia_font(run: Any) -> str | None:
    rpr = run._element.rPr
    if rpr is None or rpr.rFonts is None:
        return None
    return rpr.rFonts.get(qn("w:eastAsia"))


def _header_footer_facts(document: Any) -> list[dict[str, object]]:
    facts: list[dict[str, object]] = []
    for index, section in enumerate(document.sections):
        header_text = "\n".join(paragraph.text.strip() for paragraph in section.header.paragraphs if paragraph.text.strip())
        footer_text = "\n".join(paragraph.text.strip() for paragraph in section.footer.paragraphs if paragraph.text.strip())
        if header_text or footer_text:
            facts.append({"section_index": index, "header_text": header_text, "footer_text": footer_text})
    return facts


def _length_cm(value: Any) -> float | None:
    return round(value.cm, 2) if value is not None else None


def _length_pt(value: Any) -> float | None:
    return round(value.pt, 2) if value is not None else None
