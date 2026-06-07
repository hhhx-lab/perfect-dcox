from __future__ import annotations

from dataclasses import dataclass
import json
from pathlib import Path
from typing import Protocol
from uuid import uuid4

from docx import Document
from pydantic import ValidationError
import yaml

from app.core.config import Settings
from app.documents.converter import DocumentConversionError, convert_doc_to_docx
from app.models import ExtractionEvidence, ExtractionSourceType, FileRecord, ProfileExtractionRecord, UncertainItem
from app.profiles.models import FormatProfile
from app.storage.repository import JsonMetadataRepository

SUPPORTED_RULE_SOURCE_EXTENSIONS = {".doc", ".docx"}


class ExtractionSourceError(ValueError):
    pass


@dataclass(frozen=True)
class ResolvedExtractionSource:
    source_type: ExtractionSourceType
    text: str
    file_record: FileRecord | None = None


@dataclass(frozen=True)
class ParsedAgentExtraction:
    profile_draft: FormatProfile
    uncertain_items: list[UncertainItem]
    evidence: list[ExtractionEvidence]


class RuleExtractionProvider(Protocol):
    def extract(self, source_text: str, source_meta: dict[str, str]) -> str:
        """Return raw structured Agent output for downstream parsing."""


class ConfiguredLLMRuleExtractionProvider:
    def __init__(self, settings: Settings) -> None:
        self.settings = settings

    def extract(self, source_text: str, source_meta: dict[str, str]) -> str:
        if not self.settings.llm_api_key or not self.settings.llm_model:
            raise ExtractionSourceError("LLM_API_KEY and LLM_MODEL are required for profile extraction.")
        raise ExtractionSourceError("Live LLM extraction provider is not implemented in this local MVP.")


def resolve_extraction_source(
    repository: JsonMetadataRepository,
    file_id: str | None,
    natural_language: str | None,
) -> ResolvedExtractionSource:
    text = (natural_language or "").strip()
    if file_id:
        record = repository.get_file(file_id)
        if record is None:
            raise ExtractionSourceError(f"Rule source file not found: {file_id}")
        suffix = Path(record.filename).suffix.lower()
        if suffix not in SUPPORTED_RULE_SOURCE_EXTENSIONS:
            raise ExtractionSourceError("Only .doc and .docx rule source files are supported.")
        return ResolvedExtractionSource(source_type="document", text="", file_record=record)
    if text:
        return ResolvedExtractionSource(source_type="natural_language", text=text)
    raise ExtractionSourceError("Either file_id or natural_language is required for profile extraction.")


def extract_rule_source_text(record: FileRecord, work_dir: Path, soffice_bin: str | None) -> str:
    input_path = Path(record.storage_path)
    try:
        docx_path = convert_doc_to_docx(input_path, work_dir, soffice_bin)
    except DocumentConversionError as exc:
        raise ExtractionSourceError(str(exc)) from exc

    try:
        document = Document(docx_path)
    except Exception as exc:  # noqa: BLE001 - python-docx exposes multiple low-level parse exceptions.
        raise ExtractionSourceError(f"Rule document text extraction failed: {exc}") from exc

    parts: list[str] = []
    parts.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)
    if not parts:
        raise ExtractionSourceError("Rule document does not contain extractable text.")
    return "\n".join(parts)


def parse_agent_extraction_output(raw_output: str) -> ParsedAgentExtraction:
    try:
        parsed = json.loads(raw_output)
    except json.JSONDecodeError:
        try:
            parsed = yaml.safe_load(raw_output)
        except yaml.YAMLError as exc:
            raise ExtractionSourceError(f"Agent output must be valid JSON or YAML: {exc}") from exc
    if not isinstance(parsed, dict):
        raise ExtractionSourceError("Agent output must be a structured object.")

    missing = [key for key in ("profile_draft", "uncertain_items", "evidence") if key not in parsed]
    if missing:
        raise ExtractionSourceError(f"Agent output missing required section(s): {', '.join(missing)}")
    try:
        profile = FormatProfile.model_validate(parsed["profile_draft"])
    except ValidationError as exc:
        raise ExtractionSourceError(f"Agent profile_draft failed schema validation: {exc}") from exc
    try:
        uncertain_items = [UncertainItem.model_validate(item) for item in parsed["uncertain_items"]]
        evidence = [ExtractionEvidence.model_validate(item) for item in parsed["evidence"]]
    except ValidationError as exc:
        raise ExtractionSourceError(f"Agent review metadata failed validation: {exc}") from exc
    if not evidence:
        raise ExtractionSourceError("Agent output evidence must contain at least one item.")
    return ParsedAgentExtraction(profile_draft=profile, uncertain_items=uncertain_items, evidence=evidence)


class ProfileExtractionService:
    def __init__(
        self,
        repository: JsonMetadataRepository,
        storage_root: Path,
        soffice_bin: str | None,
        provider: RuleExtractionProvider,
    ) -> None:
        self.repository = repository
        self.storage_root = storage_root
        self.soffice_bin = soffice_bin
        self.provider = provider

    def create_extraction(self, file_id: str | None, natural_language: str | None) -> ProfileExtractionRecord:
        source = resolve_extraction_source(self.repository, file_id, natural_language)
        record = ProfileExtractionRecord(
            extraction_id=f"extract_{uuid4().hex}",
            source_type=source.source_type,
            file_id=source.file_record.file_id if source.file_record else None,
            natural_language=source.text or None,
            status="queued",
        )
        return self.repository.add_profile_extraction(record)

    def process_extraction(self, extraction_id: str) -> ProfileExtractionRecord:
        record = self.repository.get_profile_extraction(extraction_id)
        if record is None:
            raise ExtractionSourceError(f"Profile extraction job not found: {extraction_id}")

        record.status = "running"
        record.error_message = None
        self.repository.update_profile_extraction(record)
        try:
            source_text = self._source_text_for_record(record)
            source_meta = {
                "source_type": record.source_type,
                "file_id": record.file_id or "",
                "extraction_id": record.extraction_id,
            }
            raw_output = self.provider.extract(source_text, source_meta)
            parsed = parse_agent_extraction_output(raw_output)
        except ExtractionSourceError as exc:
            record.status = "failed"
            record.error_message = str(exc)
            return self.repository.update_profile_extraction(record)

        record.status = "completed"
        record.profile_draft = parsed.profile_draft
        record.uncertain_items = parsed.uncertain_items
        record.evidence = parsed.evidence
        record.error_message = None
        return self.repository.update_profile_extraction(record)

    def _source_text_for_record(self, record: ProfileExtractionRecord) -> str:
        if record.source_type == "natural_language":
            if not record.natural_language or not record.natural_language.strip():
                raise ExtractionSourceError("Natural-language extraction source is empty.")
            return record.natural_language.strip()
        if not record.file_id:
            raise ExtractionSourceError("Document extraction source is missing file_id.")
        file_record = self.repository.get_file(record.file_id)
        if file_record is None:
            raise ExtractionSourceError(f"Rule source file not found: {record.file_id}")
        return extract_rule_source_text(file_record, self.storage_root / "work" / record.extraction_id, self.soffice_bin)
