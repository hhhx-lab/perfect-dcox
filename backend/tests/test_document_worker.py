from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.core.config import Settings
from app.documents.formatter import format_docx_with_profile
from app.documents.service import DocumentFormattingError, DocumentFormattingService
from app.jobs.worker import process_next_queued_job, process_placeholder_job
from app.main import create_app
from app.models import FileRecord, JobRecord
from app.profiles.models import ProfileCapabilityCoverage, ProfileUnsupportedRule
from app.profiles.seed import load_builtin_profiles
from app.quality.delivery_gate import InternalDeliveryGateService
from app.quality.final_layout_review import FinalLayoutReviewResult
from app.quality.inspection import inspect_docx_quality
from app.storage.local import LocalFileStorage
from app.storage.repository import JsonMetadataRepository
from tests.document_fixtures import add_ooxml_features, create_minimal_thesis_docx


def build_client(tmp_path: Path) -> TestClient:
    return TestClient(create_app(Settings(FILE_STORAGE_ROOT=tmp_path)))


def test_register_generated_output_file_metadata(tmp_path: Path) -> None:
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    storage = LocalFileStorage(tmp_path)
    source = tmp_path / "formatted.docx"
    source.write_bytes(b"formatted-docx")

    record = storage.store_generated_file(
        source,
        filename="formatted.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(record)

    loaded = repository.get_file(record.file_id)
    assert loaded is not None
    assert loaded.filename == "formatted.docx"
    assert loaded.size == len(b"formatted-docx")
    assert Path(loaded.storage_path).exists()

    client = build_client(tmp_path)
    response = client.get(f"/api/files/{record.file_id}")
    assert response.status_code == 200
    assert response.json()["filename"] == "formatted.docx"


def test_formatting_service_creates_docx_output_record(tmp_path: Path) -> None:
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    storage = LocalFileStorage(tmp_path)
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    input_record = storage.store_generated_file(
        source,
        filename="input.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(input_record)
    repository.save_profile_version(load_builtin_profiles()["ecnu_thesis"])
    service = DocumentFormattingService(repository, storage, soffice_bin=None)

    outputs = service.format_job(input_record.file_id, "ecnu_thesis", "1.0.0")

    assert len(outputs) == 1
    assert outputs[0].filename.endswith(".docx")
    assert repository.get_file(outputs[0].file_id) is not None
    formatter_registry = service.last_delivery_gate_summary["compile"]["formatter_registry"]
    assert "_apply_body_paragraph" in formatter_registry["executed_appliers"]
    assert "_apply_page_settings" in formatter_registry["executed_appliers"]
    assert "body.font.color" in formatter_registry["executed_field_paths"]
    assert "headings.font.color" in formatter_registry["executed_field_paths"]
    assert isinstance(formatter_registry["not_executed_field_paths"], list)
    registry = service.last_delivery_gate_summary["docx"]["rule_registry"]
    assert registry["all_covered"] is True
    assert registry["dispatch"]["all_executed"] is True
    assert "docx.body.font.color" in registry["dispatch"]["executed_check_keys"]


def test_formatting_service_applies_template_policy_and_inherits_header_footer(tmp_path: Path) -> None:
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    storage = LocalFileStorage(tmp_path)
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    input_record = storage.store_generated_file(
        source,
        filename="input.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(input_record)

    template_path = tmp_path / "template.docx"
    template_doc = Document()
    template_doc.sections[0].header.paragraphs[0].text = "模板页眉"
    template_doc.sections[0].footer.paragraphs[0].text = "模板页脚"
    template_doc.add_paragraph("模板固定封面")
    template_doc.add_paragraph("{{BODY}}")
    template_doc.add_paragraph("{{UNRESOLVED}}")
    template_doc.save(template_path)
    template_record = storage.store_generated_file(
        template_path,
        filename="template.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(template_record)

    base_profile = load_builtin_profiles()["ecnu_thesis"]
    profile = base_profile.model_copy(
        deep=True,
        update={
            "id": "template-profile",
            "template_binding": base_profile.template_binding.model_copy(
                update={
                    "body_slot": "{{BODY}}",
                    "inherit_header_footer": True,
                    "placeholder_policy": "remove",
                }
            ),
        },
    )
    repository.save_profile_version(profile)
    service = DocumentFormattingService(repository, storage, soffice_bin=None)

    outputs = service.format_job(input_record.file_id, profile.id, profile.version, template_file_id=template_record.file_id)

    output_doc = Document(outputs[0].storage_path)
    text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
    header_text = "\n".join(paragraph.text for paragraph in output_doc.sections[0].header.paragraphs)
    footer_text = "\n".join(paragraph.text for paragraph in output_doc.sections[0].footer.paragraphs)
    assert "模板固定封面" in text
    assert "这是一段正文内容，用于验证格式化后文本不会丢失。" in text
    assert "{{UNRESOLVED}}" not in text
    assert "模板页眉" in header_text
    assert "模板页脚" in footer_text
    assert service.last_delivery_gate_summary["compile"]["template_applied"] is True
    assert service.last_delivery_gate_summary["compile"]["header_footer_inherited"] is True
    assert service.last_delivery_gate_summary["docx"]["passed"] is True


def test_formatting_service_blocks_template_placeholder_residue(tmp_path: Path) -> None:
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    storage = LocalFileStorage(tmp_path)
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    input_record = storage.store_generated_file(
        source,
        filename="input.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(input_record)

    template_path = tmp_path / "template-with-residue.docx"
    template_doc = Document()
    template_doc.add_paragraph("模板固定封面")
    template_doc.add_paragraph("{{BODY}}")
    template_doc.add_paragraph("{{UNRESOLVED}}")
    template_doc.save(template_path)
    template_record = storage.store_generated_file(
        template_path,
        filename="template-with-residue.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(template_record)

    base_profile = load_builtin_profiles()["ecnu_thesis"]
    profile = base_profile.model_copy(deep=True)
    profile.id = "template-residue-profile"
    profile.template_binding.placeholder_policy = "preserve"
    repository.save_profile_version(profile)
    service = DocumentFormattingService(repository, storage, soffice_bin=None)

    try:
        service.format_job(input_record.file_id, profile.id, profile.version, template_file_id=template_record.file_id)
    except DocumentFormattingError as error:
        assert "placeholder" in str(error).lower()
        assert "{{UNRESOLVED}}" in str(error)
    else:
        raise AssertionError("Expected internal delivery gate to block unresolved template placeholders.")


def test_delivery_gate_blocks_missing_registered_docx_verifier_output(tmp_path: Path, monkeypatch) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)

    def fake_inspect_docx_quality(path, profile, *, inherited_header_footer=False):
        return [
            issue
            for issue in inspect_docx_quality(path, profile, inherited_header_footer=inherited_header_footer)
            if issue.check_key != "docx.figure.size"
        ]

    monkeypatch.setattr("app.quality.delivery_gate.inspect_docx_quality", fake_inspect_docx_quality)

    gate = InternalDeliveryGateService().validate_docx(formatted, profile, tmp_path)

    assert gate.passed is False
    assert "docx.figure.size" in (gate.failure_reason or "")
    assert any(issue.check_key == "profile.rule_registry.verifier_coverage" for issue in gate.issues)
    assert gate.public_summary()["rule_registry"]["missing_field_paths"] == ["figure.size_rules"]


def test_formatting_service_pdf_export_failure_is_diagnostic(tmp_path: Path) -> None:
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    storage = LocalFileStorage(tmp_path)
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    input_record = storage.store_generated_file(
        source,
        filename="input.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(input_record)
    repository.save_profile_version(load_builtin_profiles()["ecnu_thesis"])
    service = DocumentFormattingService(repository, storage, soffice_bin=None)

    try:
        service.format_job(input_record.file_id, "ecnu_thesis", "1.0.0", include_pdf=True)
    except DocumentFormattingError as error:
        assert "SOFFICE_BIN" in str(error)
    else:
        raise AssertionError("Expected PDF export failure when SOFFICE_BIN is missing.")


def test_formatting_service_runs_final_layout_review_for_pdf(tmp_path: Path, monkeypatch) -> None:
    class PassingReviewer:
        def review_pdf(self, payload):
            return FinalLayoutReviewResult(passed=True, summary=f"reviewed {payload.profile_id}", issues=[])

    def fake_export_docx_to_pdf(input_path, output_dir, soffice_bin):
        pdf_path = output_dir / f"{input_path.stem}.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n% fake final layout review fixture\n%%EOF\n")
        return pdf_path

    monkeypatch.setattr("app.documents.service.export_docx_to_pdf", fake_export_docx_to_pdf)
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    storage = LocalFileStorage(tmp_path)
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    input_record = storage.store_generated_file(
        source,
        filename="input.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(input_record)
    base_profile = load_builtin_profiles()["ecnu_thesis"]
    profile = base_profile.model_copy(deep=True)
    profile.id = "final-review-profile"
    profile.delivery_gate.require_pdf_inspection = False
    profile.llm_final_review.enabled = True
    profile.llm_final_review.required = True
    repository.save_profile_version(profile)
    service = DocumentFormattingService(
        repository,
        storage,
        soffice_bin="/fake/soffice",
        final_layout_reviewer=PassingReviewer(),
    )

    outputs = service.format_job(input_record.file_id, profile.id, profile.version, include_pdf=True)

    assert len(outputs) == 2
    assert outputs[1].filename.endswith(".pdf")
    assert service.last_delivery_gate_summary["llm_layout_review"]["passed"] is True
    assert service.last_delivery_gate_summary["llm_layout_review"]["remaining_issue_count"] == 0


def test_formatting_service_blocks_required_final_layout_review_without_reviewer(tmp_path: Path, monkeypatch) -> None:
    def fake_export_docx_to_pdf(input_path, output_dir, soffice_bin):
        pdf_path = output_dir / f"{input_path.stem}.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n% fake final layout review fixture\n%%EOF\n")
        return pdf_path

    monkeypatch.setattr("app.documents.service.export_docx_to_pdf", fake_export_docx_to_pdf)
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    storage = LocalFileStorage(tmp_path)
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    input_record = storage.store_generated_file(
        source,
        filename="input.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(input_record)
    base_profile = load_builtin_profiles()["ecnu_thesis"]
    profile = base_profile.model_copy(deep=True)
    profile.id = "required-final-review-profile"
    profile.delivery_gate.require_pdf_inspection = False
    profile.llm_final_review.enabled = True
    profile.llm_final_review.required = True
    repository.save_profile_version(profile)
    service = DocumentFormattingService(repository, storage, soffice_bin="/fake/soffice")

    try:
        service.format_job(input_record.file_id, profile.id, profile.version, include_pdf=True)
    except DocumentFormattingError as error:
        assert "Final LLM layout review failed" in str(error)
        assert "not configured" in str(error)
    else:
        raise AssertionError("Expected final LLM layout review to block without a reviewer.")


def test_worker_runs_document_engine_for_profile_referenced_job(tmp_path: Path) -> None:
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    storage = LocalFileStorage(tmp_path)
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    input_record = storage.store_generated_file(
        source,
        filename="input.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(input_record)
    repository.save_profile_version(load_builtin_profiles()["ecnu_thesis"])
    job = repository.add_job(
        JobRecord(
            job_id="job_profiled",
            job_type="placeholder_format",
            input_file_id=input_record.file_id,
            profile_id="ecnu_thesis",
            profile_version="1.0.0",
        )
    )

    completed = process_placeholder_job(repository, job.job_id, storage=storage, soffice_bin=None)

    assert completed.status == "completed"
    assert completed.progress == 100
    assert completed.output_file_ids
    assert repository.get_file(completed.output_file_ids[0]) is not None


def test_profiled_job_api_runs_document_engine_immediately(tmp_path: Path) -> None:
    client = build_client(tmp_path)
    source = create_minimal_thesis_docx(tmp_path / "input.docx")

    uploaded = client.post(
        "/api/files",
        files={
            "file": (
                "input.docx",
                source.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    file_id = uploaded.json()["file_id"]

    created = client.post(
        "/api/jobs",
        json={
            "input_file_id": file_id,
            "profile_id": "ecnu_thesis",
            "profile_version": "1.0.0",
        },
    )

    assert created.status_code == 201
    payload = created.json()
    assert payload["status"] == "completed"
    assert payload["progress"] == 100
    assert payload["output_file_ids"]
    output = client.get(f"/api/files/{payload['output_file_ids'][0]}")
    assert output.status_code == 200
    assert output.json()["filename"].endswith("-formatted.docx")


def test_batch_api_processes_multiple_profiled_documents_with_delivery_manifest(tmp_path: Path) -> None:
    client = build_client(tmp_path)
    source_a = create_minimal_thesis_docx(tmp_path / "input-a.docx")
    source_b = create_minimal_thesis_docx(tmp_path / "input-b.docx")
    uploaded_ids: list[str] = []
    for path in (source_a, source_b):
        response = client.post(
            "/api/files",
            files={
                "file": (
                    path.name,
                    path.read_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert response.status_code == 201
        uploaded_ids.append(response.json()["file_id"])

    created = client.post(
        "/api/batches",
        json={
            "profile_id": "ecnu_thesis",
            "profile_version": "1.0.0",
            "input_file_ids": uploaded_ids,
            "output_formats": ["docx", "pdf"],
            "auto_quality": True,
        },
    )

    assert created.status_code == 201
    payload = created.json()
    assert payload["batch_id"].startswith("batch_")
    assert payload["status"] == "completed"
    assert payload["delivery_manifest_id"].startswith("manifest_")
    assert payload["manifest_download_url"].endswith(f"/api/batches/{payload['batch_id']}/manifest")
    assert len(payload["job_ids"]) == 2
    assert len(payload["items"]) == 2
    for item in payload["items"]:
        assert item["input_file_id"] in uploaded_ids
        assert item["job_id"].startswith("job_")
        assert item["final_docx_file_id"].startswith("file_")
        assert item["final_pdf_file_id"].startswith("file_")
        assert item["quality_report_id"] is None
        assert item["fix_loop_ids"] == []
        assert item["delivery_gate_summary"]["docx"]["passed"] is True
        assert item["delivery_status"] == "completed"
        assert item["download_urls"]["docx"].endswith(f"/api/files/{item['final_docx_file_id']}/download")
        assert item["download_urls"]["pdf"].endswith(f"/api/files/{item['final_pdf_file_id']}/download")

    loaded = client.get(f"/api/batches/{payload['batch_id']}")
    assert loaded.status_code == 200
    assert loaded.json() == payload
    manifest = client.get(f"/api/batches/{payload['batch_id']}/manifest")
    assert manifest.status_code == 200
    assert manifest.json()["batch_id"] == payload["batch_id"]


def test_batch_api_maps_quality_failed_job_without_user_visible_report(tmp_path: Path, monkeypatch) -> None:
    client = build_client(tmp_path)
    source = create_minimal_thesis_docx(tmp_path / "broken-margin.docx")
    uploaded = client.post(
        "/api/files",
        files={
            "file": (
                source.name,
                source.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert uploaded.status_code == 201

    def fake_process_placeholder_job(repository, job_id, storage=None, soffice_bin=None, final_layout_reviewer=None):
        job = repository.get_job(job_id)
        assert job is not None
        job.status = "quality_failed"
        job.progress = 100
        job.current_step = "Internal delivery gate failed"
        job.error_message = "Internal delivery gate failed: margin mismatch"
        job.delivery_gate_summary = {"docx": {"passed": False, "failure_reason": "margin mismatch"}}
        return repository.update_job(job)

    monkeypatch.setattr("app.api.batches.process_placeholder_job", fake_process_placeholder_job)

    created = client.post(
        "/api/batches",
        json={
            "profile_id": "ecnu_thesis",
            "profile_version": "1.0.0",
            "input_file_ids": [uploaded.json()["file_id"]],
            "output_formats": ["docx"],
            "auto_quality": True,
            "auto_fix": True,
        },
    )

    assert created.status_code == 201
    payload = created.json()
    item = payload["items"][0]
    assert payload["status"] == "quality_failed"
    assert item["delivery_status"] == "manual_review_required"
    assert item["quality_report_id"] is None
    assert item["fix_loop_ids"] == []
    assert item["failure_reason"] == "Internal delivery gate failed: margin mismatch"
    assert item["delivery_gate_summary"]["docx"]["passed"] is False


def test_worker_requests_pdf_output_when_soffice_is_configured(tmp_path: Path, monkeypatch) -> None:
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    storage = LocalFileStorage(tmp_path)
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    input_record = storage.store_generated_file(
        source,
        filename="input.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(input_record)
    job = repository.add_job(
            JobRecord(
                job_id="job_pdf",
                job_type="placeholder_format",
                input_file_id=input_record.file_id,
                profile_id="ecnu_thesis",
                profile_version="1.0.0",
                output_formats=["docx", "pdf"],
            )
        )
    captured: dict[str, bool] = {}

    class FakeDocumentFormattingService:
        def __init__(self, repository, storage, soffice_bin, final_layout_reviewer=None) -> None:
            self.soffice_bin = soffice_bin

        def format_job(self, input_file_id, profile_id, profile_version, include_pdf=False, template_file_id=None):
            captured["include_pdf"] = include_pdf
            return [
                FileRecord(
                    file_id="file_docx",
                    filename="input-formatted.docx",
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    size=1,
                    sha256="0" * 64,
                    storage_path=str(tmp_path / "input-formatted.docx"),
                ),
                FileRecord(
                    file_id="file_pdf",
                    filename="input-formatted.pdf",
                    mime_type="application/pdf",
                    size=1,
                    sha256="1" * 64,
                    storage_path=str(tmp_path / "input-formatted.pdf"),
                ),
            ]

    monkeypatch.setattr("app.jobs.worker.DocumentFormattingService", FakeDocumentFormattingService)

    completed = process_placeholder_job(repository, job.job_id, storage=storage, soffice_bin="/usr/bin/soffice")

    assert captured["include_pdf"] is True
    assert completed.status == "completed"
    assert completed.output_file_ids == ["file_docx", "file_pdf"]


def test_formatting_service_blocks_profiles_with_unsupported_rules(tmp_path: Path) -> None:
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    storage = LocalFileStorage(tmp_path)
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    input_record = storage.store_generated_file(
        source,
        filename="input.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(input_record)
    profile = load_builtin_profiles()["ecnu_thesis"].model_copy(deep=True)
    profile.id = "unsupported-profile"
    profile.unsupported_rules = [
        ProfileUnsupportedRule(
            field_path="cover",
            message="封面规则当前不能自动验证。",
            suggestion="请绑定可执行模板或人工复核。",
        )
    ]
    repository.save_profile_version(profile)
    service = DocumentFormattingService(repository, storage, soffice_bin=None)

    try:
        service.format_job(input_record.file_id, profile.id, profile.version)
    except DocumentFormattingError as error:
        assert "unsupported rules" in str(error)
    else:
        raise AssertionError("Expected internal delivery gate to block unsupported rules.")


def test_delivery_gate_allows_inapplicable_notes_and_appendix_unsupported_rules(tmp_path: Path) -> None:
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    profile = load_builtin_profiles()["ecnu_thesis"].model_copy(deep=True)
    profile.id = "conditional-unsupported-profile"
    profile.unsupported_rules = [
        ProfileUnsupportedRule(
            field_path="notes",
            message="脚注/尾注规则当前不能自动执行。",
        ),
        ProfileUnsupportedRule(
            field_path="appendix",
            message="附录规则当前不能自动执行。",
        ),
    ]
    profile.capability_coverage = [
        ProfileCapabilityCoverage(
            field_path="notes",
            formatter="unsupported",
            qc="unsupported",
            unsupported_behavior="block",
        ),
        ProfileCapabilityCoverage(
            field_path="appendix",
            formatter="unsupported",
            qc="unsupported",
            unsupported_behavior="block",
        ),
    ]

    result = InternalDeliveryGateService().validate_docx(source, profile, tmp_path)

    assert result.passed is True


def test_delivery_gate_blocks_unknown_footnote_rule_when_document_contains_footnotes(tmp_path: Path) -> None:
    source = add_ooxml_features(create_minimal_thesis_docx(tmp_path / "input.docx"), footnote=True)
    profile = load_builtin_profiles()["ecnu_thesis"].model_copy(deep=True)
    profile.id = "unsupported-footnotes-profile"
    profile.unsupported_rules = [
        ProfileUnsupportedRule(
            field_path="footnotes.custom_separator",
            message="脚注分隔线自定义规则当前不能自动执行。",
        )
    ]

    result = InternalDeliveryGateService().validate_docx(source, profile, tmp_path)

    assert result.passed is False
    assert result.failure_reason == "Profile contains unsupported rules that cannot be verified."
    assert result.issues[0].profile_rule_ref == "footnotes.custom_separator"


def test_delivery_gate_ignores_stale_notes_unsupported_rule_after_notes_support(tmp_path: Path) -> None:
    source = add_ooxml_features(create_minimal_thesis_docx(tmp_path / "input.docx"), footnote=True)
    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", load_builtin_profiles()["ecnu_thesis"])
    profile = load_builtin_profiles()["ecnu_thesis"].model_copy(deep=True)
    profile.id = "stale-notes-unsupported-profile"
    profile.unsupported_rules = [
        ProfileUnsupportedRule(
            field_path="notes",
            message="旧 Agent Profile 曾标记脚注/尾注规则不可执行。",
        )
    ]

    result = InternalDeliveryGateService().validate_docx(formatted, profile, tmp_path)

    assert result.passed is True


def test_delivery_gate_blocks_unknown_appendix_rule_when_document_contains_appendix(tmp_path: Path) -> None:
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    document = Document(source)
    document.add_paragraph("附录 A 访谈提纲")
    document.save(source)
    profile = load_builtin_profiles()["ecnu_thesis"].model_copy(deep=True)
    profile.id = "unsupported-appendix-extra-profile"
    profile.unsupported_rules = [
        ProfileUnsupportedRule(
            field_path="appendix.figure_numbering",
            message="附录图表独立编号规则当前不能自动执行。",
        )
    ]

    result = InternalDeliveryGateService().validate_docx(source, profile, tmp_path)

    assert result.passed is False
    assert result.failure_reason == "Profile contains unsupported rules that cannot be verified."
    assert result.issues[0].profile_rule_ref == "appendix.figure_numbering"


def test_delivery_gate_ignores_stale_appendix_unsupported_rule_after_appendix_support(tmp_path: Path) -> None:
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    document = Document(source)
    document.add_paragraph("附录 A 访谈提纲")
    document.add_paragraph("这是附录正文。")
    document.save(source)
    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", load_builtin_profiles()["ecnu_thesis"])
    profile = load_builtin_profiles()["ecnu_thesis"].model_copy(deep=True)
    profile.id = "stale-appendix-unsupported-profile"
    profile.unsupported_rules = [
        ProfileUnsupportedRule(
            field_path="appendix",
            message="旧 Agent Profile 曾标记附录规则不可执行。",
        )
    ]

    result = InternalDeliveryGateService().validate_docx(formatted, profile, tmp_path)

    assert result.passed is True


def test_formatting_service_blocks_unsupported_capability_coverage(tmp_path: Path) -> None:
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    storage = LocalFileStorage(tmp_path)
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    input_record = storage.store_generated_file(
        source,
        filename="input.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(input_record)
    profile = load_builtin_profiles()["ecnu_thesis"].model_copy(deep=True)
    profile.id = "unsupported-coverage-profile"
    profile.capability_coverage = [
        ProfileCapabilityCoverage(
            field_path="figures.placement.anchored_layout",
            formatter="extract_only",
            qc="unsupported",
            unsupported_behavior="block",
            note="Agent extracted a figure placement rule that cannot be executed yet.",
        )
    ]
    repository.save_profile_version(profile)
    service = DocumentFormattingService(repository, storage, soffice_bin=None)

    try:
        service.format_job(input_record.file_id, profile.id, profile.version)
    except DocumentFormattingError as error:
        assert "unsupported rules" in str(error)
    else:
        raise AssertionError("Expected internal delivery gate to block unsupported capability coverage.")


def test_worker_reports_formatting_failure(tmp_path: Path) -> None:
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    storage = LocalFileStorage(tmp_path)
    repository.save_profile_version(load_builtin_profiles()["ecnu_thesis"])
    job = repository.add_job(
        JobRecord(
            job_id="job_missing_input",
            job_type="placeholder_format",
            input_file_id="file_missing",
            profile_id="ecnu_thesis",
            profile_version="1.0.0",
        )
    )

    failed = process_placeholder_job(repository, job.job_id, storage=storage, soffice_bin=None)

    assert failed.status == "failed"
    assert failed.progress == 100
    assert "Input file" in (failed.error_message or "")


def test_next_queued_job_processes_profiled_formatting_job(tmp_path: Path) -> None:
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    storage = LocalFileStorage(tmp_path)
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    input_record = storage.store_generated_file(
        source,
        filename="input.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(input_record)
    repository.save_profile_version(load_builtin_profiles()["ecnu_thesis"])
    repository.add_job(
        JobRecord(
            job_id="job_next",
            job_type="placeholder_format",
            input_file_id=input_record.file_id,
            profile_id="ecnu_thesis",
            profile_version="1.0.0",
        )
    )

    completed = process_next_queued_job(repository, storage=storage, soffice_bin=None)

    assert completed is not None
    assert completed.status == "completed"
    assert completed.output_file_ids


def test_worker_reports_missing_profile_version(tmp_path: Path) -> None:
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    storage = LocalFileStorage(tmp_path)
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    input_record = storage.store_generated_file(
        source,
        filename="input.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(input_record)
    job = repository.add_job(
        JobRecord(
            job_id="job_missing_profile",
            job_type="placeholder_format",
            input_file_id=input_record.file_id,
            profile_id="missing",
            profile_version="1.0.0",
        )
    )

    failed = process_placeholder_job(repository, job.job_id, storage=storage, soffice_bin=None)

    assert failed.status == "failed"
    assert "Profile version not found" in (failed.error_message or "")
