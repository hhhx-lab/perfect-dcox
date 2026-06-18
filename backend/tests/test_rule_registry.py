from docx import Document

from app.documents.rule_registry import (
    VERIFIED_PROFILE_FIELDS_DETAIL_KEY,
    RuleSpec,
    blocking_unsupported_capabilities,
    build_capability_coverage,
    docx_formatter_field_paths_by_applier,
    find_supported_rule_specs_without_handlers,
    find_supported_docx_verifier_keys_missing_from_issues,
    execute_docx_rule_verifiers,
    resolve_rule_applier,
    resolve_rule_verifier,
    summarize_docx_formatter_dispatch,
    supported_docx_formatter_applier_names,
    supported_docx_verifier_check_keys,
    registered_rule_specs,
    verify_docx_rule_registry_coverage,
)
from app.documents.formatter import FormatterExecutionTrace, format_docx_with_profile, formatter_pipeline_applier_names
from app.models import ExtractionEvidence, QualityIssue
from app.profiles.models import ProfileCapabilityCoverage
from app.profiles.seed import load_builtin_profiles
from app.quality.inspection import inspect_docx_quality
from tests.document_fixtures import create_minimal_thesis_docx


def test_rule_registry_has_handlers_for_every_supported_rule() -> None:
    specs = registered_rule_specs()

    missing = find_supported_rule_specs_without_handlers(specs)

    assert missing == []
    body_color = next(spec for spec in specs if spec.field_path == "body.font.color")
    assert body_color.formatter == "supported"
    assert body_color.qc == "supported"
    assert body_color.applier
    assert body_color.verifier


def test_supported_rule_handlers_resolve_to_callables() -> None:
    specs = registered_rule_specs()

    missing_appliers = [
        spec.field_path
        for spec in specs
        if spec.formatter == "supported" and spec.applier and not callable(resolve_rule_applier(spec))
    ]
    missing_verifiers = [
        spec.field_path
        for spec in specs
        if spec.qc == "supported" and spec.verifier and not callable(resolve_rule_verifier(spec))
    ]

    assert missing_appliers == []
    assert missing_verifiers == []


def test_supported_docx_formatter_appliers_are_declared_in_pipeline() -> None:
    registry_appliers = set(supported_docx_formatter_applier_names())
    pipeline_appliers = set(formatter_pipeline_applier_names())

    assert sorted(registry_appliers - pipeline_appliers) == []


def test_formatter_trace_records_registered_applier_execution(tmp_path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    source = create_minimal_thesis_docx(tmp_path / "source.docx")
    trace = FormatterExecutionTrace()

    format_docx_with_profile(source, tmp_path / "formatted.docx", profile, trace=trace)
    summary = trace.public_summary()
    fields_by_applier = docx_formatter_field_paths_by_applier()

    assert "body.font.color" in fields_by_applier["_apply_body_paragraph"]
    assert "_apply_page_settings" in summary["executed_appliers"]
    assert "_apply_body_paragraph" in summary["executed_appliers"]
    assert "_apply_table_rules" in summary["executed_appliers"]
    assert "body.font.color" in summary["executed_field_paths"]
    assert "figure.size_rules" in summary["executed_field_paths"]
    assert any(item["applier"] == "_apply_body_paragraph" and item["call_count"] > 0 for item in summary["items"])


def test_formatter_dispatch_summary_reports_unexecuted_fields() -> None:
    result = summarize_docx_formatter_dispatch(
        {
            "_apply_body_paragraph": 2,
            "_apply_page_settings": 1,
            "_unexpected_formatter_hook": 1,
        }
    )
    summary = result.public_summary()

    assert result.all_registered_appliers_executed is False
    assert "body.font.color" in summary["executed_field_paths"]
    assert "headings.font.color" in summary["not_executed_field_paths"]
    assert "_apply_heading_paragraph" in summary["missing_registered_appliers"]
    assert summary["unexpected_appliers"] == ["_unexpected_formatter_hook"]


def test_docx_rule_verifier_wrapper_returns_matching_quality_issue(tmp_path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    docx_path = create_minimal_thesis_docx(tmp_path / "source.docx")
    body_color = next(spec for spec in registered_rule_specs() if spec.field_path == "body.font.color")

    verifier = resolve_rule_verifier(body_color)
    issue = verifier(docx_path, profile) if verifier else None

    assert issue is not None
    assert issue.check_key == "docx.body.font.color"


def test_supported_docx_verifier_keys_are_present_in_quality_output(tmp_path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    source = create_minimal_thesis_docx(tmp_path / "source.docx")
    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)
    issues = inspect_docx_quality(formatted, profile)

    missing = find_supported_docx_verifier_keys_missing_from_issues(issues)

    assert "docx.figure.size" in supported_docx_verifier_check_keys()
    assert missing == []


def test_registry_executes_registered_docx_verifiers(tmp_path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    source = create_minimal_thesis_docx(tmp_path / "source.docx")
    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)

    result = execute_docx_rule_verifiers(formatted, profile)
    by_check_key = {item.check_key: item for item in result.dispatches}

    assert result.all_executed is True
    assert "docx.body.font.color" in by_check_key
    assert by_check_key["docx.body.font.color"].executed is True
    assert by_check_key["docx.template.body_slot"].executed is True
    assert any(issue.check_key == "docx.template.placeholders" for issue in result.issues)


def test_docx_registry_verification_maps_supported_fields_to_quality_issues(tmp_path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    source = create_minimal_thesis_docx(tmp_path / "source.docx")
    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)
    issues = inspect_docx_quality(formatted, profile)

    result = verify_docx_rule_registry_coverage(issues)
    by_field = {item.field_path: item for item in result.verifications}
    body_issue = next(issue for issue in issues if issue.check_key == "docx.body.font.color")

    assert result.all_covered is True
    assert "body.font.color" in body_issue.details[VERIFIED_PROFILE_FIELDS_DETAIL_KEY]
    assert by_field["body.font.color"].check_key == "docx.body.font.color"
    assert by_field["body.font.color"].issue_ids == ("docx_body_font_color",)
    assert by_field["figure.size_rules"].check_key == "docx.figure.size"
    assert by_field["figure.size_rules"].covered is True
    assert by_field["headings.font.color"].check_key == "docx.heading.font.color"
    assert by_field["headings.font.color"].issue_ids == ("docx_heading_font_color",)
    assert by_field["headings.pagination"].check_key == "docx.heading.pagination"
    assert by_field["headings.pagination"].covered is True
    assert by_field["header_footer.header_text"].check_key == "docx.header_footer.header_text"
    assert by_field["header_footer.header_text"].covered is True
    assert by_field["header_footer.page_number_format"].check_key == "docx.page_number.format"
    assert by_field["header_footer.page_number_format"].covered is True
    assert by_field["table.border_style"].check_key == "docx.table.border_style"
    assert by_field["table.header_repeat"].check_key == "docx.table.header_repeat"
    assert by_field["table.caption.position"].check_key == "docx.table.caption.position"
    assert by_field["figure.caption.position"].check_key == "docx.figure.caption.position"
    assert by_field["toc.title"].check_key == "docx.toc.title"
    assert by_field["toc.include_levels"].check_key == "docx.toc.include_levels"
    assert by_field["template_binding.body_slot"].check_key == "docx.template.body_slot"
    assert by_field["template_binding.placeholder_policy"].check_key == "docx.template.placeholders"


def test_template_binding_fields_have_field_level_quality_evidence(tmp_path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    source = tmp_path / "template-residue.docx"
    document = Document()
    document.add_paragraph("封面")
    document.add_paragraph(profile.template_binding.body_slot or "{{BODY}}")
    document.add_paragraph("{{UNRESOLVED}}")
    document.save(source)

    issues = inspect_docx_quality(source, profile)
    by_key = {issue.check_key: issue for issue in issues}

    assert by_key["docx.template.body_slot"].status == "warning"
    assert by_key["docx.template.placeholders"].status == "warning"
    assert "template_binding.body_slot" in by_key["docx.template.body_slot"].details[VERIFIED_PROFILE_FIELDS_DETAIL_KEY]
    assert "template_binding.placeholder_policy" in by_key["docx.template.placeholders"].details[
        VERIFIED_PROFILE_FIELDS_DETAIL_KEY
    ]


def test_docx_registry_verification_requires_field_level_issue_evidence() -> None:
    specs = [
        RuleSpec(
            field_path="body.font.color",
            formatter="supported",
            qc="supported",
            applier="_apply_body_paragraph",
            verifier="docx.body.style",
        )
    ]
    unannotated_issue = QualityIssue(
        issue_id="docx_body_style",
        status="pass",
        check_key="docx.body.style",
        title="DOCX body paragraph style matches the profile.",
    )

    missing = verify_docx_rule_registry_coverage([unannotated_issue], specs)
    annotated = unannotated_issue.model_copy(
        update={"details": {VERIFIED_PROFILE_FIELDS_DETAIL_KEY: ["body.font.color"]}}
    )
    covered = verify_docx_rule_registry_coverage([annotated], specs)

    assert missing.all_covered is False
    assert missing.missing_verifier_keys == ("docx.body.style",)
    assert missing.missing_field_paths == ("body.font.color",)
    assert covered.all_covered is True


def test_docx_registry_verification_reports_missing_field_paths(tmp_path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    source = create_minimal_thesis_docx(tmp_path / "source.docx")
    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)
    issues = [
        issue
        for issue in inspect_docx_quality(formatted, profile)
        if issue.check_key != "docx.figure.size"
    ]

    result = verify_docx_rule_registry_coverage(issues)

    assert result.all_covered is False
    assert result.missing_verifier_keys == ("docx.figure.size",)
    assert result.missing_field_paths == ("figure.size_rules",)
    assert result.public_summary()["missing_field_paths"] == ["figure.size_rules"]


def test_capability_coverage_uses_registry_and_blocks_unknown_fields() -> None:
    profile = load_builtin_profiles()["ecnu_thesis"].model_copy(deep=True)
    evidence = [
        ExtractionEvidence(
            field_path="body.font.color",
            source="natural_language",
            quote="正文黑色",
            confidence=0.9,
        ),
        ExtractionEvidence(
            field_path="decorative.watermark.opacity",
            source="rule_document",
            quote="水印透明度 35%",
            confidence=0.8,
        ),
    ]

    coverage = build_capability_coverage(profile, evidence, locked_fields=["body.font.color"])
    by_path = {item.field_path: item for item in coverage}

    assert by_path["body.font.color"].formatter == "supported"
    assert by_path["body.font.color"].qc == "supported"
    assert by_path["body.font.color"].locked_by_user is True
    assert by_path["decorative.watermark.opacity"].formatter == "unsupported"
    assert by_path["decorative.watermark.opacity"].qc == "unsupported"
    assert by_path["decorative.watermark.opacity"].unsupported_behavior == "block"


def test_capability_coverage_normalizes_indexed_profile_field_paths() -> None:
    profile = load_builtin_profiles()["ecnu_thesis"].model_copy(deep=True)
    evidence = [
        ExtractionEvidence(
            field_path="headings[1].font",
            source="rule_document",
            quote="一级标题黑体小三",
            confidence=0.88,
        )
    ]

    coverage = build_capability_coverage(profile, evidence, locked_fields=["headings[1].font"])
    by_path = {item.field_path: item for item in coverage}

    assert by_path["headings[1].font"].formatter == "supported"
    assert by_path["headings[1].font"].qc == "supported"
    assert by_path["headings[1].font"].locked_by_user is True


def test_blocking_unsupported_capabilities_ignores_stale_indexed_heading_coverage() -> None:
    profile = load_builtin_profiles()["ecnu_thesis"].model_copy(deep=True)
    profile.capability_coverage = [
        ProfileCapabilityCoverage(
            field_path="headings[1].font",
            formatter="unsupported",
            qc="unsupported",
            unsupported_behavior="block",
            note="Generated before indexed heading fields were normalized.",
        ),
        ProfileCapabilityCoverage(
            field_path="decorative.watermark.opacity",
            formatter="unsupported",
            qc="unsupported",
            unsupported_behavior="block",
        ),
    ]

    blocking = blocking_unsupported_capabilities(profile)

    assert [item.field_path for item in blocking] == ["decorative.watermark.opacity"]
