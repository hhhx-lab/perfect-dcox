import base64
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, RGBColor
from docx.shared import Pt

from app.documents.formatter import format_docx_with_profile
from app.documents.ooxml import inspect_ooxml_features
from app.profiles.seed import load_builtin_profiles
from app.quality.inspection import inspect_docx_quality
from tests.document_fixtures import add_ooxml_features, create_minimal_thesis_docx, read_docx_text

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def test_formatter_applies_ecnu_page_body_and_heading_rules(tmp_path: Path) -> None:
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    output = tmp_path / "formatted.docx"
    profile = load_builtin_profiles()["ecnu_thesis"]

    formatted = format_docx_with_profile(source, output, profile)

    document = Document(formatted)
    section = document.sections[0]
    assert round(section.top_margin.cm, 1) == 2.5
    assert round(section.bottom_margin.cm, 1) == 2.0
    assert round(section.left_margin.cm, 1) == 3.0
    assert round(section.right_margin.cm, 1) == 2.5

    body = document.paragraphs[1]
    assert round(body.paragraph_format.first_line_indent.cm, 2) == 0.74
    assert body.paragraph_format.line_spacing == 1.5
    assert body.runs[0].font.name == "Times New Roman"
    assert body.runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")) == "SimSun"

    heading = document.paragraphs[0]
    assert heading.style.name == "Heading 1"
    assert heading.runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert "PAGE" in section.footer.paragraphs[0]._p.xml
    formatted_text = read_docx_text(formatted)
    assert "这是一段正文内容，用于验证格式化后文本不会丢失。" in formatted_text
    assert "表 1 Sample table" in formatted_text


def test_formatter_applies_note_formatting_and_quality_checks_it(tmp_path: Path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    source = add_ooxml_features(create_minimal_thesis_docx(tmp_path / "input.docx"), footnote=True)
    raw_issues = inspect_docx_quality(source, profile)
    raw_notes_issue = next(issue for issue in raw_issues if issue.check_key == "docx.notes")

    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)
    formatted_issues = inspect_docx_quality(formatted, profile)
    formatted_notes_issue = next(issue for issue in formatted_issues if issue.check_key == "docx.notes")

    assert raw_notes_issue.status == "fail"
    assert formatted_notes_issue.status == "pass"
    assert formatted_notes_issue.details["footnote_count"] == 1


def test_formatter_applies_appendix_formatting_and_quality_checks_it(tmp_path: Path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    document = Document(source)
    document.add_paragraph("附录 A 访谈提纲")
    document.add_paragraph("这是附录正文。")
    document.save(source)
    raw_issues = inspect_docx_quality(source, profile)
    raw_appendix_issue = next(issue for issue in raw_issues if issue.check_key == "docx.appendix")

    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)
    formatted_issues = inspect_docx_quality(formatted, profile)
    formatted_appendix_issue = next(issue for issue in formatted_issues if issue.check_key == "docx.appendix")

    assert raw_appendix_issue.status == "fail"
    assert formatted_appendix_issue.status == "pass"
    assert formatted_appendix_issue.details["heading_count"] == 1
    assert formatted_appendix_issue.details["body_count"] == 1


def test_formatter_applies_profile_font_color(tmp_path: Path) -> None:
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    output = tmp_path / "formatted.docx"
    profile = load_builtin_profiles()["ecnu_thesis"]
    profile.body.font.color = "000000"
    for heading in profile.headings:
        heading.font.color = "000000"

    formatted = format_docx_with_profile(source, output, profile)

    document = Document(formatted)
    body = document.paragraphs[1]
    heading = document.paragraphs[0]
    assert body.runs[0].font.color.rgb == RGBColor(0, 0, 0)
    assert heading.runs[0].font.color.rgb == RGBColor(0, 0, 0)


def test_formatter_applies_profile_page_setup_header_and_page_number_toggle(tmp_path: Path) -> None:
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    output = tmp_path / "formatted.docx"
    base_profile = load_builtin_profiles()["ecnu_thesis"]
    profile = base_profile.model_copy(
        update={
            "page": base_profile.page.model_copy(update={"size": "Letter", "orientation": "landscape"}),
            "header_footer": base_profile.header_footer.model_copy(
                update={
                    "header_text": "测试模板页眉",
                    "header_alignment": "right",
                    "footer_page_number": False,
                }
            ),
        }
    )

    formatted = format_docx_with_profile(source, output, profile)

    document = Document(formatted)
    section = document.sections[0]
    assert section.orientation == WD_ORIENT.LANDSCAPE
    assert round(section.page_width.cm, 1) == 27.9
    assert round(section.page_height.cm, 1) == 21.6
    assert "测试模板页眉" in "\n".join(paragraph.text for paragraph in section.header.paragraphs)
    assert all("PAGE" not in paragraph._p.xml for paragraph in section.footer.paragraphs)
    assert inspect_ooxml_features(formatted).has_update_fields is True


def test_formatter_replaces_manual_toc_and_infers_table_captions(tmp_path: Path) -> None:
    source = tmp_path / "manual-toc.docx"
    document = Document()
    document.add_heading("论文题目", level=1)
    document.add_paragraph("目录")
    document.add_paragraph("摘要")
    document.add_paragraph("1、引言")
    document.add_paragraph("参考文献")
    document.add_paragraph("摘要")
    document.add_paragraph("摘要正文。")
    document.add_paragraph("1、引言")
    document.add_paragraph("表 1 给出一个课程化比较。")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "示例"
    document.save(source)

    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", load_builtin_profiles()["ecnu_thesis"])

    features = inspect_ooxml_features(formatted)
    formatted_doc = Document(formatted)
    text = read_docx_text(formatted)
    assert features.has_update_fields is True
    assert features.toc_field_count >= 1
    assert any("摘要" in item and "1、引言" in item for item in text)
    assert "表 1 课程化比较" in text
    caption = next(paragraph for paragraph in formatted_doc.paragraphs if paragraph.text == "表 1 课程化比较")
    assert caption.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_formatter_applies_special_paragraphs_and_table_borders(tmp_path: Path) -> None:
    source = create_minimal_thesis_docx(tmp_path / "input.docx")
    output = tmp_path / "formatted.docx"
    profile = load_builtin_profiles()["ecnu_thesis"]

    formatted = format_docx_with_profile(source, output, profile)

    document = Document(formatted)
    table_caption = next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith("表 1"))
    figure_caption = next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith("图 1"))
    equation = next(paragraph for paragraph in document.paragraphs if "E = mc" in paragraph.text)
    reference = next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith("[1]"))
    assert table_caption.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert figure_caption.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert equation.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert reference.paragraph_format.first_line_indent.cm < 0
    assert reference.paragraph_format.left_indent.cm > 0

    borders = document.tables[0]._tbl.tblPr.xpath("./w:tblBorders")
    assert borders
    assert borders[0].xpath("./*[local-name()='top']")
    assert borders[0].xpath("./*[local-name()='bottom']")


def test_formatter_preserves_toc_items_as_body_font_not_abstract_font(tmp_path: Path) -> None:
    source = tmp_path / "toc.docx"
    document = Document()
    document.add_heading("国内 RISC-V 架构现状及发展趋势", level=1)
    document.add_paragraph("目录")
    document.add_paragraph("摘要")
    document.add_paragraph("1、引言")
    document.add_paragraph("2、计算机系统结构视角下的 RISC-V")
    document.add_paragraph("参考文献")
    document.add_paragraph("致谢")
    document.add_paragraph("摘要")
    document.add_paragraph("摘要正文。")
    document.add_paragraph("关键词：RISC-V")
    document.add_paragraph("Abstract")
    document.add_paragraph("English abstract.")
    document.add_paragraph("Keywords: RISC-V")
    document.add_paragraph("1、引言")
    document.add_paragraph("正文段落。")
    document.save(source)

    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", load_builtin_profiles()["ecnu_thesis"])

    formatted_doc = Document(formatted)
    toc_field = next(paragraph for paragraph in formatted_doc.paragraphs if "TOC" in paragraph._p.xml)
    body = next(paragraph for paragraph in formatted_doc.paragraphs if paragraph.text == "正文段落。")
    assert "1、引言" in toc_field.text
    assert not any(
        paragraph.text in {"2、计算机系统结构视角下的 RISC-V", "参考文献", "致谢"} and paragraph._p.xml != toc_field._p.xml
        for paragraph in formatted_doc.paragraphs[:8]
    )
    assert body.runs[0].font.size == Pt(12)


def test_formatter_preserves_empty_and_unrecognized_paragraphs(tmp_path: Path) -> None:
    source = tmp_path / "input.docx"
    document = Document()
    document.add_paragraph("")
    document.add_paragraph("一个无法归类但必须保留的段落")
    document.save(source)
    output = tmp_path / "formatted.docx"

    formatted = format_docx_with_profile(source, output, load_builtin_profiles()["ecnu_thesis"])

    assert read_docx_text(formatted) == ["", "一个无法归类但必须保留的段落"]


def test_formatter_uses_document_structure_instead_of_template_specific_regexes(tmp_path: Path) -> None:
    source = tmp_path / "mixed.docx"
    document = Document()
    document.add_heading("国内 RISC-V 架构现状及发展趋势", level=1)
    document.add_paragraph("目录")
    document.add_paragraph("摘要")
    document.add_paragraph("1、引言")
    document.add_paragraph("2、计算机系统结构视角下的 RISC-V")
    document.add_paragraph("参考文献")
    document.add_paragraph("致谢")
    document.add_paragraph("摘要")
    document.add_paragraph("RISC-V 是近年来计算机系统结构领域最受关注的开放指令集之一。")
    document.add_paragraph("关键词：RISC-V；计算机系统结构；国产处理器")
    document.add_paragraph("Abstract")
    document.add_paragraph("RISC-V has become one of the most influential open instruction set architectures.")
    document.add_paragraph("Keywords: RISC-V; computer architecture")
    document.add_paragraph("1、引言")
    document.add_paragraph("RISC-V 的出现，使指令集架构重新成为系统结构课程中适合讨论的主题。")
    document.add_paragraph("2.1 ISA 与微体系结构的分层")
    document.add_paragraph("系统结构课程常用 CPU 执行时间公式说明这一点：")
    equation = document.add_paragraph("E = mc^2")
    equation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("表 1 给出一个课程化比较。")
    document.add_paragraph("表 1 RISC-V、Arm 与 x86 对比")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "维度"
    table.cell(0, 1).text = "RISC-V"
    table.cell(1, 0).text = "生态"
    table.cell(1, 1).text = "开放标准"
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] RISC-V International. Specifications [E]. https://riscv.org/specifications, 2026.")
    document.add_paragraph("[2] XUANTIE-RV. OpenC910 Core [E]. https://github.com/XUANTIE-RV/openc910.")
    document.add_paragraph("致谢")
    document.save(source)
    output = tmp_path / "formatted.docx"

    formatted = format_docx_with_profile(source, output, load_builtin_profiles()["ecnu_thesis"])

    formatted_doc = Document(formatted)
    by_text = {paragraph.text: paragraph for paragraph in formatted_doc.paragraphs if paragraph.text}
    toc_field = next(paragraph for paragraph in formatted_doc.paragraphs if "TOC" in paragraph._p.xml)
    abstract_body = by_text["RISC-V 是近年来计算机系统结构领域最受关注的开放指令集之一。"]
    body = by_text["RISC-V 的出现，使指令集架构重新成为系统结构课程中适合讨论的主题。"]
    keywords = by_text["关键词：RISC-V；计算机系统结构；国产处理器"]
    chapter_heading = next(
        paragraph
        for paragraph in formatted_doc.paragraphs
        if paragraph.text.endswith("引言") and paragraph.style and paragraph.style.name.startswith("Heading")
    )
    section_heading = by_text["2.1 ISA 与微体系结构的分层"]
    table_prose = by_text["表 1 给出一个课程化比较。"]
    table_caption = by_text["表 1 RISC-V、Arm 与 x86 对比"]
    reference = by_text["[1] RISC-V International. Specifications [E]. https://riscv.org/specifications, 2026."]

    assert "摘要" in toc_field.text
    assert "1、引言" in toc_field.text

    assert round(body.paragraph_format.first_line_indent.cm, 2) == 0.74
    assert body.paragraph_format.left_indent is None
    assert body.paragraph_format.line_spacing == 1.5
    assert body.runs[0].font.size == Pt(12)
    assert body.runs[0].font.bold is False

    assert abstract_body.paragraph_format.left_indent is None
    assert abstract_body.runs[0].font.size == Pt(10.5)
    assert abstract_body.runs[0].font.bold is False
    assert keywords.alignment != WD_ALIGN_PARAGRAPH.CENTER

    assert chapter_heading.style.name == "Heading 2"
    assert section_heading.style.name == "Heading 2"
    assert chapter_heading.runs[0].font.bold is True

    assert table_prose.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert table_prose.runs[0].font.size == Pt(12)
    assert table_caption.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert table_caption.runs[0].font.size == Pt(10.5)

    assert reference.alignment != WD_ALIGN_PARAGRAPH.CENTER
    assert reference.paragraph_format.first_line_indent.cm < 0
    assert reference.paragraph_format.left_indent.cm > 0


def test_formatter_applies_advanced_numbering_grid_caption_and_unit_rules(tmp_path: Path) -> None:
    source = tmp_path / "advanced.docx"
    document = Document()
    document.add_heading("绪论", level=1)
    document.add_paragraph("宽度为１００mm，长度为１０厘米，费用为５元。")
    document.add_paragraph("表 1 表格示例。")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "A"
    document.add_paragraph("图 1 插图示例。")
    image_path = tmp_path / "figure.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO9r9fkAAAAASUVORK5CYII="
        )
    )
    document.add_picture(str(image_path), width=Mm(80))
    document.save(source)

    profile = load_builtin_profiles()["ecnu_thesis"].model_copy(
        update={
            "numbering": load_builtin_profiles()["ecnu_thesis"].numbering.model_copy(
                update={"enabled": True, "heading_pattern": "第%1章", "restart_per_section": False}
            ),
            "unit_rules": load_builtin_profiles()["ecnu_thesis"].unit_rules.model_copy(
                update={"enforce_consistency": True, "unit_spacing": "space", "normalize_fullwidth_numbers": True}
            ),
            "document_grid": load_builtin_profiles()["ecnu_thesis"].document_grid.model_copy(
                update={"enabled": True, "type": "line_and_character", "characters_per_line": 40, "lines_per_page": 28, "snap_to_grid": True}
            ),
            "header_footer": load_builtin_profiles()["ecnu_thesis"].header_footer.model_copy(
                update={"page_number_format": "roman_upper", "page_number_start": 3, "footer_page_number": True}
            ),
            "table": load_builtin_profiles()["ecnu_thesis"].table.model_copy(
                update={
                    "caption": load_builtin_profiles()["ecnu_thesis"].table.caption.model_copy(
                        update={"bilingual": True, "english_prefix": "Table", "separator": " / "}
                    ),
                    "enforce_caption_above": True,
                }
            ),
            "figure": load_builtin_profiles()["ecnu_thesis"].figure.model_copy(
                update={
                    "caption": load_builtin_profiles()["ecnu_thesis"].figure.caption.model_copy(
                        update={"bilingual": True, "english_prefix": "Figure", "separator": " / "}
                    ),
                    "enforce_caption_below": True,
                }
            ),
        }
    )

    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)

    formatted_doc = Document(formatted)
    paragraphs = [paragraph.text for paragraph in formatted_doc.paragraphs if paragraph.text.strip()]
    assert any(text.startswith("第1章") for text in paragraphs)
    assert any("100 mm" in text and "10 cm" in text and "5 元" in text for text in paragraphs)
    assert any(text.startswith("表 / 1 / 表格示例") for text in paragraphs)
    assert any(text.startswith("Table / 1 / 表格示例") for text in paragraphs)
    assert any(text.startswith("图 / 1 / 插图示例") for text in paragraphs)
    assert any(text.startswith("Figure / 1 / 插图示例") for text in paragraphs)
    body_order: list[tuple[str, str]] = []
    for child in formatted_doc.element.body:
        if child.tag == qn("w:tbl"):
            body_order.append(("table", ""))
        elif child.tag == qn("w:p"):
            text = "".join(node.text or "" for node in child.xpath(".//w:t"))
            body_order.append(("image" if child.xpath(".//*[local-name()='inline']") else "paragraph", text))
    table_index = next(index for index, (kind, _) in enumerate(body_order) if kind == "table")
    table_zh_index = next(index for index, (_, text) in enumerate(body_order) if text.startswith("表 / 1 /"))
    table_en_index = next(index for index, (_, text) in enumerate(body_order) if text.startswith("Table / 1 /"))
    image_index = next(index for index, (kind, _) in enumerate(body_order) if kind == "image")
    figure_zh_index = next(index for index, (_, text) in enumerate(body_order) if text.startswith("图 / 1 /"))
    figure_en_index = next(index for index, (_, text) in enumerate(body_order) if text.startswith("Figure / 1 /"))
    assert table_zh_index < table_en_index < table_index
    assert image_index < figure_zh_index < figure_en_index

    features = inspect_ooxml_features(formatted)
    assert features.document_grid_types[0] == "linesAndChars"
    assert features.page_number_formats[0] == "upperRoman"
    assert features.page_number_starts[0] == 3
    assert features.inline_image_width_mm and round(features.inline_image_width_mm[0], 0) == 100
    assert "PAGE" in formatted_doc.sections[0].footer.paragraphs[0]._p.xml
    assert "ROMAN" in formatted_doc.sections[0].footer.paragraphs[0]._p.xml


def test_formatter_moves_existing_table_caption_group_next_to_table(tmp_path: Path) -> None:
    source = tmp_path / "separated-table-caption.docx"
    document = Document()
    document.add_heading("绪论", level=1)
    document.add_paragraph("表 1 国内 RISC-V 典型案例分层比较")
    document.add_paragraph("Table 1 Layered Comparison of Representative Domestic RISC-V Cases")
    document.add_paragraph("为了避免把不同层级案例混在一起，本文采用分层框架。")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "示例"
    document.save(source)
    base_profile = load_builtin_profiles()["ecnu_thesis"]
    profile = base_profile.model_copy(
        update={
            "table": base_profile.table.model_copy(
                update={
                    "caption": base_profile.table.caption.model_copy(
                        update={"bilingual": True, "english_prefix": "Table"}
                    )
                }
            )
        }
    )

    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)

    formatted_doc = Document(formatted)
    body_order: list[tuple[str, str]] = []
    for child in formatted_doc.element.body:
        if child.tag == qn("w:tbl"):
            body_order.append(("table", ""))
        elif child.tag == qn("w:p"):
            text = "".join(node.text or "" for node in child.xpath(".//w:t")).strip()
            if text:
                body_order.append(("paragraph", text))
    table_index = next(index for index, (kind, _) in enumerate(body_order) if kind == "table")
    assert body_order[table_index - 2][1].startswith("表 1")
    assert body_order[table_index - 1][1].startswith("Table 1")
    by_check = {issue.check_key: issue for issue in inspect_docx_quality(formatted, profile)}
    assert by_check["docx.table.caption.position"].status == "pass"
    assert by_check["docx.table.caption.bilingual"].status == "pass"


def test_formatter_does_not_number_toc_title_on_second_pass(tmp_path: Path) -> None:
    source = tmp_path / "toc-numbering.docx"
    document = Document()
    document.add_heading("目录", level=2)
    document.add_paragraph("摘要")
    document.add_paragraph("1、引言")
    document.add_paragraph("摘要")
    document.add_paragraph("摘要正文。")
    document.add_paragraph("1、引言")
    document.add_paragraph("正文段落。")
    document.save(source)

    base_profile = load_builtin_profiles()["ecnu_thesis"]
    profile = base_profile.model_copy(
        update={"numbering": base_profile.numbering.model_copy(update={"enabled": True, "heading_pattern": "%1"})}
    )

    second_pass = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)

    paragraphs = [paragraph.text.strip() for paragraph in Document(second_pass).paragraphs if paragraph.text.strip()]
    assert profile.toc.title in paragraphs
    assert f"1 {profile.toc.title}" not in paragraphs


def test_formatter_normalizes_unit_rules_inside_table_cells(tmp_path: Path) -> None:
    source = tmp_path / "table-units.docx"
    document = Document()
    document.add_heading("第一章 绪论", level=1)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "宽度"
    table.cell(0, 1).text = "费用"
    table.cell(1, 0).text = "１００mm"
    table.cell(1, 1).text = "５元"
    document.save(source)
    base_profile = load_builtin_profiles()["ecnu_thesis"]
    profile = base_profile.model_copy(
        update={
            "unit_rules": base_profile.unit_rules.model_copy(
                update={"enforce_consistency": True, "unit_spacing": "space", "normalize_fullwidth_numbers": True}
            )
        }
    )

    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)

    formatted_doc = Document(formatted)
    assert formatted_doc.tables[0].cell(1, 0).text == "100 mm"
    assert formatted_doc.tables[0].cell(1, 1).text == "5 元"


def test_formatter_honors_caption_numbering_toc_and_equation_options(tmp_path: Path) -> None:
    source = tmp_path / "numbering-options.docx"
    document = Document()
    document.add_heading("论文题目", level=1)
    document.add_paragraph("目录")
    document.add_paragraph("1、绪论")
    document.add_paragraph("摘要")
    document.add_paragraph("摘要正文。")
    document.add_paragraph("关键词：格式")
    document.add_paragraph("1、绪论")
    document.add_paragraph("1.1 方法")
    document.add_paragraph("E = mc^2")
    document.add_paragraph("表 1 方法对比。")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "A"
    document.add_paragraph("图 1 方法流程。")
    image_path = tmp_path / "figure.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO9r9fkAAAAASUVORK5CYII="
        )
    )
    document.add_picture(str(image_path), width=Mm(40))
    document.save(source)

    base_profile = load_builtin_profiles()["ecnu_thesis"]
    profile = base_profile.model_copy(
        update={
            "toc": base_profile.toc.model_copy(update={"right_align_page_numbers": False}),
            "equations": base_profile.equations.model_copy(update={"numbering": "left"}),
            "table": base_profile.table.model_copy(
                update={
                    "caption": base_profile.table.caption.model_copy(
                        update={"numbering": "chapter", "separator": " "}
                    )
                }
            ),
            "figure": base_profile.figure.model_copy(
                update={
                    "caption": base_profile.figure.caption.model_copy(
                        update={"numbering": "section", "separator": " "}
                    )
                }
            ),
        }
    )

    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)

    formatted_doc = Document(formatted)
    paragraphs = [paragraph.text for paragraph in formatted_doc.paragraphs if paragraph.text.strip()]
    features = inspect_ooxml_features(formatted)
    assert any(r'\p " "' in instruction for instruction in features.toc_instructions)
    assert any(text.startswith("(1) E = mc^2") for text in paragraphs)
    assert any(text.startswith("表 1-1 方法对比") for text in paragraphs)
    assert any(text.startswith("图 1.1-1 方法流程") for text in paragraphs)


def test_formatter_skips_uncaptioned_front_matter_images(tmp_path: Path) -> None:
    source = tmp_path / "cover-image.docx"
    image_path = tmp_path / "figure.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO9r9fkAAAAASUVORK5CYII="
        )
    )
    document = Document()
    document.add_picture(str(image_path), width=Mm(20))
    document.add_heading("论文题目", level=1)
    document.add_paragraph("姓名：测试")
    document.add_paragraph("1、引言")
    document.add_paragraph("图 1 正文插图。")
    document.add_picture(str(image_path), width=Mm(40))
    document.save(source)

    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", load_builtin_profiles()["ecnu_thesis"])

    formatted_doc = Document(formatted)
    paragraphs = [paragraph.text for paragraph in formatted_doc.paragraphs if paragraph.text.strip()]
    figure_captions = [text for text in paragraphs if text.startswith("图 ")]
    assert figure_captions == ["图 1 正文插图"]
    title_index = paragraphs.index("论文题目")
    assert all(not text.startswith("图 ") for text in paragraphs[:title_index])


def test_formatter_does_not_crash_when_heading_style_is_missing(tmp_path: Path) -> None:
    source = tmp_path / "missing-heading-style.docx"
    document = Document()
    document.add_heading("论文题目", level=1)
    document.add_paragraph("2.1 方法")
    document.add_paragraph("正文内容。")
    document.save(source)
    _remove_style_from_docx(source, "Heading2")

    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", load_builtin_profiles()["ecnu_thesis"])

    formatted_doc = Document(formatted)
    heading = next(paragraph for paragraph in formatted_doc.paragraphs if paragraph.text == "2.1 方法")
    assert heading.style.name == "Heading 2"
    assert heading._p.xpath("./w:pPr/w:outlineLvl[@w:val='1']")


def _remove_style_from_docx(path: Path, style_id: str) -> None:
    with ZipFile(path) as source_package:
        styles_root = ET.fromstring(source_package.read("word/styles.xml"))
        for style in list(styles_root.findall(f"{{{W_NS}}}style")):
            if style.get(f"{{{W_NS}}}styleId") == style_id:
                styles_root.remove(style)
        with NamedTemporaryFile(delete=False, suffix=".docx", dir=path.parent) as tmp_file:
            tmp_path = Path(tmp_file.name)
        try:
            with ZipFile(tmp_path, "w") as target_package:
                for item in source_package.infolist():
                    if item.filename == "word/styles.xml":
                        continue
                    target_package.writestr(item, source_package.read(item.filename))
                ET.register_namespace("w", W_NS)
                target_package.writestr(
                    "word/styles.xml",
                    ET.tostring(styles_root, encoding="utf-8", xml_declaration=True),
                )
            tmp_path.replace(path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink()


def qn(name: str) -> str:
    from docx.oxml.ns import qn as docx_qn

    return docx_qn(name)
