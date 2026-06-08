import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, RGBColor
from fastapi.testclient import TestClient
from pydantic import ValidationError

from app.core.config import Settings
from app.documents.formatter import format_docx_with_profile
from app.main import create_app
from app.models import (
    FixAction,
    IssueExplanation,
    FixLoopRecord,
    FixPlan,
    QualityIssue,
    QualityReport,
    QualitySummary,
)
from app.profiles.seed import load_builtin_profiles
from app.quality.inspection import inspect_docx_quality, inspect_pdf_quality
from app.quality.fix_planning import FixPlanSafetyError, FixPlanService, validate_fix_plan
from app.quality.service import QualityReportService
from app.storage.repository import JsonMetadataRepository
from app.storage.local import LocalFileStorage
from tests.document_fixtures import add_ooxml_features, create_minimal_thesis_docx


def test_quality_report_models_serialize_status_groups_and_issue_metadata() -> None:
    issue = QualityIssue(
        issue_id="issue_margin_top",
        status="fail",
        severity="high",
        check_key="docx.page.margin.top",
        title="Top margin does not match profile.",
        description="Expected 2.5 cm, found 1.9 cm.",
        profile_rule_ref="page.margins_cm.top",
        location="section[0]",
        recommendation="Reapply the selected profile margins.",
        fixable=True,
    )
    summary = QualitySummary.from_issues([issue])
    report = QualityReport(
        report_id="qr_123",
        job_id="job_123",
        profile_id="ecnu_thesis",
        profile_version="1.0.0",
        output_file_ids=["file_docx"],
        summary=summary,
        issues=[issue],
    )

    payload = report.model_dump(mode="json")

    assert payload["summary"]["counts"]["fail"] == 1
    assert payload["summary"]["remaining_issue_count"] == 1
    assert payload["summary"]["all_compliant"] is False
    assert payload["issues"][0]["profile_rule_ref"] == "page.margins_cm.top"
    assert payload["issues"][0]["location"] == "section[0]"
    assert payload["issues"][0]["recommendation"] == "Reapply the selected profile margins."
    assert payload["created_at"]
    assert payload["updated_at"]


def test_quality_summary_counts_all_supported_statuses() -> None:
    issues = [
        QualityIssue(issue_id="pass_1", status="pass", check_key="docx.open", title="Openable"),
        QualityIssue(issue_id="fixed_1", status="fixed", check_key="docx.table", title="Table fixed"),
        QualityIssue(issue_id="warning_1", status="warning", check_key="pdf.blank", title="Possible blank page"),
        QualityIssue(issue_id="fail_1", status="fail", check_key="docx.font", title="Font mismatch"),
        QualityIssue(
            issue_id="unsupported_1",
            status="unsupported",
            check_key="docx.page_number",
            title="Page number unsupported",
        ),
    ]

    summary = QualitySummary.from_issues(issues)

    assert summary.counts == {"pass": 1, "fixed": 1, "warning": 1, "fail": 1, "unsupported": 1}
    assert summary.remaining_issue_count == 3
    assert summary.all_compliant is False


def test_fix_plan_and_fix_loop_models_preserve_confirmation_and_lineage() -> None:
    action = FixAction(
        action="reapply_profile_formatting",
        target_issue_ids=["issue_margin_top"],
        params={"scope": "document"},
        requires_user_confirmation=True,
    )
    plan = FixPlan(
        fix_plan_id="fp_123",
        report_id="qr_original",
        actions=[action],
        manual_review_issue_ids=["issue_page_number"],
    )
    loop = FixLoopRecord(
        fix_loop_id="fl_123",
        original_report_id="qr_original",
        fix_plan_id="fp_123",
        selected_issue_ids=["issue_margin_top"],
        selected_actions=[action],
        status="pending_confirmation",
        new_job_id=None,
        new_output_file_ids=[],
        updated_report_id=None,
    )

    assert plan.requires_user_confirmation is True
    assert loop.original_report_id == "qr_original"
    assert loop.updated_report_id is None
    assert loop.selected_actions[0].target_issue_ids == ["issue_margin_top"]


def test_quality_and_fix_models_reject_unknown_statuses() -> None:
    with pytest.raises(ValidationError):
        QualityIssue(issue_id="bad", status="ok", check_key="docx.open", title="Bad status")

    with pytest.raises(ValidationError):
        FixLoopRecord(
            fix_loop_id="fl_bad",
            original_report_id="qr_original",
            fix_plan_id="fp_123",
            selected_issue_ids=["issue_margin_top"],
            selected_actions=[],
            status="done",
        )


def test_repository_persists_quality_reports_and_fix_loop_lineage(tmp_path) -> None:
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    issue = QualityIssue(
        issue_id="issue_font",
        status="warning",
        check_key="docx.body.font",
        title="Body font needs review.",
        profile_rule_ref="body.font",
        location="paragraph[2]",
        recommendation="Apply body paragraph style.",
        fixable=True,
    )
    report = QualityReport(
        report_id="qr_123",
        job_id="job_123",
        profile_id="ecnu_thesis",
        profile_version="1.0.0",
        output_file_ids=["file_docx"],
        summary=QualitySummary.from_issues([issue]),
        issues=[issue],
    )
    action = FixAction(action="apply_body_paragraph_style", target_issue_ids=["issue_font"])
    loop = FixLoopRecord(
        fix_loop_id="fl_123",
        original_report_id="qr_123",
        fix_plan_id="fp_123",
        selected_issue_ids=["issue_font"],
        selected_actions=[action],
        status="confirmed",
    )

    repository.add_quality_report(report)
    repository.add_quality_fix_loop(loop)
    loaded_report = JsonMetadataRepository(tmp_path / "metadata.json").get_quality_report("qr_123")
    loaded_loop = JsonMetadataRepository(tmp_path / "metadata.json").get_quality_fix_loop("fl_123")

    assert loaded_report == report
    assert loaded_loop == loop
    assert loaded_report.summary.counts["warning"] == 1

    loaded_loop.status = "completed"
    loaded_loop.new_job_id = "job_fix_123"
    loaded_loop.new_output_file_ids = ["file_fixed_docx"]
    loaded_loop.updated_report_id = "qr_fixed"
    updated_loop = repository.update_quality_fix_loop(loaded_loop)

    reloaded_loop = JsonMetadataRepository(tmp_path / "metadata.json").get_quality_fix_loop("fl_123")
    assert reloaded_loop == updated_loop
    assert reloaded_loop.original_report_id == "qr_123"
    assert reloaded_loop.updated_report_id == "qr_fixed"
    assert updated_loop.updated_at >= loop.updated_at


def test_repository_handles_legacy_metadata_without_quality_collections(tmp_path) -> None:
    metadata_path = tmp_path / "metadata.json"
    metadata_path.write_text(
        '{"files": {}, "jobs": {}, "profiles": {}, "profile_versions": {}, "profile_extractions": {}}',
        encoding="utf-8",
    )
    repository = JsonMetadataRepository(metadata_path)

    assert repository.list_quality_reports() == []
    assert repository.get_quality_report("qr_missing") is None
    assert repository.list_quality_fix_loops() == []
    assert repository.get_quality_fix_loop("fl_missing") is None


def test_quality_report_serializes_issues_by_status_without_hiding_unsupported() -> None:
    issues = [
        QualityIssue(issue_id="issue_pass", status="pass", check_key="docx.open", title="DOCX opens"),
        QualityIssue(issue_id="issue_warning", status="warning", check_key="pdf.blank", title="Blank page warning"),
        QualityIssue(issue_id="issue_fail", status="fail", check_key="docx.font", title="Font mismatch"),
        QualityIssue(
            issue_id="issue_unsupported",
            status="unsupported",
            check_key="docx.page_number",
            title="Page number cannot be judged",
            recommendation="Review page numbering manually.",
        ),
    ]
    report = QualityReport(
        report_id="qr_grouped",
        job_id="job_grouped",
        profile_id="ecnu_thesis",
        profile_version="1.0.0",
        output_file_ids=["file_docx"],
        summary=QualitySummary.from_issues(issues),
        issues=issues,
    )

    payload = report.model_dump(mode="json")

    assert [issue.issue_id for issue in report.issues_by_status["warning"]] == ["issue_warning"]
    assert payload["issues_by_status"]["unsupported"][0]["issue_id"] == "issue_unsupported"
    assert payload["summary"]["remaining_issue_count"] == 3
    assert payload["summary"]["all_compliant"] is False


def test_docx_quality_inspection_passes_profiled_output_and_page_numbers(tmp_path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    source = create_minimal_thesis_docx(tmp_path / "source.docx")
    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)

    issues = inspect_docx_quality(formatted, profile)
    by_key = {issue.check_key: issue for issue in issues}

    assert by_key["docx.page.margins"].status == "pass"
    assert by_key["docx.page.setup"].status == "pass"
    assert by_key["docx.header_footer"].status == "pass"
    assert by_key["docx.body.style"].status == "pass"
    assert by_key["docx.heading.style"].status == "pass"
    assert by_key["docx.table.borders"].status == "pass"
    assert by_key["docx.captions"].status == "pass"
    assert by_key["docx.raw_latex"].status == "pass"
    assert by_key["docx.page_number"].status == "pass"
    assert by_key["docx.ooxml.features"].status == "pass"
    assert by_key["docx.fields.update_policy"].status == "pass"
    assert by_key["docx.toc.fields"].status == "pass"
    assert by_key["docx.notes"].status == "pass"
    assert by_key["docx.visuals.caption_pairing"].status == "pass"


def test_docx_quality_inspection_honors_profile_header_and_page_number_settings(tmp_path) -> None:
    base_profile = load_builtin_profiles()["ecnu_thesis"]
    profile = base_profile.model_copy(
        update={
            "page": base_profile.page.model_copy(update={"size": "Letter", "orientation": "landscape"}),
            "header_footer": base_profile.header_footer.model_copy(
                update={
                    "header_text": "测试模板页眉",
                    "footer_page_number": False,
                    "header_alignment": "right",
                }
            ),
        }
    )
    source = create_minimal_thesis_docx(tmp_path / "source.docx")
    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)

    issues = inspect_docx_quality(formatted, profile)
    by_key = {issue.check_key: issue for issue in issues}

    assert by_key["docx.page.setup"].status == "pass"
    assert by_key["docx.header_footer"].status == "pass"
    assert by_key["docx.page_number"].status == "pass"


def test_docx_quality_inspection_detects_page_setup_and_header_footer_failures(tmp_path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"].model_copy(
        update={
            "header_footer": load_builtin_profiles()["ecnu_thesis"].header_footer.model_copy(
                update={"header_text": "必须出现的页眉", "footer_page_number": True}
            )
        }
    )
    source = create_minimal_thesis_docx(tmp_path / "source.docx")
    document = Document(source)
    document.sections[0].page_width = Cm(21.59)
    document.sections[0].page_height = Cm(27.94)
    broken = tmp_path / "broken.docx"
    document.save(broken)

    issues = inspect_docx_quality(broken, profile)
    by_key = {issue.check_key: issue for issue in issues}

    assert by_key["docx.page.setup"].status == "fail"
    assert by_key["docx.header_footer"].status == "fail"
    assert by_key["docx.header_footer"].fixable is True


def test_docx_quality_inspection_fails_missing_page_numbers(tmp_path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    source = create_minimal_thesis_docx(tmp_path / "source.docx")

    issues = inspect_docx_quality(source, profile)
    by_key = {issue.check_key: issue for issue in issues}

    assert by_key["docx.page_number"].status == "fail"
    assert by_key["docx.page_number"].fixable is True


def test_docx_quality_inspection_detects_margin_and_latex_failures(tmp_path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    source = create_minimal_thesis_docx(tmp_path / "source.docx")
    document = Document(source)
    document.sections[0].top_margin = Cm(1.0)
    document.add_paragraph(r"Inline residue: $E = mc^2$")
    broken = tmp_path / "broken.docx"
    document.save(broken)

    issues = inspect_docx_quality(broken, profile)
    by_key = {issue.check_key: issue for issue in issues}

    assert by_key["docx.page.margins"].status == "fail"
    assert by_key["docx.page.margins"].profile_rule_ref == "page.margins_cm"
    assert by_key["docx.raw_latex"].status == "fail"
    assert by_key["docx.raw_latex"].location == "paragraph[8]"


def test_docx_quality_inspection_detects_body_font_color_mismatch(tmp_path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    profile.body.font.color = "000000"
    source = create_minimal_thesis_docx(tmp_path / "source.docx")
    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)
    broken = Document(formatted)
    broken.paragraphs[1].runs[0].font.color.rgb = RGBColor(255, 0, 0)
    broken_path = tmp_path / "broken.docx"
    broken.save(broken_path)

    issues = inspect_docx_quality(broken_path, profile)
    by_key = {issue.check_key: issue for issue in issues}

    assert by_key["docx.body.style"].status == "warning"
    assert "body run color expected 000000" in (by_key["docx.body.style"].description or "")


def test_docx_caption_inspection_distinguishes_not_applicable_from_missing_captions(tmp_path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    no_visuals = tmp_path / "no-visuals.docx"
    document = Document()
    document.add_heading("第一章 绪论", level=1)
    document.add_paragraph("这是一份没有表格和图片的正文。")
    document.save(no_visuals)

    table_without_caption = tmp_path / "table-without-caption.docx"
    document = Document()
    document.add_heading("第一章 绪论", level=1)
    document.add_paragraph("下表展示示例数据。")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "示例"
    document.save(table_without_caption)

    no_visuals_by_key = {issue.check_key: issue for issue in inspect_docx_quality(no_visuals, profile)}
    table_by_key = {issue.check_key: issue for issue in inspect_docx_quality(table_without_caption, profile)}

    assert no_visuals_by_key["docx.captions"].status == "pass"
    assert table_by_key["docx.captions"].status == "warning"
    assert table_by_key["docx.captions"].fixable is False


def test_docx_quality_inspection_fails_closed_for_complex_ooxml_features(tmp_path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    source = create_minimal_thesis_docx(tmp_path / "source.docx")
    complex_docx = add_ooxml_features(
        source,
        toc_field=True,
        footnote=True,
        anchored_image=True,
        numbering=True,
        update_fields=False,
    )

    issues = inspect_docx_quality(complex_docx, profile)
    by_key = {issue.check_key: issue for issue in issues}
    summary = QualitySummary.from_issues(issues)

    assert by_key["docx.ooxml.features"].status == "pass"
    assert by_key["docx.ooxml.features"].details["footnote_count"] == 1
    assert by_key["docx.ooxml.features"].details["anchored_image_count"] == 1
    assert by_key["docx.fields.update_policy"].status == "warning"
    assert by_key["docx.toc.fields"].status == "warning"
    assert by_key["docx.notes"].status == "unsupported"
    assert by_key["docx.visuals.caption_pairing"].status == "unsupported"
    assert by_key["docx.numbering"].status == "pass"
    assert summary.all_compliant is False


def test_docx_quality_inspection_detects_role_style_misclassification(tmp_path) -> None:
    profile = load_builtin_profiles()["ecnu_thesis"]
    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("1、引言")
    document.add_paragraph("第一段正文已经正确应用正文样式。")
    document.add_paragraph("第二段正文被错误套成参考文献悬挂缩进。")
    document.add_paragraph("表 1 给出一个课程化比较。")
    document.add_paragraph("表 1 RISC-V、Arm 与 x86 对比")
    document.save(source)

    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)
    broken = Document(formatted)
    broken.paragraphs[0].style = broken.styles["Normal"]
    broken.paragraphs[2].paragraph_format.left_indent = Cm(0.74)
    broken.paragraphs[2].paragraph_format.first_line_indent = Cm(-0.74)
    broken.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    broken_path = tmp_path / "broken.docx"
    broken.save(broken_path)

    issues = inspect_docx_quality(broken_path, profile)
    by_key = {issue.check_key: issue for issue in issues}

    role_issue = by_key["docx.role_styles"]
    assert role_issue.status == "warning"
    assert role_issue.fixable is False
    assert "paragraph[1] heading is not using a Word heading style" in role_issue.description
    assert "paragraph[3] body has reference-style hanging indent" in role_issue.description
    assert "paragraph[2] body is centered like a caption or equation" in role_issue.description


def test_pdf_quality_inspection_passes_basic_deliverability(tmp_path) -> None:
    pdf = tmp_path / "deliverable.pdf"
    pdf.write_bytes(
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >> endobj\n"
        b"4 0 obj << /Length 48 >> stream\n"
        b"BT /F1 12 Tf 72 720 Td (Hello thesis text) Tj ET\n"
        b"endstream endobj\n"
        b"%%EOF"
    )

    issues = inspect_pdf_quality(pdf)
    by_key = {issue.check_key: issue for issue in issues}

    assert by_key["pdf.openability"].status == "pass"
    assert by_key["pdf.page_count"].status == "pass"
    assert by_key["pdf.page_count"].details["page_count"] == 1
    assert by_key["pdf.text_extractability"].status == "pass"
    assert by_key["pdf.blank_pages"].status == "pass"


def test_pdf_quality_inspection_flags_unreadable_and_blank_pdf(tmp_path) -> None:
    corrupt = tmp_path / "corrupt.pdf"
    corrupt.write_bytes(b"not a pdf")
    blank = tmp_path / "blank.pdf"
    blank.write_bytes(
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >> endobj\n"
        b"4 0 obj << /Length 0 >> stream\n\nendstream endobj\n"
        b"%%EOF"
    )

    corrupt_by_key = {issue.check_key: issue for issue in inspect_pdf_quality(corrupt)}
    blank_by_key = {issue.check_key: issue for issue in inspect_pdf_quality(blank)}

    assert corrupt_by_key["pdf.openability"].status == "fail"
    assert corrupt_by_key["pdf.page_count"].status == "fail"
    assert blank_by_key["pdf.openability"].status == "pass"
    assert blank_by_key["pdf.text_extractability"].status == "fail"
    assert blank_by_key["pdf.blank_pages"].status == "warning"


def test_quality_report_service_generates_and_persists_status_summary(tmp_path) -> None:
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    storage = LocalFileStorage(tmp_path)
    profile = load_builtin_profiles()["ecnu_thesis"]
    repository.save_profile_version(profile)
    source = create_minimal_thesis_docx(tmp_path / "source.docx")
    formatted = format_docx_with_profile(source, tmp_path / "formatted.docx", profile)
    docx_record = storage.store_generated_file(
        formatted,
        filename="formatted.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(docx_record)
    pdf = tmp_path / "blank.pdf"
    pdf.write_bytes(
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >> endobj\n"
        b"4 0 obj << /Length 0 >> stream\n\nendstream endobj\n"
        b"%%EOF"
    )
    pdf_record = storage.store_generated_file(pdf, filename="formatted.pdf", mime_type="application/pdf")
    repository.add_file(pdf_record)

    report = QualityReportService(repository).create_report(
        profile_id="ecnu_thesis",
        profile_version="1.0.0",
        output_file_ids=[docx_record.file_id, pdf_record.file_id],
        job_id="job_quality",
    )

    assert repository.get_quality_report(report.report_id) == report
    assert report.job_id == "job_quality"
    assert report.summary.counts["pass"] >= 6
    assert report.summary.counts["warning"] >= 1
    assert report.summary.counts["fail"] >= 1
    assert report.summary.counts["unsupported"] == 0
    assert report.summary.all_compliant is False
    assert any(issue.check_key == "pdf.text_extractability" for issue in report.issues_by_status["fail"])


def test_fix_plan_service_explains_issues_and_uses_whitelisted_actions() -> None:
    report = QualityReport(
        report_id="qr_fixable",
        job_id="job_quality",
        profile_id="ecnu_thesis",
        profile_version="1.0.0",
        output_file_ids=["file_docx"],
        summary=QualitySummary.from_issues([]),
        issues=[
            QualityIssue(
                issue_id="issue_body",
                status="warning",
                check_key="docx.body.style",
                title="Body style needs review.",
                description="Line spacing mismatch.",
                profile_rule_ref="body",
                location="paragraph[2]",
                recommendation="Apply body paragraph style.",
                fixable=True,
            ),
            QualityIssue(
                issue_id="issue_page_number",
                status="unsupported",
                check_key="docx.page_number",
                title="Page number cannot be judged.",
                recommendation="Review manually.",
                fixable=False,
            ),
        ],
    )
    report.summary = QualitySummary.from_issues(report.issues)

    plan = FixPlanService().create_fix_plan(report)

    assert isinstance(plan.explanations[0], IssueExplanation)
    assert plan.explanations[0].issue_id == "issue_body"
    assert plan.explanations[0].automatic_repair_allowed is True
    assert plan.actions == [
        FixAction(
            action="apply_body_paragraph_style",
            target_issue_ids=["issue_body"],
            params={"check_key": "docx.body.style", "profile_rule_ref": "body"},
            requires_user_confirmation=True,
        )
    ]
    unsupported = next(item for item in plan.explanations if item.issue_id == "issue_page_number")
    assert unsupported.automatic_repair_allowed is False
    assert "cannot judge" in unsupported.manual_review_guidance.lower()
    assert "issue_page_number" in plan.manual_review_issue_ids


def test_fix_plan_validation_rejects_unsafe_actions_and_missing_targets() -> None:
    safe_plan = FixPlan(
        fix_plan_id="fp_safe",
        report_id="qr_safe",
        actions=[FixAction(action="apply_table_borders", target_issue_ids=["issue_table"])],
    )
    validate_fix_plan(safe_plan, known_issue_ids={"issue_table"})

    unsafe_payloads = [
        {"action": "rewrite_thesis_argument", "target_issue_ids": ["issue_table"]},
        {"action": "edit_formula_content", "target_issue_ids": ["issue_table"]},
        {"action": "edit_reference_content", "target_issue_ids": ["issue_table"]},
        {"action": "apply_table_borders", "target_issue_ids": []},
        {"action": "apply_table_borders", "target_issue_ids": ["missing_issue"]},
    ]
    for payload in unsafe_payloads:
        with pytest.raises(FixPlanSafetyError):
            validate_fix_plan(
                {
                    "fix_plan_id": "fp_unsafe",
                    "report_id": "qr_safe",
                    "actions": [payload],
                },
                known_issue_ids={"issue_table"},
            )


def test_fix_plan_service_is_deterministic_without_llm_configuration() -> None:
    report = QualityReport(
        report_id="qr_no_llm",
        profile_id="ecnu_thesis",
        profile_version="1.0.0",
        output_file_ids=["file_docx"],
        summary=QualitySummary.from_issues([]),
        issues=[
            QualityIssue(
                issue_id="issue_heading",
                status="fail",
                check_key="docx.heading.style",
                title="Heading style mismatch.",
                description="Heading alignment is not centered.",
                profile_rule_ref="headings[1]",
                fixable=True,
            )
        ],
    )
    report.summary = QualitySummary.from_issues(report.issues)

    first = FixPlanService(llm_configured=False).create_fix_plan(report)
    second = FixPlanService(llm_configured=False).create_fix_plan(report)

    assert [action.model_dump() for action in first.actions] == [action.model_dump() for action in second.actions]
    assert first.explanation == "Deterministic fallback fix plan generated from quality issue metadata."


def test_fix_loop_execute_rewrites_docx_and_updates_lineage(tmp_path) -> None:
    repository = JsonMetadataRepository(tmp_path / "metadata.json")
    storage = LocalFileStorage(tmp_path)
    profile = load_builtin_profiles()["ecnu_thesis"]
    repository.save_profile_version(profile)
    source = create_minimal_thesis_docx(tmp_path / "source.docx")
    broken_doc = Document(source)
    broken_doc.sections[0].top_margin = Cm(1.0)
    broken_path = tmp_path / "broken-output.docx"
    broken_doc.save(broken_path)
    broken_record = storage.store_generated_file(
        broken_path,
        filename="broken-output.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repository.add_file(broken_record)
    report = QualityReportService(repository).create_report(
        profile_id="ecnu_thesis",
        profile_version="1.0.0",
        output_file_ids=[broken_record.file_id],
        job_id="job_original",
    )
    assert next(issue for issue in report.issues if issue.check_key == "docx.page.margins").status == "fail"
    client = TestClient(create_app(Settings(FILE_STORAGE_ROOT=tmp_path, SOFFICE_BIN=None)))
    plan_response = client.post(f"/api/quality-reports/{report.report_id}/fix-plan")
    assert plan_response.status_code == 201
    selected_issue_id = next(
        action["target_issue_ids"][0]
        for action in plan_response.json()["actions"]
        if action["action"] == "reapply_profile_formatting"
    )
    loop_response = client.post(
        f"/api/quality-reports/{report.report_id}/fix-loops",
        json={"fix_plan_id": plan_response.json()["fix_plan_id"], "selected_issue_ids": [selected_issue_id]},
    )
    assert loop_response.status_code == 201

    executed = client.post(
        f"/api/quality-reports/{report.report_id}/fix-loops/{loop_response.json()['fix_loop_id']}/execute"
    )

    assert executed.status_code == 200
    loop = executed.json()
    assert loop["status"] == "completed"
    assert loop["new_job_id"].startswith("job_fix_")
    assert len(loop["new_output_file_ids"]) == 1
    assert loop["updated_report_id"].startswith("qr_")
    updated_report = repository.get_quality_report(loop["updated_report_id"])
    assert updated_report is not None
    assert updated_report.job_id == loop["new_job_id"]
    assert next(issue for issue in updated_report.issues if issue.check_key == "docx.page.margins").status == "pass"
    fixed_file = repository.get_file(loop["new_output_file_ids"][0])
    assert fixed_file is not None
    assert fixed_file.filename.endswith("-fixed.docx")
