from __future__ import annotations

from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

ECNU_RULE_TEXT = (
    "毕业论文（设计）一律用A4纸张电脑打印。左侧装订。\n"
    "“纸型” ──主要选用“A4”，“纵向”，个别页面可以采用“A4”，“横向”。\n"
    "“文档网格” ──一律使用“无网格”。\n"
    "“页边距” ──上：2.5cm，下：2.0 cm，左：3.0cm，右：2.5 cm。装订线位置居左。\n"
    "论文（设计）题目居中，每段落首行缩进2字符。\n"
    "“行距”一律为1.5倍。\n"
    "外文字体：一律为Times New Roman\n"
    "中外文均是五号，中文使用宋体。\n"
    "正文一般用宋体小四号字打印。中文题名用黑体小三打印，外文题名用小三打印。"
    "文章中的各段标题用黑体、小四号字打印，并且前后要一致。\n"
    "（理科）中文各层次系统为：第一层：1、 2、 3、 ……；"
    "第二层：1.1、2.1、3.1……；第三层：1.1.1、2.1.1、3.1.1……。\n"
    "每页要插入阿拉伯数字页码，置于页面底端居中。\n"
    "表名放置在表格正上方，中外文对照；图名放置在图件的正下方，中外文对照。"
    "表格一览表采用“三线表”形式，公式应独立成行居中斜体排版。\n"
    "参考文献具体格式如下。"
)


def create_minimal_thesis_docx(path: Path) -> Path:
    document = Document()
    document.add_heading("第一章 绪论", level=1)
    document.add_paragraph("这是一段正文内容，用于验证格式化后文本不会丢失。")
    document.add_paragraph("Table 1 Sample table")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "Value A"
    table.cell(1, 1).text = "Value B"
    document.add_paragraph("图 1 示例图片标题")
    document.add_paragraph("E = mc^2")
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] Author. Title.")
    document.save(path)
    return path


def create_ecnu_rule_docx(path: Path) -> Path:
    document = Document()
    for line in ECNU_RULE_TEXT.splitlines():
        document.add_paragraph(line)
    document.save(path)
    return path


def read_docx_text(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(path).paragraphs]


def add_ooxml_features(
    path: Path,
    *,
    toc_field: bool = False,
    footnote: bool = False,
    anchored_image: bool = False,
    numbering: bool = False,
    update_fields: bool = False,
) -> Path:
    with ZipFile(path) as source_package:
        document_root = ET.fromstring(source_package.read("word/document.xml"))
        body = document_root.find(f"{{{W_NS}}}body")
        if body is None:
            raise ValueError("document body missing")
        first_paragraph = body.find(f"{{{W_NS}}}p")
        if first_paragraph is None:
            first_paragraph = ET.Element(f"{{{W_NS}}}p")
            body.insert(0, first_paragraph)
        if toc_field:
            first_paragraph.insert(0, _run_with_instr_text(' TOC \\o "1-3" \\h \\z \\u '))
        if footnote:
            first_paragraph.append(_run_with_footnote_reference("2"))
        if anchored_image:
            first_paragraph.append(_run_with_anchor())
        if numbering:
            ppr = first_paragraph.find(f"{{{W_NS}}}pPr")
            if ppr is None:
                ppr = ET.Element(f"{{{W_NS}}}pPr")
                first_paragraph.insert(0, ppr)
            num_pr = ET.SubElement(ppr, f"{{{W_NS}}}numPr")
            ilvl = ET.SubElement(num_pr, f"{{{W_NS}}}ilvl")
            ilvl.set(f"{{{W_NS}}}val", "0")
            num_id = ET.SubElement(num_pr, f"{{{W_NS}}}numId")
            num_id.set(f"{{{W_NS}}}val", "1")

        settings_root = (
            ET.fromstring(source_package.read("word/settings.xml"))
            if "word/settings.xml" in source_package.namelist()
            else ET.Element(f"{{{W_NS}}}settings")
        )
        if update_fields:
            update = settings_root.find(f"{{{W_NS}}}updateFields")
            if update is None:
                update = ET.SubElement(settings_root, f"{{{W_NS}}}updateFields")
            update.set(f"{{{W_NS}}}val", "true")

        with NamedTemporaryFile(delete=False, suffix=".docx", dir=path.parent) as tmp_file:
            tmp_path = Path(tmp_file.name)
        try:
            with ZipFile(tmp_path, "w") as target_package:
                copied: set[str] = set()
                for item in source_package.infolist():
                    if item.filename in copied or item.filename in {"word/document.xml", "word/settings.xml", "word/footnotes.xml"}:
                        continue
                    copied.add(item.filename)
                    target_package.writestr(item, source_package.read(item.filename))
                target_package.writestr("word/document.xml", _xml_bytes(document_root))
                target_package.writestr("word/settings.xml", _xml_bytes(settings_root))
                if footnote:
                    target_package.writestr("word/footnotes.xml", _footnotes_xml())
            tmp_path.replace(path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink()
    return path


def _run_with_instr_text(text: str) -> ET.Element:
    run = ET.Element(f"{{{W_NS}}}r")
    instr = ET.SubElement(run, f"{{{W_NS}}}instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = text
    return run


def _run_with_footnote_reference(note_id: str) -> ET.Element:
    run = ET.Element(f"{{{W_NS}}}r")
    ref = ET.SubElement(run, f"{{{W_NS}}}footnoteReference")
    ref.set(f"{{{W_NS}}}id", note_id)
    return run


def _run_with_anchor() -> ET.Element:
    run = ET.Element(f"{{{W_NS}}}r")
    drawing = ET.SubElement(run, f"{{{W_NS}}}drawing")
    ET.SubElement(drawing, f"{{{WP_NS}}}anchor")
    return run


def _footnotes_xml() -> bytes:
    root = ET.Element(f"{{{W_NS}}}footnotes")
    note = ET.SubElement(root, f"{{{W_NS}}}footnote")
    note.set(f"{{{W_NS}}}id", "2")
    paragraph = ET.SubElement(note, f"{{{W_NS}}}p")
    run = ET.SubElement(paragraph, f"{{{W_NS}}}r")
    text = ET.SubElement(run, f"{{{W_NS}}}t")
    text.text = "脚注内容"
    return _xml_bytes(root)


def _xml_bytes(root: ET.Element) -> bytes:
    ET.register_namespace("w", W_NS)
    ET.register_namespace("wp", WP_NS)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)
